"""
Editor letter generator for creating a response letter to the editor.
"""

import logging
import os
import datetime
from typing import Dict, Any, List, Optional, Tuple
from docx import Document

from src.core.context import RevisionContext
from src.utils.document_processor import DocumentProcessor


class EditorLetterGenerator:
    """
    Creates a response letter to the editor addressing reviewer comments.
    
    This class is responsible for generating a letter to the editor that responds
    to each reviewer comment and explains how the paper was revised to address
    the concerns.
    """
    
    def __init__(self, context: RevisionContext):
        """
        Initialize the editor letter generator.
        
        Args:
            context: The shared revision context
        """
        self.context = context
        self.logger = context.logger or logging.getLogger(__name__)
    
    def create_editor_letter(
        self,
        reviewer_comments: List[Dict[str, Any]],
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: Optional[str] = None
    ) -> str:
        """
        Create letter to the editor.
        
        Args:
            reviewer_comments: Analysis of reviewer comments
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved (optional)
            
        Returns:
            Path to the created document
        """
        self.logger.info("Creating letter to editor")
        
        # Determine output path if not provided
        if output_path is None:
            output_path = self.context.get_output_path("editor_letter.docx")
        
        # Process changes to create responses to reviewers
        reviewer_responses = []
        
        for i, reviewer in enumerate(reviewer_comments, 1):
            # Group changes by reason to identify which ones address this reviewer's comments
            changes_by_reason = {}
            for old_text, new_text, reason, _ in changes:
                if f"Reviewer {i}" in reason or f"reviewer {i}" in reason.lower():
                    if reason not in changes_by_reason:
                        changes_by_reason[reason] = []
                    changes_by_reason[reason].append((old_text, new_text))
            
            # Extract main comments from the reviewer
            main_concerns = reviewer.get("main_concerns", [])
            required_changes = reviewer.get("required_changes", [])
            
            # Create a list of comment-response pairs
            comments = []
            
            # Add responses to main concerns
            for j, concern in enumerate(main_concerns, 1):
                response_text = "We appreciate this concern and have addressed it in our revision."
                changes_text = "No specific changes were made for this comment."
                
                # Find changes related to this concern
                for reason, change_list in changes_by_reason.items():
                    if any(keyword in concern.lower() for keyword in reason.lower().split()):
                        changes_text = f"We made {len(change_list)} changes to address this concern, including: "
                        changes_text += ", ".join([f"replacing '{old[:20]}...' with '{new[:20]}...'" 
                                                 for old, new in change_list[:2]])
                        if len(change_list) > 2:
                            changes_text += f", and {len(change_list) - 2} more changes."
                
                comments.append({
                    "comment": concern,
                    "response": response_text,
                    "changes": changes_text
                })
            
            # Add responses to required changes
            for j, required in enumerate(required_changes, 1):
                response_text = "We have implemented this required change in our revision."
                changes_text = "No specific changes were made for this comment."
                
                # Find changes related to this requirement
                for reason, change_list in changes_by_reason.items():
                    if any(keyword in required.lower() for keyword in reason.lower().split()):
                        changes_text = f"We made {len(change_list)} changes to address this requirement, including: "
                        changes_text += ", ".join([f"replacing '{old[:20]}...' with '{new[:20]}...'" 
                                                 for old, new in change_list[:2]])
                        if len(change_list) > 2:
                            changes_text += f", and {len(change_list) - 2} more changes."
                
                comments.append({
                    "comment": required,
                    "response": response_text,
                    "changes": changes_text
                })
            
            reviewer_responses.append({
                "reviewer": i,
                "comments": comments
            })
        
        # Create the editor letter using DocumentProcessor if original docx available
        if os.path.exists(self.context.original_docx_path):
            doc_processor = DocumentProcessor(self.context.original_docx_path)
            
            # Get process summary for inclusion in the letter
            process_summary = None
            if hasattr(self.context, 'workflow_db') and self.context.run_id:
                try:
                    process_summary = self.context.workflow_db.get_review_process_summary(self.context.run_id)
                except Exception as e:
                    self.logger.warning(f"Error getting review process summary: {e}")
            
            editor_letter_path = doc_processor.create_editor_letter(reviewer_responses, output_path, process_summary)
            
        else:
            # Create manually if original docx not available
            doc = Document()
            
            # Add header information
            doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("To: The Editor")
            doc.add_paragraph("Subject: Revised manuscript submission")
            doc.add_paragraph()
            
            # Add salutation
            doc.add_paragraph("Dear Editor,")
            
            # Add introduction
            doc.add_paragraph(
                "Thank you for the opportunity to revise our manuscript. We have carefully addressed all the comments " +
                "provided by the reviewers and made the necessary changes to improve the quality of our paper. " +
                "We believe that the revised version addresses all the concerns raised and significantly improves the manuscript."
            )
            
            # Add reviewer responses
            doc.add_heading('Responses to Reviewer Comments', 1)
            
            for i, response in enumerate(reviewer_responses, 1):
                doc.add_heading(f"Reviewer {i}", 2)
                
                for j, comment in enumerate(response['comments'], 1):
                    p = doc.add_paragraph()
                    p.add_run(f"Comment {j}: ").bold = True
                    p.add_run(comment['comment'])
                    
                    p = doc.add_paragraph()
                    p.add_run("Response: ").bold = True
                    p.add_run(comment['response'])
                    
                    p = doc.add_paragraph()
                    p.add_run("Changes made: ").bold = True
                    p.add_run(comment['changes'])
                    
                    doc.add_paragraph()
            
            # Add closing
            doc.add_paragraph(
                "We hope that the revised manuscript now meets the standards for publication. " +
                "We look forward to your feedback and are available to address any additional questions or concerns."
            )
            
            # Add revision process disclosure
            doc.add_heading('REVISION PROCESS DISCLOSURE', 1)
            
            # Add information about the LLM-assisted revision process
            doc.add_paragraph(
                "This revision was created with the assistance of AI-based language models " +
                f"({self.context.provider.capitalize()} {self.context.model_name}) to improve the quality " +
                "and clarity of the manuscript. All changes were carefully reviewed by the authors to ensure " +
                "accuracy and alignment with the scientific content."
            )
            
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph("The Authors")
            
            # Save the document
            doc.save(output_path)
            editor_letter_path = output_path
        
        self.context.process_statistics["files_created"] = self.context.process_statistics.get("files_created", 0) + 1
        
        self.logger.info(f"Editor letter created at {editor_letter_path}")
        return editor_letter_path