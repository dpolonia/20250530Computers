"""
Assessment document generator for creating a document assessing the revision impact.
"""

import logging
import os
import json
from typing import Dict, Any, List, Optional, Tuple
from docx import Document

from src.core.context import RevisionContext
from src.core.json_utils import parse_json_safely
from src.core.constants import SYSTEM_PROMPTS


class AssessmentGenerator:
    """
    Creates an assessment document evaluating the impact of revisions.
    
    This class is responsible for generating a document that assesses the impact
    of the proposed changes on the paper's quality, how well they address
    reviewer concerns, and any remaining issues that might need attention.
    """
    
    def __init__(self, context: RevisionContext):
        """
        Initialize the assessment generator.
        
        Args:
            context: The shared revision context
        """
        self.context = context
        self.logger = context.logger or logging.getLogger(__name__)
    
    def create_assessment(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        paper_analysis: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:
        """
        Create assessment document.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            paper_analysis: Analysis of the original paper
            output_path: Path where the document should be saved (optional)
            
        Returns:
            Path to the created document
        """
        self.logger.info("Creating assessment document")
        
        # Determine output path if not provided
        if output_path is None:
            output_path = self.context.get_output_path("assessment.docx")
        
        # Create a new document
        doc = Document()
        
        # Add heading
        doc.add_heading('Revision Assessment', 0)
        
        # Add introduction
        doc.add_paragraph(
            'This document assesses the changes made to the paper in response to reviewer comments. ' +
            'It evaluates the impact of these changes and identifies any issues that may need to be addressed manually.'
        )
        
        # Summarize changes
        doc.add_heading('Summary of Changes', 1)
        
        changes_by_type = {}
        for _, _, reason, _ in changes:
            change_type = reason.split(' ')[0] if reason else 'Other'
            if change_type not in changes_by_type:
                changes_by_type[change_type] = 0
            changes_by_type[change_type] += 1
        
        p = doc.add_paragraph('Changes by type:')
        for change_type, count in changes_by_type.items():
            p = doc.add_paragraph(f'{change_type}: {count} changes', style='List Bullet')
        
        # Generate impact assessment
        doc.add_heading('Impact Assessment', 1)
        
        # Use LLM to generate impact assessment
        paper_summary = {
            "title": paper_analysis.get("title", "Unknown Title"),
            "objectives": paper_analysis.get("objectives", "Unknown Objectives"),
            "methodology": paper_analysis.get("methodology", "Unknown Methodology"),
            "findings": paper_analysis.get("findings", "Unknown Findings")
        }
        
        changes_summary = []
        for old_text, new_text, reason, _ in changes[:5]:  # Use first 5 changes as examples
            summary = {
                "reason": reason,
                "old_text_preview": old_text[:50] + "..." if len(old_text) > 50 else old_text,
                "new_text_preview": new_text[:50] + "..." if len(new_text) > 50 else new_text
            }
            changes_summary.append(summary)
        
        prompt = f"""
        I'm assessing the impact of revisions made to a scientific paper.
        
        Paper Summary:
        {json.dumps(paper_summary, indent=2)}
        
        Sample of Changes Made:
        {json.dumps(changes_summary, indent=2)}
        
        Total changes: {len(changes)}
        
        Provide an assessment of:
        1. Overall impact of the changes on the paper's quality
        2. How well the changes address likely reviewer concerns
        3. Potential remaining issues that might need manual attention
        4. Specific areas where the paper has been strengthened
        5. Recommendations for any final manual adjustments
        
        IMPORTANT: Format the response as a valid JSON object with these five sections. Do not include any explanatory text before or after the JSON. The response should begin with '{{' and end with '}}'.
        
        Example format:
        {{
          "overall_impact": "The changes have significantly improved the paper's clarity and scientific rigor...",
          "reviewer_concerns": "The changes effectively address reviewer concerns about methodology and results...",
          "remaining_issues": ["Minor formatting inconsistencies", "Some references need manual verification"],
          "strengthened_areas": ["Methodology section", "Results interpretation", "Literature review"],
          "recommendations": ["Perform a final proofreading", "Check reference formatting consistency"]
        }}
        """
        
        # Get assessment from LLM
        assessment_json = self._get_completion(
            prompt=prompt,
            system_prompt=SYSTEM_PROMPTS["assessment"],
            max_tokens=2000 if self.context.optimize_costs else 3000
        )
        
        try:
            # Parse the LLM response
            assessment = parse_json_safely(assessment_json)
            
            # Add overall impact
            doc.add_heading('Overall Impact', 2)
            doc.add_paragraph(assessment.get("overall_impact", "No assessment available."))
            
            # Add reviewer concerns addressed
            doc.add_heading('Reviewer Concerns Addressed', 2)
            doc.add_paragraph(assessment.get("reviewer_concerns", "No assessment available."))
            
            # Add remaining issues
            doc.add_heading('Remaining Issues', 2)
            remaining_issues = assessment.get("remaining_issues", "No issues identified.")
            if isinstance(remaining_issues, list):
                for issue in remaining_issues:
                    doc.add_paragraph(issue, style='List Bullet')
            else:
                doc.add_paragraph(remaining_issues)
            
            # Add strengthened areas
            doc.add_heading('Strengthened Areas', 2)
            strengthened_areas = assessment.get("strengthened_areas", "No areas identified.")
            if isinstance(strengthened_areas, list):
                for area in strengthened_areas:
                    doc.add_paragraph(area, style='List Bullet')
            else:
                doc.add_paragraph(strengthened_areas)
            
            # Add recommendations
            doc.add_heading('Recommendations', 2)
            recommendations = assessment.get("recommendations", "No recommendations.")
            if isinstance(recommendations, list):
                for rec in recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
            else:
                doc.add_paragraph(recommendations)
            
        except ValueError as e:
            # Fallback if LLM didn't return valid JSON
            self.logger.warning(f"LLM didn't return valid JSON for assessment. Error: {e}. Using basic assessment.")
            
            doc.add_heading('Overall Impact', 2)
            doc.add_paragraph('The changes have addressed many of the issues identified in the review process.')
            
            doc.add_heading('Remaining Tasks', 2)
            doc.add_paragraph('The following tasks should be completed manually:')
            doc.add_paragraph('1. Review all changes for accuracy and consistency.', style='List Bullet')
            doc.add_paragraph('2. Check references for proper formatting.', style='List Bullet')
            doc.add_paragraph('3. Ensure all reviewer comments have been addressed.', style='List Bullet')
        
        # Add conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(
            'This assessment provides an overview of the changes made and their impact. ' +
            'Authors should review the revised paper carefully before submission.'
        )
        
        # Save the document
        doc.save(output_path)
        self.context.process_statistics["files_created"] = self.context.process_statistics.get("files_created", 0) + 1
        
        self.logger.info(f"Assessment document created at {output_path}")
        return output_path
    
    def _get_completion(self, prompt: str, system_prompt: str, max_tokens: int) -> str:
        """
        Get a completion from the LLM with appropriate error handling.
        
        Args:
            prompt: The prompt to send to the LLM
            system_prompt: The system prompt to use
            max_tokens: Maximum number of tokens to generate
            
        Returns:
            The LLM response text
        """
        try:
            # Ensure LLM client is initialized
            if not self.context.llm_client:
                self.context.setup_llm_client()
                
            # Get completion from LLM
            response = self.context.llm_client.get_completion(
                prompt=prompt,
                system_prompt=system_prompt,
                max_tokens=max_tokens
            )
            
            # Update budget
            tokens_used = self.context.llm_client.total_tokens_used
            cost = self.context.llm_client.total_cost
            self.context.update_budget(tokens_used, cost)
            
            return response
            
        except Exception as e:
            self.logger.error(f"Error getting completion: {e}")
            raise