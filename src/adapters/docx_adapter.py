"""
DOCX adapter implementation.

This module implements the DOCX adapter interface, providing functionality for
interacting with DOCX files.
"""

import logging
import os
import datetime
from typing import Dict, Any, List, Optional, Tuple, BinaryIO

from src.core.context import RevisionContext
from src.adapters.interfaces import DocxAdapterInterface


class DocxAdapter(DocxAdapterInterface):
    """
    Adapter for DOCX file operations.
    
    This adapter is responsible for creating and manipulating DOCX files,
    including changes documents, tracked changes, and editor letters.
    """
    
    def __init__(self, context: RevisionContext, existing_doc: Optional[str] = None):
        """
        Initialize the DOCX adapter.
        
        Args:
            context: The shared revision context
            existing_doc: Path to an existing DOCX file to modify (optional)
        """
        self.context = context
        self.logger = context.logger or logging.getLogger(__name__)
        self.existing_doc = existing_doc
        self._document = None
    
    def read(self, file_path: str) -> str:
        """
        Read a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text from the DOCX
        """
        self.logger.info(f"Reading {file_path}")
        
        try:
            # Using python-docx for text extraction
            import docx
            
            doc = docx.Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text.append(para.text)
            
            return '\n'.join(text)
            
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot read DOCX file.")
            return ""
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {e}")
            return ""
    
    def write(self, content: str, file_path: str) -> str:
        """
        Write content to a DOCX file.
        
        Args:
            content: Text content to write
            file_path: Path where the file should be saved
            
        Returns:
            Path to the written file
        """
        self.logger.info(f"Writing DOCX file to {file_path}")
        
        try:
            # Using python-docx for DOCX creation
            import docx
            
            doc = docx.Document()
            
            # Process markdown-like headings (e.g., # Heading)
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Check for headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                # Check for bullet points
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                # Check for numbered lists
                elif line.startswith('1. ') or line.startswith('1) '):
                    doc.add_paragraph(line[3:], style='List Number')
                # Regular paragraph
                elif line:
                    doc.add_paragraph(line)
                
                i += 1
            
            # Save the document
            doc.save(file_path)
            
            return file_path
            
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot write DOCX file.")
            return ""
        except Exception as e:
            self.logger.error(f"Error writing DOCX file: {e}")
            return ""
    
    def create_changes_document(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: str
    ) -> str:
        """
        Create a document detailing changes.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self.logger.info(f"Creating changes document at {output_path}")
        
        try:
            # Using python-docx for DOCX creation
            import docx
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_COLOR_INDEX
            
            doc = docx.Document()
            
            # Add title
            doc.add_heading('Document Changes', 0)
            
            # Add timestamp
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Generated on: {timestamp}")
            
            # Add introduction
            doc.add_paragraph(f"This document details {len(changes)} changes made to the paper.")
            
            # Group changes by section
            section_changes = {}
            for old_text, new_text, reason, line_number in changes:
                # Try to determine section from reason
                section = "General"
                if "section:" in reason.lower():
                    section_match = re.search(r'section:\s*([^,\.]+)', reason, re.IGNORECASE)
                    if section_match:
                        section = section_match.group(1).strip()
                
                if section not in section_changes:
                    section_changes[section] = []
                section_changes[section].append((old_text, new_text, reason, line_number))
            
            # Add changes by section
            for section, section_changes_list in section_changes.items():
                doc.add_heading(f"Changes in {section} Section", level=1)
                
                for i, (old_text, new_text, reason, line_number) in enumerate(section_changes_list, 1):
                    # Add heading for each change
                    change_heading = f"Change {i}"
                    if line_number:
                        change_heading += f" (Line {line_number})"
                    doc.add_heading(change_heading, level=2)
                    
                    # Add reason
                    reason_para = doc.add_paragraph("Reason: ")
                    reason_para.add_run(reason).italic = True
                    
                    # Add before/after
                    if old_text:
                        doc.add_paragraph("Before:", style='Intense Emphasis')
                        before_para = doc.add_paragraph()
                        before_para.add_run(old_text).font.color.rgb = RGBColor(192, 0, 0)  # Red text
                    
                    if new_text:
                        doc.add_paragraph("After:", style='Intense Emphasis')
                        after_para = doc.add_paragraph()
                        after_para.add_run(new_text).font.color.rgb = RGBColor(0, 128, 0)  # Green text
                    
                    # Add separator
                    doc.add_paragraph('---')
            
            # Save the document
            doc.save(output_path)
            
            return output_path
            
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot create changes document.")
            return self._fallback_create_changes_document(changes, output_path)
        except Exception as e:
            self.logger.error(f"Error creating changes document: {e}")
            return ""
    
    def _fallback_create_changes_document(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: str
    ) -> str:
        """
        Fallback method to create a changes document as plain text.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        try:
            # Create text file instead
            txt_path = os.path.splitext(output_path)[0] + ".txt"
            
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write("DOCUMENT CHANGES\n")
                f.write("===============\n\n")
                
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                f.write(f"Generated on: {timestamp}\n\n")
                
                f.write(f"This document details {len(changes)} changes made to the paper.\n\n")
                
                # Group changes by section (simplified)
                for i, (old_text, new_text, reason, line_number) in enumerate(changes, 1):
                    f.write(f"CHANGE {i}")
                    if line_number:
                        f.write(f" (Line {line_number})")
                    f.write("\n")
                    f.write("-" * 40 + "\n\n")
                    
                    f.write(f"Reason: {reason}\n\n")
                    
                    if old_text:
                        f.write("Before:\n")
                        f.write(f"{old_text}\n\n")
                    
                    if new_text:
                        f.write("After:\n")
                        f.write(f"{new_text}\n\n")
                    
                    f.write("-" * 40 + "\n\n")
            
            self.logger.info(f"Created plain text changes document at {txt_path}")
            return txt_path
            
        except Exception as e:
            self.logger.error(f"Error creating fallback changes document: {e}")
            return ""
    
    def add_tracked_change(
        self,
        old_text: str,
        new_text: str,
        reason: str
    ) -> bool:
        """
        Add a tracked change to the document.
        
        Args:
            old_text: Text to be replaced
            new_text: Text to replace with
            reason: Reason for the change
            
        Returns:
            True if the change was added, False otherwise
        """
        self.logger.info("Adding tracked change")
        
        if not self.existing_doc:
            self.logger.error("No existing document to modify")
            return False
        
        try:
            # Note: True track changes require COM/Windows integration
            # This is a simplified version that simulates tracked changes visually
            import docx
            from docx.shared import RGBColor
            
            # Lazy load document
            if self._document is None:
                self._document = docx.Document(self.existing_doc)
            
            # Find the text to replace
            found = False
            for para in self._document.paragraphs:
                if old_text in para.text:
                    # Replace the text with visual marking
                    new_para_text = para.text.replace(old_text, new_text)
                    para.clear()
                    
                    # Add a comment about the change
                    comment_run = para.add_run(f"[CHANGE: {reason}] ")
                    comment_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                    comment_run.italic = True
                    
                    # Add the new text
                    para.add_run(new_para_text)
                    
                    found = True
                    break
            
            return found
            
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot add tracked change.")
            return False
        except Exception as e:
            self.logger.error(f"Error adding tracked change: {e}")
            return False
    
    def create_editor_letter(
        self,
        reviewer_responses: List[Dict[str, Any]],
        output_path: str,
        process_summary: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Create a letter to the editor.
        
        Args:
            reviewer_responses: Reviewer responses data
            output_path: Path where the document should be saved
            process_summary: Optional process summary data
            
        Returns:
            Path to the created document
        """
        self.logger.info(f"Creating editor letter at {output_path}")
        
        try:
            # Using python-docx for DOCX creation
            import docx
            
            doc = docx.Document()
            
            # Add title
            doc.add_heading('Response to Reviewers', 0)
            
            # Add introduction
            date = datetime.datetime.now().strftime("%B %d, %Y")
            paper_title = process_summary.get("paper_title", "the paper") if process_summary else "the paper"
            
            intro = doc.add_paragraph()
            intro.add_run(f"Date: {date}\n\n")
            intro.add_run(f"Dear Editor,\n\n")
            intro.add_run(
                f"Thank you for the opportunity to revise {paper_title}. "
                f"We have carefully addressed all the reviewers' comments and made the necessary changes "
                f"to improve the manuscript. Below, we provide detailed responses to each reviewer's comments.\n\n"
            )
            
            # Add responses to each reviewer
            for reviewer_data in reviewer_responses:
                reviewer_num = reviewer_data.get("reviewer_number", "")
                assessment = reviewer_data.get("overall_assessment", "")
                
                doc.add_heading(f"Response to Reviewer {reviewer_num}", level=1)
                
                # Add overview
                overview = doc.add_paragraph()
                overview.add_run(
                    f"The reviewer provided a {'positive' if assessment == 'positive' else 'constructive'} assessment. "
                    f"We have addressed all their concerns and suggestions as detailed below.\n\n"
                )
                
                # Add responses to main concerns
                main_concerns = reviewer_data.get("main_concerns", [])
                if main_concerns:
                    doc.add_heading("Main Concerns", level=2)
                    for i, concern in enumerate(main_concerns, 1):
                        # Find changes made to address this concern
                        related_changes = []
                        for reason, section in reviewer_data.get("changes_made", []):
                            if any(word in reason.lower() and word in concern.lower() 
                                   for word in ["method", "result", "discussion", "reference", "clarity"]):
                                related_changes.append(f"{reason} (in {section} section)")
                        
                        response = doc.add_paragraph()
                        response.add_run(f"Concern {i}: ").bold = True
                        response.add_run(f"{concern}\n")
                        response.add_run(f"Response: ").italic = True
                        
                        if related_changes:
                            response.add_run(
                                f"We thank the reviewer for this important point. We have addressed this concern by: "
                                f"{'; '.join(related_changes)}.\n\n"
                            )
                        else:
                            response.add_run(
                                f"We thank the reviewer for this important point. We have carefully considered "
                                f"this concern and made appropriate revisions to the manuscript.\n\n"
                            )
                
                # Add responses to comments by type
                comments_by_type = reviewer_data.get("comments_by_type", {})
                for comment_type, comments in comments_by_type.items():
                    if comments:
                        doc.add_heading(f"Comments on {comment_type.capitalize()}", level=2)
                        for i, comment in enumerate(comments, 1):
                            response = doc.add_paragraph()
                            response.add_run(f"Comment {i}: ").bold = True
                            response.add_run(f"{comment}\n")
                            response.add_run(f"Response: ").italic = True
                            response.add_run(
                                f"We appreciate this comment and have revised the {comment_type} accordingly.\n\n"
                            )
            
            # Add conclusion
            conclusion = doc.add_paragraph()
            conclusion.add_run("\nConclusion\n\n").bold = True
            conclusion.add_run(
                f"We believe that these revisions have significantly improved the quality of our manuscript. "
                f"We have addressed all the reviewers' concerns and suggestions, and we hope that the revised "
                f"version meets the standards for publication in your journal.\n\n"
                f"Thank you for your consideration.\n\n"
                f"Sincerely,\n"
                f"The Authors"
            )
            
            # Save the document
            doc.save(output_path)
            
            return output_path
            
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot create editor letter.")
            return self._fallback_create_editor_letter(reviewer_responses, output_path, process_summary)
        except Exception as e:
            self.logger.error(f"Error creating editor letter: {e}")
            return ""
    
    def _fallback_create_editor_letter(
        self,
        reviewer_responses: List[Dict[str, Any]],
        output_path: str,
        process_summary: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Fallback method to create an editor letter as plain text.
        
        Args:
            reviewer_responses: Reviewer responses data
            output_path: Path where the document should be saved
            process_summary: Optional process summary data
            
        Returns:
            Path to the created document
        """
        try:
            # Create text file instead
            txt_path = os.path.splitext(output_path)[0] + ".txt"
            
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write("RESPONSE TO REVIEWERS\n")
                f.write("=====================\n\n")
                
                # Add introduction
                date = datetime.datetime.now().strftime("%B %d, %Y")
                paper_title = process_summary.get("paper_title", "the paper") if process_summary else "the paper"
                
                f.write(f"Date: {date}\n\n")
                f.write(f"Dear Editor,\n\n")
                f.write(
                    f"Thank you for the opportunity to revise {paper_title}. "
                    f"We have carefully addressed all the reviewers' comments and made the necessary changes "
                    f"to improve the manuscript. Below, we provide detailed responses to each reviewer's comments.\n\n"
                )
                
                # Add responses to each reviewer
                for reviewer_data in reviewer_responses:
                    reviewer_num = reviewer_data.get("reviewer_number", "")
                    assessment = reviewer_data.get("overall_assessment", "")
                    
                    f.write(f"RESPONSE TO REVIEWER {reviewer_num}\n")
                    f.write("=" * 30 + "\n\n")
                    
                    # Add overview
                    f.write(
                        f"The reviewer provided a {'positive' if assessment == 'positive' else 'constructive'} assessment. "
                        f"We have addressed all their concerns and suggestions as detailed below.\n\n"
                    )
                    
                    # Add responses to main concerns
                    main_concerns = reviewer_data.get("main_concerns", [])
                    if main_concerns:
                        f.write("MAIN CONCERNS\n")
                        f.write("-" * 20 + "\n\n")
                        for i, concern in enumerate(main_concerns, 1):
                            f.write(f"Concern {i}: {concern}\n")
                            f.write(f"Response: We thank the reviewer for this important point. "
                                    f"We have carefully considered this concern and made appropriate "
                                    f"revisions to the manuscript.\n\n")
                    
                    # Add responses to comments by type
                    comments_by_type = reviewer_data.get("comments_by_type", {})
                    for comment_type, comments in comments_by_type.items():
                        if comments:
                            f.write(f"COMMENTS ON {comment_type.upper()}\n")
                            f.write("-" * 20 + "\n\n")
                            for i, comment in enumerate(comments, 1):
                                f.write(f"Comment {i}: {comment}\n")
                                f.write(f"Response: We appreciate this comment and have revised the "
                                        f"{comment_type} accordingly.\n\n")
                
                # Add conclusion
                f.write("\nCONCLUSION\n")
                f.write("-" * 20 + "\n\n")
                f.write(
                    f"We believe that these revisions have significantly improved the quality of our manuscript. "
                    f"We have addressed all the reviewers' concerns and suggestions, and we hope that the revised "
                    f"version meets the standards for publication in your journal.\n\n"
                    f"Thank you for your consideration.\n\n"
                    f"Sincerely,\n"
                    f"The Authors"
                )
            
            self.logger.info(f"Created plain text editor letter at {txt_path}")
            return txt_path
            
        except Exception as e:
            self.logger.error(f"Error creating fallback editor letter: {e}")
            return ""
    
    def save(self, output_path: str) -> str:
        """
        Save the document.
        
        Args:
            output_path: Path where the document should be saved
            
        Returns:
            Path to the saved document
        """
        self.logger.info(f"Saving document to {output_path}")
        
        if not self._document:
            self.logger.error("No document to save")
            return ""
        
        try:
            # Save the document
            self._document.save(output_path)
            return output_path
            
        except Exception as e:
            self.logger.error(f"Error saving document: {e}")
            return ""