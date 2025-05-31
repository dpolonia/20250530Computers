"""
PDF adapter implementation.

This module implements the PDF adapter interface, providing functionality for
interacting with PDF files.
"""

import logging
import os
import re
from typing import Dict, Any, List, Optional, Tuple

from src.core.context import RevisionContext
from src.adapters.interfaces import PDFAdapterInterface


class PDFAdapter(PDFAdapterInterface):
    """
    Adapter for PDF file operations.
    
    This adapter is responsible for extracting text, metadata, tables, figures,
    and sections from PDF files, as well as converting PDFs to DOCX format.
    """
    
    def __init__(self, context: RevisionContext):
        """
        Initialize the PDF adapter.
        
        Args:
            context: The shared revision context
        """
        self.context = context
        self.logger = context.logger or logging.getLogger(__name__)
    
    def read(self, file_path: str) -> str:
        """
        Read a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text from the PDF
        """
        return self.extract_text(file_path)
    
    def write(self, content: Any, file_path: str) -> str:
        """
        Write content to a PDF file.
        
        This is a stub implementation as direct PDF writing is not supported.
        
        Args:
            content: Content to write
            file_path: Path where the file should be saved
            
        Returns:
            Path to the written file
        """
        self.logger.warning("Direct PDF writing is not supported. Use a document service instead.")
        return ""
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        self.logger.info(f"Extracting text from {file_path}")
        
        try:
            # Using PyPDF2 for text extraction
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
            
            return text
            
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Using fallback method.")
            return self._fallback_extract_text(file_path)
        except Exception as e:
            self.logger.error(f"Error extracting text: {e}")
            return ""
    
    def _fallback_extract_text(self, file_path: str) -> str:
        """
        Fallback method for text extraction using subprocess.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        try:
            import subprocess
            
            # Try using pdftotext if available
            try:
                output = subprocess.run(
                    ['pdftotext', file_path, '-'],
                    capture_output=True,
                    text=True,
                    check=True
                )
                return output.stdout
            except (subprocess.SubprocessError, FileNotFoundError):
                # If pdftotext not available, try using pdf2txt.py from pdfminer
                try:
                    output = subprocess.run(
                        ['pdf2txt.py', file_path],
                        capture_output=True,
                        text=True,
                        check=True
                    )
                    return output.stdout
                except (subprocess.SubprocessError, FileNotFoundError):
                    self.logger.error("No PDF extraction tools available")
                    return ""
                
        except Exception as e:
            self.logger.error(f"Error in fallback text extraction: {e}")
            return ""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted metadata
        """
        self.logger.info(f"Extracting metadata from {file_path}")
        
        try:
            # Using PyPDF2 for metadata extraction
            import PyPDF2
            
            metadata = {}
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        # Clean up metadata keys
                        clean_key = key
                        if key.startswith('/'):
                            clean_key = key[1:]
                        metadata[clean_key.lower()] = value
                
                # Add page count
                metadata['pages'] = len(reader.pages)
            
            # Extract title and author from first page if not in metadata
            if 'title' not in metadata or not metadata['title']:
                first_page_text = reader.pages[0].extract_text()
                title_match = re.search(r'^(.+?)\n', first_page_text)
                if title_match:
                    metadata['title'] = title_match.group(1).strip()
            
            if 'author' not in metadata or not metadata['author']:
                first_page_text = reader.pages[0].extract_text()
                # Look for typical author patterns (simplified)
                author_match = re.search(r'\n([A-Z][a-z]+ [A-Z][a-z]+(?:,? (?:and )?[A-Z][a-z]+ [A-Z][a-z]+)*)', first_page_text)
                if author_match:
                    metadata['author'] = author_match.group(1).strip()
            
            return metadata
            
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Using empty metadata.")
            return {}
        except Exception as e:
            self.logger.error(f"Error extracting metadata: {e}")
            return {}
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of extracted tables
        """
        self.logger.info(f"Extracting tables from {file_path}")
        
        try:
            # Extract text first
            text = self.extract_text(file_path)
            
            # Simple heuristic to identify tables (could be improved with proper table extraction libraries)
            tables = []
            table_pattern = re.compile(r'Table\s+(\d+)[.:]?\s*([^\n]+)', re.IGNORECASE)
            
            for match in table_pattern.finditer(text):
                table_num = match.group(1)
                caption = match.group(2).strip()
                
                tables.append({
                    "number": int(table_num),
                    "caption": caption,
                    "content": []  # Placeholder for actual table content
                })
            
            return tables
            
        except Exception as e:
            self.logger.error(f"Error extracting tables: {e}")
            return []
    
    def extract_figures(self, file_path: str) -> List[Tuple[str, Optional[str]]]:
        """
        Extract figures from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of tuples (caption, path)
        """
        self.logger.info(f"Extracting figures from {file_path}")
        
        try:
            # Extract text first
            text = self.extract_text(file_path)
            
            # Simple heuristic to identify figures (could be improved with proper figure extraction libraries)
            figures = []
            figure_pattern = re.compile(r'(?:Figure|Fig\.)\s+(\d+)[.:]?\s*([^\n]+)', re.IGNORECASE)
            
            for match in figure_pattern.finditer(text):
                figure_num = match.group(1)
                caption = match.group(2).strip()
                
                figures.append((caption, None))  # No path for figures extracted this way
            
            return figures
            
        except Exception as e:
            self.logger.error(f"Error extracting figures: {e}")
            return []
    
    def extract_sections(self, file_path: str) -> Dict[str, str]:
        """
        Extract sections from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary mapping section names to content
        """
        self.logger.info(f"Extracting sections from {file_path}")
        
        try:
            # Extract text first
            text = self.extract_text(file_path)
            
            # Simple heuristic to identify sections (could be improved with more sophisticated methods)
            sections = {}
            
            # Define common section header patterns
            section_patterns = [
                # Numbered sections (e.g., "1. Introduction")
                r'(\d+\.?\s+[A-Z][a-zA-Z\s]+)\n',
                # Uppercase sections (e.g., "INTRODUCTION")
                r'([A-Z]{2,}[A-Z\s]+)\n',
                # Title case sections (e.g., "Introduction")
                r'\n([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\n'
            ]
            
            # Combine patterns
            combined_pattern = '|'.join(section_patterns)
            matches = re.finditer(combined_pattern, text)
            
            # Extract section headers and their positions
            section_positions = []
            for match in matches:
                for i, group in enumerate(match.groups()):
                    if group is not None:
                        section_positions.append((match.start(), group.strip()))
                        break
            
            # Sort by position
            section_positions.sort()
            
            # Extract section content
            for i, (pos, name) in enumerate(section_positions):
                # Get start position (after section header)
                start = pos + len(name)
                
                # Get end position (start of next section or end of text)
                end = section_positions[i + 1][0] if i < len(section_positions) - 1 else len(text)
                
                # Extract content
                content = text[start:end].strip()
                
                # Add to sections dictionary
                sections[name] = content
            
            # If no sections found, create a default "Content" section
            if not sections:
                sections["Content"] = text
            
            return sections
            
        except Exception as e:
            self.logger.error(f"Error extracting sections: {e}")
            return {"Content": text} if 'text' in locals() else {"Content": ""}
    
    def extract_references(self, file_path: str) -> List[str]:
        """
        Extract references from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of reference strings
        """
        self.logger.info(f"Extracting references from {file_path}")
        
        try:
            # Extract text first
            text = self.extract_text(file_path)
            
            # Look for references section
            references_section = ""
            ref_headers = ["References", "REFERENCES", "Bibliography", "BIBLIOGRAPHY"]
            
            for header in ref_headers:
                ref_match = re.search(rf'{header}\s*\n(.*?)(?:\n\s*(?:[A-Z]{2,}|\d+\.)\s|\Z)', 
                                     text, re.DOTALL)
                if ref_match:
                    references_section = ref_match.group(1)
                    break
            
            if not references_section:
                # Try to find the last section, which is often references
                sections = self.extract_sections(file_path)
                section_names = list(sections.keys())
                if section_names and any(ref in section_names[-1].lower() for ref in ["reference", "bibliography"]):
                    references_section = sections[section_names[-1]]
            
            # Parse individual references
            references = []
            
            # Look for numbered references (e.g., "[1] Author, Title...")
            numbered_refs = re.findall(r'\[\d+\]\s+([^\[\n]+)', references_section)
            if numbered_refs:
                references.extend(numbered_refs)
            else:
                # Try another common format (Author et al., year)
                author_year_refs = re.findall(r'([A-Z][a-z]+(?:,?\s+(?:et al\.|and)\s+[A-Z][a-z]+)*,?\s+\(\d{4}\)[^\.]+\.)', 
                                             references_section)
                if author_year_refs:
                    references.extend(author_year_refs)
                else:
                    # Fallback: split by newlines and filter
                    lines = references_section.split('\n')
                    for line in lines:
                        line = line.strip()
                        if len(line) > 20 and (',' in line or '.' in line):
                            references.append(line)
            
            return [ref.strip() for ref in references if ref.strip()]
            
        except Exception as e:
            self.logger.error(f"Error extracting references: {e}")
            return []
    
    def pdf_to_docx(self, file_path: str) -> str:
        """
        Convert a PDF file to DOCX format.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Path to the created DOCX file
        """
        self.logger.info(f"Converting {file_path} to DOCX")
        
        try:
            # Generate output path
            base_name = os.path.splitext(file_path)[0]
            output_path = f"{base_name}.docx"
            
            # Check if docx file already exists
            if os.path.exists(output_path):
                self.logger.info(f"DOCX file already exists at {output_path}")
                return output_path
            
            # Try using pdf2docx library
            try:
                from pdf2docx import Converter
                
                cv = Converter(file_path)
                cv.convert(output_path)
                cv.close()
                
                self.logger.info(f"PDF converted to DOCX at {output_path}")
                return output_path
                
            except ImportError:
                self.logger.warning("pdf2docx not installed. Using fallback method.")
                
                # Fallback: Extract text and create a simple DOCX
                from docx import Document
                
                # Extract text and metadata
                text = self.extract_text(file_path)
                metadata = self.extract_metadata(file_path)
                
                # Create document
                doc = Document()
                
                # Add title
                title = metadata.get('title', 'Untitled Document')
                doc.add_heading(title, 0)
                
                # Add content
                for paragraph in text.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                # Save document
                doc.save(output_path)
                
                self.logger.info(f"Simple DOCX created at {output_path}")
                return output_path
                
        except Exception as e:
            self.logger.error(f"Error converting PDF to DOCX: {e}")
            return ""
    
    def close(self) -> None:
        """Clean up resources."""
        pass