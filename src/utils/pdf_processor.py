"""Utilities for processing PDF files."""

import os
import tempfile
import fitz  # PyMuPDF
import docx
from docx import Document
from typing import Dict, List, Tuple, Optional

class PDFProcessor:
    """Class for handling PDF document processing and text extraction."""
    
    def __init__(self, pdf_path: str):
        """Initialize with path to PDF file.
        
        Args:
            pdf_path: Path to the PDF file
        """
        self.pdf_path = pdf_path
        self.doc = None
        self.text = ""
        self.sections = {}
        self._load_document()
    
    def _load_document(self):
        """Load the PDF document."""
        try:
            self.doc = fitz.open(self.pdf_path)
            self._extract_text()
        except Exception as e:
            print(f"Error loading PDF {self.pdf_path}: {e}")
            raise
    
    def _extract_text(self):
        """Extract text from the PDF document."""
        text = ""
        for page in self.doc:
            text += page.get_text()
        self.text = text
        
    def extract_text(self, output_path: str = None) -> str:
        """Extract text from the PDF document.
        
        Args:
            output_path: Optional path to save the extracted text to a file
            
        Returns:
            Extracted text as a string
        """
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.text)
        return self.text
    
    def extract_sections(self) -> Dict[str, str]:
        """Extract sections from the PDF document.
        
        Returns:
            Dictionary with section names as keys and section content as values
        """
        # This is a simplified implementation - actual section extraction
        # would require more sophisticated parsing based on document structure
        sections = {}
        current_section = "Abstract"
        current_content = []
        
        lines = self.text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Identify potential section headers (uppercase, short lines)
            if (line.isupper() or (line[0].isupper() and len(line.split()) <= 5 and 
                any(kw in line.lower() for kw in ['introduction', 'method', 'result', 'discussion', 
                                                'conclusion', 'reference', 'abstract']))):
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                    current_content = []
                current_section = line
            else:
                current_content.append(line)
        
        # Add the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
            
        self.sections = sections
        return sections
    
    def extract_tables(self) -> List[str]:
        """Extract tables from the PDF document.
        
        Returns:
            List of extracted table content as strings
        """
        tables = []
        for page in self.doc:
            # Find tables based on heuristics (this is simplified)
            # In a real implementation, we would use more sophisticated table detection
            blocks = page.get_text("blocks")
            for block in blocks:
                # Check if block might be a table based on layout
                if len(block) >= 5:  # Block with coordinates
                    x0, y0, x1, y1, text, block_type, block_no = block[:7]
                    if text.count("|") > 3 or text.count("\t") > 3:
                        tables.append(text)
        return tables
    
    def extract_figures(self) -> List[Tuple[str, bytes]]:
        """Extract figures from the PDF document.
        
        Returns:
            List of tuples with figure captions and image data
        """
        figures = []
        for page_num, page in enumerate(self.doc):
            # Extract images
            img_list = page.get_images(full=True)
            for img_index, img in enumerate(img_list):
                xref = img[0]
                base_img = self.doc.extract_image(xref)
                image_bytes = base_img["image"]
                
                # Find nearby caption text (simplified approach)
                page_text = page.get_text()
                caption_candidates = [line for line in page_text.split('\n') 
                                     if "figure" in line.lower() or "fig" in line.lower()]
                caption = caption_candidates[0] if caption_candidates else f"Figure {page_num+1}-{img_index+1}"
                
                figures.append((caption, image_bytes))
        return figures
    
    def extract_references(self) -> List[str]:
        """Extract references from the PDF document.
        
        Returns:
            List of reference strings
        """
        references = []
        ref_section = None
        
        # Find the references section
        for section_name, content in self.sections.items():
            if "reference" in section_name.lower() or "bibliography" in section_name.lower():
                ref_section = content
                break
        
        if ref_section:
            # Simple heuristic: Split by line and look for numbered references
            lines = ref_section.split('\n')
            current_ref = ""
            
            for line in lines:
                if line.strip().startswith("[") or line.strip()[0].isdigit():
                    if current_ref:
                        references.append(current_ref.strip())
                    current_ref = line
                else:
                    current_ref += " " + line
            
            # Add the last reference
            if current_ref:
                references.append(current_ref.strip())
        
        return references
    
    def pdf_to_docx(self, output_path: Optional[str] = None) -> str:
        """Convert PDF to DOCX format.
        
        Args:
            output_path: Path for the output DOCX file. If None, a temporary file is created.
            
        Returns:
            Path to the created DOCX file
        """
        if not output_path:
            fd, output_path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
        
        doc = Document()
        
        # Add document content section by section
        for section_name, content in self.sections.items():
            doc.add_heading(section_name, level=1)
            doc.add_paragraph(content)
        
        # Add figures and tables (simplified)
        figures = self.extract_figures()
        for caption, _ in figures:
            p = doc.add_paragraph()
            p.add_run(caption).italic = True
        
        doc.save(output_path)
        return output_path
    
    def get_page_count(self) -> int:
        """Get the number of pages in the PDF document.
        
        Returns:
            Number of pages in the document
        """
        if self.doc:
            return len(self.doc)
        return 0
    
    def save_first_page_as_image(self, output_path: str, dpi: int = 150):
        """Save the first page of the PDF as an image.
        
        Args:
            output_path: Path to save the image
            dpi: Resolution in dots per inch
        """
        if not self.doc or len(self.doc) == 0:
            return
            
        # Get the first page
        page = self.doc[0]
        
        # Render the page to a pixmap
        pixmap = page.get_pixmap(dpi=dpi)
        
        # Save the pixmap as a PNG
        pixmap.save(output_path)
    
    def close(self):
        """Close the PDF document."""
        if self.doc:
            self.doc.close()
