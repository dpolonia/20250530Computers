"""Document processing utilities for DOCX files."""

import os
import re
import datetime
from typing import Dict, List, Optional, Tuple, Set
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

class DocumentProcessor:
    """Class for handling DOCX document processing and editing."""
    
    def __init__(self, docx_path: str):
        """Initialize with path to DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
        """
        self.docx_path = docx_path
        self.doc = None
        self._load_document()
    
    def _load_document(self):
        """Load the DOCX document."""
        try:
            self.doc = Document(self.docx_path)
        except Exception as e:
            print(f"Error loading DOCX {self.docx_path}: {e}")
            raise
    
    def get_full_text(self) -> str:
        """Get the full text of the document.
        
        Returns:
            String containing all text in the document
        """
        return "\n".join([para.text for para in self.doc.paragraphs])
        
    def extract_text(self, output_path: str = None) -> str:
        """Extract text from the document.
        
        Args:
            output_path: Optional path to save the extracted text to a file
            
        Returns:
            Extracted text as a string
        """
        text = self.get_full_text()
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        return text
    
    def extract_sections(self) -> Dict[str, str]:
        """Extract sections from the document based on headings.
        
        Returns:
            Dictionary with section names as keys and section content as values
        """
        sections = {}
        current_section = "Main"
        current_content = []
        
        for para in self.doc.paragraphs:
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                    current_content = []
                current_section = para.text
            else:
                if para.text.strip():  # Skip empty paragraphs
                    current_content.append(para.text)
        
        # Add the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
            
        return sections
    
    def add_tracked_change(self, old_text: str, new_text: str, reason: str = "") -> bool:
        """Add a tracked change to the document.
        
        Args:
            old_text: Text to be replaced
            new_text: New text to replace with
            reason: Reason for the change
            
        Returns:
            True if change was made, False otherwise
        """
        # Find and replace text with formatting to indicate tracked changes
        found = False
        
        for para in self.doc.paragraphs:
            if old_text in para.text:
                found = True
                # Clear paragraph content
                text = para.text
                para.clear()
                
                # Split text on the old_text to preserve content before and after
                parts = text.split(old_text, 1)
                
                # Add text before the change
                if parts[0]:
                    para.add_run(parts[0])
                
                # Add the deleted text (strikethrough and red)
                if old_text:
                    deleted_run = para.add_run(old_text)
                    deleted_run.font.strike = True
                    deleted_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Add comment about the change if reason provided
                    if reason:
                        comment_run = para.add_run(f" [{reason}] ")
                        comment_run.font.italic = True
                        comment_run.font.size = Pt(8)
                
                # Add the new text (green and underlined)
                if new_text:
                    new_run = para.add_run(new_text)
                    new_run.font.color.rgb = RGBColor(0, 128, 0)
                    new_run.font.underline = True
                
                # Add text after the change
                if len(parts) > 1 and parts[1]:
                    para.add_run(parts[1])
                
                break
        
        return found
    
    def create_changes_document(self, changes: List[Tuple[str, str, str, Optional[int]]], output_path: str) -> str:
        """Create a document detailing all changes.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        doc = Document()
        doc.add_heading('Change Document', 0)
        
        # Add introduction
        doc.add_paragraph('This document details all changes made to the original paper in response to reviewer comments.')
        
        # Add changes table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Line Number'
        header_cells[1].text = 'Original Text'
        header_cells[2].text = 'Revised Text'
        header_cells[3].text = 'Reason for Change'
        
        # Add changes
        for old_text, new_text, reason, line_number in changes:
            row_cells = table.add_row().cells
            row_cells[0].text = str(line_number) if line_number is not None else 'N/A'
            row_cells[1].text = old_text
            row_cells[2].text = new_text
            row_cells[3].text = reason
        
        doc.save(output_path)
        return output_path
    
    def create_revision_summary(self, issues: List[Dict], solutions: List[Dict], output_path: str) -> str:
        """Create a document summarizing the revision process.
        
        Args:
            issues: List of dictionaries with issue details
            solutions: List of dictionaries with solution details
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        doc = Document()
        doc.add_heading('Revision Summary', 0)
        
        # Add issues section
        doc.add_heading('Issues Identified', 1)
        for i, issue in enumerate(issues, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {issue['title']}").bold = True
            doc.add_paragraph(f"Description: {issue['description']}")
            doc.add_paragraph(f"Source: {issue['source']}")
            doc.add_paragraph(f"Severity: {issue['severity']}")
            doc.add_paragraph()
        
        # Add solutions section
        doc.add_heading('Proposed Solutions', 1)
        for i, solution in enumerate(solutions, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {solution['title']}").bold = True
            doc.add_paragraph(f"Implementation: {solution['implementation']}")
            doc.add_paragraph(f"Complexity: {solution['complexity']}")
            doc.add_paragraph(f"Impact: {solution['impact']}")
            doc.add_paragraph()
        
        # Add conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph('This document summarizes the issues identified in the review process and the proposed solutions. ' +
                       'These changes aim to address reviewer comments while maintaining the integrity of the research.')
        
        doc.save(output_path)
        return output_path
    
    def create_editor_letter(self, reviewer_responses: List[Dict], output_path: str, process_summary: Dict = None) -> str:
        """Create a letter to the editor with responses to reviewer comments.
        
        Args:
            reviewer_responses: List of dictionaries with response details
            output_path: Path where the document should be saved
            process_summary: Optional dictionary containing review process summary
            
        Returns:
            Path to the created document
        """
        doc = Document()
        
        # Add header information
        doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("To: The Editor, Computers (ISSN: 2073-431X)")
        doc.add_paragraph("Subject: Revised manuscript submission")
        doc.add_paragraph()
        
        # Add salutation
        doc.add_paragraph("Dear Editor,")
        
        # Add introduction
        doc.add_paragraph(
            "Thank you for the opportunity to revise our manuscript. We have carefully addressed all the comments " +
            "provided by the reviewers and made the necessary changes to improve the quality of our paper. " +
            "We believe that the revised version addresses all the concerns raised and significantly improves the manuscript."
        )
        
        # Add reviewer responses
        doc.add_heading('Responses to Reviewer Comments', 1)
        
        for i, response in enumerate(reviewer_responses, 1):
            doc.add_heading(f"Reviewer {i}", 2)
            
            for j, comment in enumerate(response['comments'], 1):
                p = doc.add_paragraph()
                p.add_run(f"Comment {j}: ").bold = True
                p.add_run(comment['comment'])
                
                p = doc.add_paragraph()
                p.add_run("Response: ").bold = True
                p.add_run(comment['response'])
                
                p = doc.add_paragraph()
                p.add_run("Changes made: ").bold = True
                p.add_run(comment['changes'])
                
                doc.add_paragraph()
        
        # Add closing
        doc.add_paragraph(
            "We hope that the revised manuscript now meets the standards for publication in Computers. " +
            "We look forward to your feedback and are available to address any additional questions or concerns."
        )
        
        # Add revision process disclaimer if provided
        if process_summary:
            doc.add_heading('REVISION PROCESS DISCLOSURE', 1)
            
            # Add detailed information about the multi-persona revision process
            doc.add_paragraph(process_summary.get("process_description", ""))
            
            # Add statistics
            stats_paragraph = doc.add_paragraph()
            stats_paragraph.add_run("REVISION STATISTICS:\n").bold = True
            stats_paragraph.add_run(f"• Reviewers: {process_summary.get('reviewer_count', 0)}\n")
            stats_paragraph.add_run(f"• Reviewer personas: {process_summary.get('total_reviewer_personas', 0)}\n")
            stats_paragraph.add_run(f"• Editors: {process_summary.get('editor_count', 0)}\n")
            stats_paragraph.add_run(f"• Fine-tuned personas used: {process_summary.get('fine_personas_used', 0)}\n")
            stats_paragraph.add_run(f"• Total reviews generated: {process_summary.get('review_count', 0)}\n")
            stats_paragraph.add_run(f"• Final decision: {process_summary.get('decision', 'Not specified')}")
            
            # Add acknowledgment
            doc.add_paragraph(
                "We acknowledge that this revision process utilized advanced AI-assisted multi-persona review " +
                "technology to ensure a comprehensive, diverse, and thorough evaluation of our manuscript. " +
                "The multiple persona approach ensures that our paper was examined from various academic " +
                "perspectives before making our final revisions."
            )
            
            # Add attribution for FinePersonas if used
            if process_summary.get('fine_personas_used', 0) > 0:
                doc.add_paragraph(
                    "This revision process utilized the FinePersonas dataset " +
                    "(https://huggingface.co/datasets/argilla/FinePersonas-v0.1) " +
                    "to enhance the diversity and expertise of the reviewer perspectives."
                )
        
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("The Authors")
        
        doc.save(output_path)
        return output_path
    
    def save(self, output_path: Optional[str] = None):
        """Save the document.
        
        Args:
            output_path: Path for saving the document. If None, original path is used.
        """
        save_path = output_path if output_path else self.docx_path
        self.doc.save(save_path)
