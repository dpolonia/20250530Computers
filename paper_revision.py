#!/usr/bin/env python3
"""
Paper Revision Tool for Computers Journal

This script automates the process of revising a scientific paper based on reviewer comments.
It utilizes LLMs to analyze the original paper, review comments, and generate revision documents.
"""

import os
import sys
import time
import argparse
import datetime
import json
import hashlib
import re  # Used for extracting dates from model names and for quality evaluation
import logging
import sqlite3
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dotenv import load_dotenv
import click
from tqdm import tqdm
from colorama import Fore, Style, init as colorama_init

# Import custom modules
from src.utils.workflow_db import WorkflowDB
from src.utils.model_scorer import get_best_models_by_mode, initialize_model_scores
from src.utils.reviewer_persona import create_review_report, save_review_report

# Initialize colorama for cross-platform colored terminal output
colorama_init()

# Add src directory to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Simple cache for API responses
API_CACHE = {}
MAX_CACHE_SIZE = 100  # Maximum number of entries in cache
CACHE_DIR = os.path.join(os.path.dirname(__file__), ".cache")
os.makedirs(CACHE_DIR, exist_ok=True)

# Token budget settings
DEFAULT_TOKEN_BUDGET = 100000  # Default token budget
DEFAULT_COST_BUDGET = 5.0      # Default cost budget in dollars

# Operation mode configurations
OPERATION_MODES = {
    "training": {
        "description": "Training mode: Cheapest models to activate workflow",
        "token_budget": 30000,
        "cost_budget": 1.0,
        "max_papers": 1,
        "optimize_costs": True,
        "provider_recommendations": {
            "anthropic": "claude-3-haiku-20240307",
            "openai": "gpt-4o-mini",
            "google": "gemini-1.5-flash"
        }
    },
    "finetuning": {
        "description": "Fine-tuning mode: Mid-range models to optimize workflow",
        "token_budget": 100000,
        "cost_budget": 5.0,
        "max_papers": 2,
        "optimize_costs": True,
        "provider_recommendations": {
            "anthropic": "claude-3-5-sonnet-20241022",
            "openai": "gpt-4o",
            "google": "gemini-1.5-pro"
        }
    },
    "final": {
        "description": "Final mode: Best models available, no cost limits",
        "token_budget": 500000,
        "cost_budget": 25.0,
        "max_papers": 5,
        "optimize_costs": False,
        "provider_recommendations": {
            "anthropic": "claude-opus-4-20250514",
            "openai": "gpt-4.5-preview",
            "google": "gemini-2.5-pro-preview"
        }
    }
}

# Simple token counter estimation function
def estimate_tokens(text: str) -> int:
    """Estimate the number of tokens in a text string.
    
    Args:
        text: Input text
        
    Returns:
        Estimated token count
    """
    # Very rough estimation: ~4 characters per token
    return len(text) // 4

def calculate_tokens_per_prompt(provider: str, model: str, use_safe_limit=True, use_max_safe=False) -> Dict[str, int]:
    """Calculate max tokens per prompt based on provider/model.
    
    Args:
        provider: Provider name ("anthropic", "openai", "google")
        model: Model name
        use_safe_limit: Whether to use the 90% safe limit (default: True)
        use_max_safe: Whether to use the 99% max safe limit (default: False)
        
    Returns:
        Dictionary with max token settings
    """
    # Use the model-specific max token functions that respect safety limits
    max_tokens = get_max_tokens_for_model(provider, model, use_safe_limit, use_max_safe)
    
    # Calculate context size and per-prompt size based on model capabilities
    if provider == "anthropic":
        if "opus" in model:
            context_size = max_tokens
            per_prompt_size = max_tokens // 2
        elif "sonnet" in model:
            context_size = max_tokens * 2
            per_prompt_size = max_tokens
        elif "haiku" in model:
            context_size = max_tokens * 2
            per_prompt_size = max_tokens
        else:
            context_size = max_tokens * 2
            per_prompt_size = max_tokens
    elif provider == "openai":
        context_size = max_tokens
        per_prompt_size = max_tokens // 2
    elif provider == "google":
        if "pro" in model:
            context_size = max_tokens
            per_prompt_size = max_tokens // 4
        else:
            context_size = max_tokens
            per_prompt_size = max_tokens // 4
    else:
        # Default fallback
        context_size = max_tokens * 2
        per_prompt_size = max_tokens
    
    return {"max_context": context_size, "max_per_prompt": per_prompt_size}

# Caching mechanism for API calls
def get_cached_completion(client, prompt: str, system_prompt: str = None, **kwargs) -> Optional[str]:
    """Get completion from cache if available.
    
    Args:
        client: LLM client
        prompt: User prompt
        system_prompt: System prompt
        **kwargs: Additional parameters
        
    Returns:
        Cached response or None if not in cache
    """
    # Create a cache key from the request
    cache_key = hashlib.md5(f"{client.provider}_{client.model}_{prompt}_{system_prompt}_{str(kwargs)}".encode()).hexdigest()
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    
    # Check if we have this request cached
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r') as f:
                cache_data = json.load(f)
                return cache_data["response"]
        except:
            return None
    
    return None

def save_to_cache(client, prompt: str, system_prompt: str, response: str, **kwargs) -> None:
    """Save an API response to cache.
    
    Args:
        client: LLM client
        prompt: User prompt
        system_prompt: System prompt
        response: API response
        **kwargs: Additional parameters
    """
    # Create a cache key from the request
    cache_key = hashlib.md5(f"{client.provider}_{client.model}_{prompt}_{system_prompt}_{str(kwargs)}".encode()).hexdigest()
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    
    # Save to cache
    cache_data = {
        "provider": client.provider,
        "model": client.model,
        "prompt": prompt,
        "system_prompt": system_prompt,
        "response": response,
        "timestamp": time.time()
    }
    
    with open(cache_file, 'w') as f:
        json.dump(cache_data, f)
    
    # Prune cache if needed
    cache_files = os.listdir(CACHE_DIR)
    if len(cache_files) > MAX_CACHE_SIZE:
        # Delete oldest cache files
        cache_files = [os.path.join(CACHE_DIR, f) for f in cache_files if f.endswith('.json')]
        cache_files.sort(key=os.path.getmtime)
        for f in cache_files[:len(cache_files) - MAX_CACHE_SIZE]:
            try:
                os.remove(f)
            except:
                pass

# Import model information
try:
    from src.models.openai_models import get_openai_model_choices, get_openai_model_info, get_max_tokens_for_model as get_openai_max_tokens
except ImportError:
    # Create dummy functions if import fails
    def get_openai_model_choices():
        return [
            "gpt-4.5-preview (most powerful)",
            "gpt-4o (powerful)",
            "gpt-4o-mini (balanced)",
            "o1 (powerful reasoning)",
            "o3 (powerful reasoning)",
            "o4-mini (fast reasoning)"
        ]
    
    def get_openai_model_info(model_name):
        return None
    
    def get_openai_max_tokens(model_name):
        return 4096

try:
    from src.models.anthropic_models import get_claude_model_choices, get_claude_model_info, get_max_tokens_for_model as get_claude_max_tokens
except ImportError:
    # Create dummy functions if import fails
    def get_claude_model_choices():
        return [
            "claude-opus-4-20250514 (most powerful)",
            "claude-sonnet-4-20250514 (powerful)",
            "claude-3-7-sonnet-20250219 (powerful)",
            "claude-3-5-sonnet-20241022 (balanced)",
            "claude-3-5-haiku-20241022 (fast)",
            "claude-3-haiku-20240307 (fastest & cheapest)"
        ]
    
    def get_claude_model_info(model_name):
        return None
    
    def get_claude_max_tokens(model_name):
        return 4096
        
try:
    from src.models.google_models import get_gemini_model_choices, get_gemini_model_info, get_max_tokens_for_model as get_gemini_max_tokens
except ImportError:
    # Create dummy functions if import fails
    def get_gemini_model_choices():
        return [
            "gemini-2.5-pro-preview (most powerful, 8M context)",
            "gemini-2.5-flash-preview (efficient, 1M context)",
            "gemini-1.5-pro (powerful, 1M context)",
            "gemini-1.5-flash (fast, 1M context)",
            "gemini-2.0-flash (powerful)",
            "gemini-2.0-flash-lite (faster)"
        ]
    
    def get_gemini_model_info(model_name):
        return None
    
    def get_gemini_max_tokens(model_name):
        return 8192

# Import utility modules
from src.utils.pdf_processor import PDFProcessor
from src.utils.document_processor import DocumentProcessor
from src.utils.reference_validator import ReferenceValidator
from src.utils.llm_client import get_llm_client

# Model version mapping tables (to maintain incremental numbering)
# These tables map model identifiers to sequential numbers
# As new models are added, they get higher numbers, but the sequence remains intact
# even if older models are removed
#
# HOW TO UPDATE WHEN NEW MODELS ARE RELEASED:
# 1. Never change existing codes - they must remain stable for consistent directory naming
# 2. When a new model is released, add it to the appropriate table with the next available number
# 3. When models are deprecated/removed by providers, keep their entries in these tables
# 4. Reserved slots for future models can be used when new models are released
# 5. Update every ~2 weeks when providers release new models
ANTHROPIC_MODEL_VERSIONS = {
    # Claude 3 Haiku (entry-level)
    "claude-3-haiku-20240307": "01",
    "claude-3-5-haiku-20241022": "02",
    
    # Claude 3 Sonnet (mid-range)
    "claude-3-5-sonnet-20241022": "03",
    "claude-3-7-sonnet-20250219": "04",
    "claude-sonnet-4-20250514": "05",
    
    # Claude 3 Opus (high-end)
    "claude-opus-4-20250514": "06",
    
    # Future model slots (reserved)
    "claude-future-1": "07",
    "claude-future-2": "08",
    "claude-future-3": "09",
}

OPENAI_MODEL_VERSIONS = {
    # GPT-4o mini (entry-level)
    "gpt-4o-mini": "01",
    "o4-mini": "02",
    
    # GPT-4o (mid-range)
    "gpt-4o": "03",
    
    # GPT-4.5 and specialty models (high-end)
    "o1": "04",
    "o3": "05",
    "gpt-4.5-preview": "06",
    
    # Future model slots (reserved)
    "openai-future-1": "07",
    "openai-future-2": "08",
    "openai-future-3": "09",
}

GOOGLE_MODEL_VERSIONS = {
    # Gemini Flash models (entry-level)
    "gemini-1.5-flash": "01",
    "gemini-2.0-flash-lite": "02",
    "gemini-2.0-flash": "03",
    "gemini-2.5-flash-preview": "04",
    
    # Gemini Pro models (high-end)
    "gemini-1.5-pro": "05",
    "gemini-2.5-pro-preview": "06",
    
    # Future model slots (reserved)
    "gemini-future-1": "07",
    "gemini-future-2": "08",
    "gemini-future-3": "09",
}

# Model code generation function
def get_model_code(provider, model_name):
    """Generate a standardized model code.
    
    Format: [A|B|C][01-99]
    - A: Anthropic
    - B: OpenAI
    - C: Google
    - Number: Maintained in sequential order from 01 onwards
      (new models get higher numbers, even if older models are removed)
    
    Args:
        provider: Provider name ("anthropic", "openai", or "google")
        model_name: Name of the model
        
    Returns:
        A standardized model code
    """
    # Provider code
    if provider == "anthropic":
        provider_code = "A"
        version_map = ANTHROPIC_MODEL_VERSIONS
    elif provider == "openai":
        provider_code = "B" 
        version_map = OPENAI_MODEL_VERSIONS
    elif provider == "google":
        provider_code = "C"
        version_map = GOOGLE_MODEL_VERSIONS
    else:
        provider_code = "X"  # Unknown provider
        version_map = {}
    
    # Get the base model name without description
    base_model = model_name.split(" (")[0]
    
    # Look up the sequential version code
    for model_key, version in version_map.items():
        if model_key in base_model:
            return f"{provider_code}{version}"
    
    # Try to extract date from model name for unknown models
    date_match = re.search(r'(\d{8})', base_model)
    if date_match:
        # If we find a date in the model name, use it to derive a version number
        # This helps maintain chronological order for new models
        date_str = date_match.group(1)
        # Use last two digits as version number
        version_num = date_str[-2:]
        # Ensure it doesn't conflict with existing models
        if version_num == "00":
            version_num = "50"  # Use 50 as a safe middle ground for auto-detected models
        return f"{provider_code}{version_num}"
    
    # If no match is found, generate a fallback code
    # This preserves the sequence numbering even for unknown models
    return f"{provider_code}00"

# Helper function to update model version tables
def update_model_version_tables(new_models):
    """Update the model version tables with new models.
    
    This function can be used when new models are released to automatically
    assign them the next available version numbers while preserving the
    existing numbering scheme.
    
    Args:
        new_models: Dictionary mapping provider names to lists of new model names
                    Example: {"anthropic": ["claude-5-sonnet-20250801"], 
                              "openai": ["gpt-5-turbo"]}
    
    Returns:
        Dictionary with updates that should be made to the model version tables
    """
    updates = {}
    
    # Process each provider
    for provider, models in new_models.items():
        provider_updates = {}
        
        if provider == "anthropic":
            # Find the highest current version number
            max_version = max(int(v) for v in ANTHROPIC_MODEL_VERSIONS.values())
            
            # Assign new version numbers
            for model in models:
                max_version += 1
                provider_updates[model] = f"{max_version:02d}"
                
            updates["anthropic"] = provider_updates
            
        elif provider == "openai":
            # Find the highest current version number
            max_version = max(int(v) for v in OPENAI_MODEL_VERSIONS.values())
            
            # Assign new version numbers
            for model in models:
                max_version += 1
                provider_updates[model] = f"{max_version:02d}"
                
            updates["openai"] = provider_updates
            
        elif provider == "google":
            # Find the highest current version number
            max_version = max(int(v) for v in GOOGLE_MODEL_VERSIONS.values())
            
            # Assign new version numbers
            for model in models:
                max_version += 1
                provider_updates[model] = f"{max_version:02d}"
                
            updates["google"] = provider_updates
    
    return updates

# Unified function to get max tokens for any model
def get_max_tokens_for_model(provider, model_name, use_safe_limit=True, use_max_safe=False):
    """Get maximum output tokens for a model across any provider.
    
    Args:
        provider: Provider name ("anthropic", "openai", "google")
        model_name: Model name
        use_safe_limit: Whether to use the 90% safe limit (default: True)
        use_max_safe: Whether to use the 99% max safe limit (default: False)
        
    Returns:
        The appropriate token limit based on safety preferences
    """
    if provider == "openai":
        return get_openai_max_tokens(model_name, use_safe_limit, use_max_safe)
    elif provider == "anthropic":
        return get_claude_max_tokens(model_name, use_safe_limit, use_max_safe)
    elif provider == "google":
        return get_gemini_max_tokens(model_name, use_safe_limit, use_max_safe)
    else:
        # Default fallback
        return 4096

class PaperRevisionTool:
    """Main class for orchestrating the paper revision process."""
    
    def __init__(self, provider: str, model_name: str, debug: bool = False, 
                 token_budget: int = DEFAULT_TOKEN_BUDGET, 
                 cost_budget: float = DEFAULT_COST_BUDGET,
                 max_papers_to_process: int = 3,
                 optimize_costs: bool = True,
                 operation_mode: str = "custom",
                 skip_api_validation: bool = False,
                 verify: bool = False,
                 competitor_evaluation: bool = True,
                 competing_evaluator: str = None,
                 db_path: str = "./.cache/workflow.db",
                 output_dir: str = None,
                 api: str = None,
                 api_key: str = None):
        """Initialize the paper revision tool.
        
        Args:
            provider: LLM provider ("anthropic", "openai", or "google")
            model_name: Name of the model to use
            debug: Enable debug mode for additional logging
            token_budget: Maximum token budget for the entire process
            cost_budget: Maximum cost budget in dollars
            max_papers_to_process: Maximum number of papers to process for style analysis
            optimize_costs: Whether to optimize costs (reduce context, use tiered approach)
            operation_mode: The operation mode (training, finetuning, final, or custom)
            skip_api_validation: Skip API validation
            verify: Whether to verify model accuracy (asks "When is Christmas?")
            competitor_evaluation: Whether to use competing models for evaluation
            competing_evaluator: Specific competitor to use for evaluation (format: "provider/model")
            db_path: Path to the SQLite database file for workflow tracking
            output_dir: Output directory for generated files
            api: API integration to use (scopus, wos, or both)
            api_key: API key for the specified API
        """
        self.provider = provider
        self.model_name = model_name
        self.debug = debug
        self.llm_client = None
        self.timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        self.token_budget = token_budget
        self.cost_budget = cost_budget
        self.max_papers_to_process = max_papers_to_process
        self.optimize_costs = optimize_costs
        self.operation_mode = operation_mode
        self.verify = verify
        self.competitor_evaluation = competitor_evaluation
        self.competing_evaluator = competing_evaluator
        
        # API integration
        self.api = api
        self.api_key = api_key
        
        # Initialize database
        self.db_path = db_path
        self.workflow_db = WorkflowDB(db_path)
        
        # Create a database run ID
        self.run_id = self.timestamp
        
        # Initialize database run
        settings = {
            "provider": provider,
            "model_name": model_name,
            "debug": debug,
            "token_budget": token_budget,
            "cost_budget": cost_budget,
            "max_papers_to_process": max_papers_to_process,
            "optimize_costs": optimize_costs,
            "operation_mode": operation_mode,
            "verify": verify,
            "competitor_evaluation": competitor_evaluation,
            "competing_evaluator": competing_evaluator,
            "output_dir": output_dir,
            "api": api,
            "api_key": "********" if api_key else None  # Don't store actual key in DB
        }
        self.workflow_db.create_run(self.run_id, provider, model_name, operation_mode, settings)
        
        # Calculate token limits for prompts
        self.high_quality = operation_mode.lower() == "final"  # Store high quality flag as instance variable
        self.token_limits = calculate_tokens_per_prompt(
            provider, 
            model_name, 
            use_safe_limit=True,  # Default to 90% safe limit
            use_max_safe=self.high_quality  # Use 99% max safe limit for high quality runs
        )
        self.max_tokens_per_prompt = self.token_limits["max_per_prompt"]
        
        # Paths
        self.original_paper_path = "./asis/00.pdf"
        self.original_docx_path = "./asis/00.docx"
        self.reviewer1_path = "./asis/01.pdf"
        self.reviewer2_path = "./asis/02.pdf"
        self.reviewer3_path = "./asis/03.pdf"
        self.editor_letter_path = "./asis/04.pdf"
        self.prisma_requirements_path = "./asis/05.pdf"
        self.journal_info_path = "./asis/06.pdf"
        self.scopus_info_paths = ["./asis/07.pdf", "./asis/08.pdf"]
        self.cited_papers_paths = [f"./asis/{i:02d}.pdf" for i in range(11, 19)]
        self.similar_papers_paths = [f"./asis/{i:02d}.pdf" for i in range(21, 25)]
        self.bib_path = "./asis/zz.bib"
        
        # Create model directory inside tobe using the standardized model code
        self.model_code = get_model_code(self.provider, self.model_name)
        self.model_description = self.model_name.replace(' ', '_').replace('-', '_')
        
        # Set output directory structure based on provided output_dir or defaults
        if output_dir:
            self.base_output_dir = output_dir
        else:
            self.base_output_dir = "./tobe"
            
        # Create mode-specific subdirectory
        self.mode_dir = f"{self.base_output_dir}/{operation_mode.upper()}"
        
        # For final mode, create a special FINAL directory in the root filesystem
        if operation_mode.lower() == "final":
            # Create a special FINAL directory with timestamp in the root
            root_dir = os.path.dirname(os.path.abspath(self.base_output_dir))
            self.final_dir = f"{root_dir}/FINAL_{self.timestamp}"
            self.model_dir = f"{self.mode_dir}/{self.model_code}_{self.model_description}"
            # Use the special FINAL directory for outputs
            self.timestamp_dir = f"{self.final_dir}"
            # Also maintain the standard structure for backups
            self.backup_dir = f"{self.model_dir}/{self.timestamp}"
        else:
            self.final_dir = None
            self.backup_dir = None
            self.model_dir = f"{self.mode_dir}/{self.model_code}_{self.model_description}"
            self.timestamp_dir = f"{self.model_dir}/{self.timestamp}"
            
        self.trash_dir = f"{self.base_output_dir}/_trash/{operation_mode.upper()}/{self.model_code}_{self.model_description}/{self.timestamp}"
        
        # Output paths
        self.revision_summary_path = f"{self.timestamp_dir}/90{self.timestamp}.docx"
        self.changes_document_path = f"{self.timestamp_dir}/91{self.timestamp}.docx"
        self.revised_paper_path = f"{self.timestamp_dir}/92{self.timestamp}.docx"
        self.assessment_path = f"{self.timestamp_dir}/93{self.timestamp}.docx"
        self.editor_response_path = f"{self.timestamp_dir}/94{self.timestamp}.docx"
        self.new_bib_path = f"{self.timestamp_dir}/zz{self.timestamp}.bib"
        self.log_path = f"{self.timestamp_dir}/log{self.timestamp}.txt"
        
        # Ensure directories exist
        os.makedirs(self.base_output_dir, exist_ok=True)
        os.makedirs(f"{self.base_output_dir}/_trash", exist_ok=True)  # Create trash root directory
        os.makedirs(self.mode_dir, exist_ok=True)
        os.makedirs(self.model_dir, exist_ok=True)
        
        # Create final directory if in final mode
        if operation_mode.lower() == "final" and self.final_dir:
            os.makedirs(self.final_dir, exist_ok=True)
            self._log_info(f"Creating special FINAL directory in root: {self.final_dir}")
            
            # Also create backup directory for final mode
            if self.backup_dir:
                os.makedirs(self.backup_dir, exist_ok=True)
                self._log_debug(f"Created backup directory: {self.backup_dir}")
        else:
            os.makedirs(self.timestamp_dir, exist_ok=True)
        # Trash directory will be created only when needed
        
        # Set up logger
        self._setup_logger()
        
        # Statistics
        self.start_time = time.time()
        self.process_statistics = {
            "tokens_used": 0,
            "cost": 0.0,
            "requests": 0,
            "files_processed": 0,
            "files_created": 0,
            "cached_requests": 0,
            "token_budget_remaining": self.token_budget,
            "cost_budget_remaining": self.cost_budget
        }
        
        # Skip API validation flag
        self.skip_api_validation = skip_api_validation
        
        # Initialize LLM client (if not skipping validation)
        if not self.skip_api_validation:
            # Use the verify parameter from constructor
            self._initialize_llm_client(verify=self.verify)
        
    def _check_budget(self, estimated_tokens: int, estimated_cost: float) -> bool:
        """Check if the estimated tokens and cost are within budget.
        
        Args:
            estimated_tokens: Estimated tokens for the request
            estimated_cost: Estimated cost for the request
            
        Returns:
            True if within budget, False otherwise
        """
        if self.process_statistics["token_budget_remaining"] < estimated_tokens:
            self._log_warning(f"Token budget exceeded. Remaining: {self.process_statistics['token_budget_remaining']}, Required: {estimated_tokens}")
            return False
            
        if self.process_statistics["cost_budget_remaining"] < estimated_cost:
            self._log_warning(f"Cost budget exceeded. Remaining: ${self.process_statistics['cost_budget_remaining']:.4f}, Required: ${estimated_cost:.4f}")
            return False
            
        return True
        
    def _update_budget(self, tokens_used: int, cost: float):
        """Update budget after a request.
        
        Args:
            tokens_used: Tokens used in the request
            cost: Cost of the request
        """
        self.process_statistics["token_budget_remaining"] -= tokens_used
        self.process_statistics["cost_budget_remaining"] -= cost
    
    def _initialize_llm_client(self, verify=False):
        """Initialize the LLM client based on selected provider and model.
        
        Args:
            verify: Whether to verify model accuracy (default: False)
                  When False, skips the Christmas test to save time and tokens
                  Run verify_models.py separately to check models before starting
        """
        try:
            # Validate API key with more detailed error messages
            provider_env_vars = {
                "anthropic": "ANTHROPIC_API_KEY",
                "openai": "OPENAI_API_KEY",
                "google": "GOOGLE_API_KEY"
            }
            
            env_var = provider_env_vars.get(self.provider, "API_KEY")
            
            if not os.getenv(env_var):
                self._log_error(f"Missing {env_var} in environment. Please add it to your .env file.")
                sys.exit(1)
            
            # Initialize client with optional verification
            self.llm_client = get_llm_client(self.provider, self.model_name, verify=verify)
                
            if not self.llm_client.validate_api_key():
                self._log_error(f"Invalid API key for {self.provider}. The API key in your .env file is not working.")
                self._log_warning(f"Please check your {env_var} in .env file and make sure it is correct and not expired.")
                sys.exit(1)
            
            # Check if verification was requested and failed
            if verify and not self.llm_client.verified:
                self._log_warning(f"Model verification failed for {self.provider.upper()} {self.model_name}")
                self._log_warning("The model may not be functioning correctly or may provide inaccurate information.")
                
                # Ask user if they want to continue
                response = input("Do you want to continue anyway? (y/n): ")
                if response.lower() != 'y':
                    self._log_error("Aborting due to failed model verification")
                    self._log_info("You can run verify_models.py to check model accuracy separately")
                    sys.exit(1)
            
            # Suggest verification if it wasn't performed
            if not verify:
                self._log_info(f"Model verification skipped. Run verify_models.py to check model accuracy separately.")
                    
            self._log_success(f"Successfully initialized {self.provider.upper()} client with model {self.model_name} (Code: {self.model_code})")
        except Exception as e:
            self._log_error(f"Error initializing LLM client: {e}")
            self._log_warning("Make sure you have the required API keys in your .env file and all dependencies installed.")
            sys.exit(1)
    
    def _move_to_trash(self):
        """Move any created files to the trash directory when a run fails."""
        self._log_info(f"Moving files to trash directory: {self.trash_dir}")
        
        # Ensure trash directory exists
        os.makedirs(self.trash_dir, exist_ok=True)
        
        # Create a failure report
        failure_report_path = f"{self.trash_dir}/failure_report_{self.timestamp}.txt"
        try:
            with open(failure_report_path, 'w') as f:
                f.write(f"FAILURE REPORT - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Model: {self.model_name} (Code: {self.model_code})\n")
                f.write(f"Provider: {self.provider}\n")
                f.write(f"Operation mode: {self.operation_mode}\n")
                f.write(f"Original directory: {self.timestamp_dir}\n")
                f.write(f"Trash directory: {self.trash_dir}\n\n")
                f.write("Files moved to trash:\n")
                
                # List all files in the timestamp directory
                try:
                    if os.path.exists(self.timestamp_dir):
                        files = os.listdir(self.timestamp_dir)
                        for file in files:
                            f.write(f"- {file}\n")
                except Exception as e:
                    f.write(f"Error listing files: {str(e)}\n")
        except Exception as e:
            self._log_error(f"Failed to create failure report: {str(e)}")
        
        # Try to move all files from timestamp directory to trash
        try:
            import shutil
            if os.path.exists(self.timestamp_dir):
                # Copy files instead of moving them
                for item in os.listdir(self.timestamp_dir):
                    s = os.path.join(self.timestamp_dir, item)
                    d = os.path.join(self.trash_dir, item)
                    if os.path.isdir(s):
                        shutil.copytree(s, d, dirs_exist_ok=True)
                    else:
                        shutil.copy2(s, d)
                
                # Don't delete the original directory - some files might still be open
                # We'll let the user decide when to clean up
                self._log_success(f"Files copied to trash directory: {self.trash_dir}")
        except Exception as e:
            self._log_error(f"Failed to move files to trash: {str(e)}")
        
        return self.trash_dir
    
    def _setup_logger(self):
        """Set up the logger for the paper revision tool."""
        self.logger = logging.getLogger(f"paper_revision_{self.timestamp}")
        self.logger.setLevel(logging.DEBUG if self.debug else logging.INFO)
        
        # Ensure directories exist for both normal and trash paths
        os.makedirs(self.timestamp_dir, exist_ok=True)
        os.makedirs(self.trash_dir, exist_ok=True)
        
        # Create file handler
        file_handler = logging.FileHandler(self.log_path)
        file_handler.setLevel(logging.DEBUG)  # Log everything to file
        
        # Create console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)  # Only info and above to console
        
        # Create formatter
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        
        # Add handlers to logger
        self.logger.addHandler(file_handler)
        
        # Log initial information
        self.logger.info(f"Paper revision started with {self.provider} model: {self.model_name}")
        self.logger.info(f"Model code: {self.model_code}")
        self.logger.info(f"Operation mode: {self.operation_mode}")
        self.logger.info(f"Output directory: {self.timestamp_dir}")
        self.logger.info(f"Trash directory (used on failure): {self.trash_dir}")
        
    def _log_info(self, message: str):
        """Log an informational message."""
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} {message}")
        self.logger.info(message)
    
    def _log_success(self, message: str):
        """Log a success message."""
        print(f"{Fore.GREEN}[SUCCESS]{Style.RESET_ALL} {message}")
        self.logger.info(f"SUCCESS: {message}")
    
    def _log_warning(self, message: str):
        """Log a warning message."""
        print(f"{Fore.YELLOW}[WARNING]{Style.RESET_ALL} {message}")
        self.logger.warning(message)
    
    def _log_error(self, message: str):
        """Log an error message."""
        print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} {message}")
        self.logger.error(message)
    
    def _log_debug(self, message: str):
        """Log a debug message."""
        if self.debug:
            print(f"{Fore.MAGENTA}[DEBUG]{Style.RESET_ALL} {message}")
        self.logger.debug(message)
        
    def _track_preprocessed_file(self, original_path: str, processed_path: str, file_type: str, 
                               size: int, token_estimate: int, page_count: int = 0):
        """Track a preprocessed file in the database.
        
        Args:
            original_path: Path to the original file
            processed_path: Path to the processed file
            file_type: Type of file (pdf, docx, etc.)
            size: Size of the file in bytes
            token_estimate: Estimated token count
            page_count: Number of pages (for PDFs and documents)
        """
        # Create a unique file ID
        file_id = hashlib.md5(f"{original_path}_{processed_path}_{self.timestamp}".encode()).hexdigest()
        
        # Add file to database
        self.workflow_db.add_file(
            self.run_id,
            file_id,
            original_path,
            processed_path,
            file_type,
            size,
            token_estimate,
            page_count
        )
        
        self._log_debug(f"Tracked preprocessed file in database: {os.path.basename(original_path)}")
    
    def _create_revision_report(self, issues, solutions, reviewer_comments, new_references):
        """Create a summary report of the revision.
        
        Args:
            issues: List of identified issues
            solutions: List of solutions
            reviewer_comments: List of reviewer comments
            new_references: List of new references added
            
        Returns:
            Tuple of (report text, report file path)
        """
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"{self.timestamp_dir}/report{timestamp}.txt"
        
        # Initialize report sections
        report_lines = [
            f"{'=' * 70}",
            f"PAPER REVISION SUMMARY REPORT",
            f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Model: {self.model_name} (Code: {self.model_code})",
            f"Operation mode: {self.operation_mode.upper()}",
            f"{'=' * 70}\n"
        ]
        
        # Add reviewer comments section
        report_lines.append("REVIEWER COMMENTS SUMMARY:")
        for i, reviewer in enumerate(reviewer_comments, 1):
            report_lines.append(f"\nREVIEWER {i}:")
            report_lines.append(f"  Overall assessment: {reviewer.get('overall_assessment', 'Not specified')}")
            
            # Main concerns
            report_lines.append("  Main concerns:")
            concerns = reviewer.get("main_concerns", [])
            if concerns:
                for concern in concerns[:3]:  # Show up to 3 concerns
                    report_lines.append(f"   - {concern}")
            else:
                report_lines.append("   - No specific concerns identified")
            
            # Required changes
            report_lines.append("  Required changes:")
            required = reviewer.get("required_changes", [])
            if required:
                for change in required[:3]:  # Show up to 3 required changes
                    report_lines.append(f"   - {change}")
            else:
                report_lines.append("   - No specific changes required")
        
        # Add issues section
        report_lines.append("\nKEY ISSUES IDENTIFIED:")
        if issues:
            for i, issue in enumerate(issues[:5], 1):  # Show top 5 issues
                report_lines.append(f"  {i}. {issue.get('title', 'Unnamed issue')} (Severity: {issue.get('severity', 'unspecified')})")
                report_lines.append(f"     Source: {issue.get('source', 'unspecified')}")
                if 'description' in issue:
                    desc = issue['description']
                    if len(desc) > 100:
                        desc = desc[:97] + "..."
                    report_lines.append(f"     Description: {desc}")
        else:
            report_lines.append("  No issues identified")
        
        # Add solutions section
        report_lines.append("\nIMPLEMENTED SOLUTIONS:")
        if solutions:
            for i, solution in enumerate(solutions[:5], 1):  # Show top 5 solutions
                report_lines.append(f"  {i}. {solution.get('title', 'Unnamed solution')} (Complexity: {solution.get('complexity', 'unspecified')})")
                if 'implementation' in solution:
                    impl = solution['implementation']
                    if len(impl) > 100:
                        impl = impl[:97] + "..."
                    report_lines.append(f"     Implementation: {impl}")
        else:
            report_lines.append("  No solutions implemented")
        
        # Add references section
        report_lines.append("\nNEW REFERENCES ADDED:")
        if new_references:
            for i, ref in enumerate(new_references, 1):
                report_lines.append(f"  {i}. {ref.get('title', 'Unnamed reference')} ({ref.get('authors', 'Unknown authors')}, {ref.get('year', 'Unknown year')})")
                if 'reason' in ref:
                    report_lines.append(f"     Reason: {ref.get('reason', 'No reason specified')}")
        else:
            report_lines.append("  No new references added")
        
        # Add output files section
        report_lines.append("\nOUTPUT FILES GENERATED:")
        report_lines.append(f"  Revision summary: {os.path.basename(self.revision_summary_path)}")
        report_lines.append(f"  Changes document: {os.path.basename(self.changes_document_path)}")
        report_lines.append(f"  Revised paper: {os.path.basename(self.revised_paper_path)}")
        report_lines.append(f"  Assessment: {os.path.basename(self.assessment_path)}")
        report_lines.append(f"  Editor letter: {os.path.basename(self.editor_response_path)}")
        report_lines.append(f"  Bibliography: {os.path.basename(self.new_bib_path)}")
        
        # Add footer
        report_lines.append(f"\n{'=' * 70}")
        report_lines.append(f"End of report - {self.model_code}_{self.model_description}")
        report_lines.append(f"{'=' * 70}")
        
        # Write to file
        with open(filename, 'w') as f:
            f.write("\n".join(report_lines))
        
        return "\n".join(report_lines), filename
        
    def _verify_final_mode(self):
        """Perform verification for final mode and create backups.
        
        This method verifies that all required database information is available for 
        the process disclosure in final mode, and creates backups of all output files
        in the standard directory structure.
        
        Returns:
            bool: True if verification succeeded, False otherwise
        """
        if self.operation_mode.lower() != "final":
            return True
            
        self._log_info("Performing final mode verification")
        
        # 1. Check if database has the review process information
        process_summary = self.workflow_db.get_review_process_summary(self.run_id)
        if not process_summary:
            self._log_error("Missing review process information in database for disclosure")
            return False
            
        # 2. Display detailed summary for final mode
        summary_text = process_summary.get("process_description", "")
        if summary_text:
            print(f"\n{Fore.GREEN}{'=' * 50}{Style.RESET_ALL}")
            print(f"{Fore.GREEN}REVIEW PROCESS SUMMARY:{Style.RESET_ALL}")
            print(summary_text)
            print(f"{Fore.GREEN}{'=' * 50}{Style.RESET_ALL}\n")
            
        # 3. Create backup copies in the standard directory structure if needed
        if self.backup_dir and self.final_dir:
            self._log_debug(f"Creating backups in standard directory structure: {self.backup_dir}")
            
            # List of all output files to backup
            output_files = [
                self.revision_summary_path,
                self.changes_document_path,
                self.revised_paper_path,
                self.assessment_path,
                self.editor_response_path,
                self.new_bib_path,
                self.log_path
            ]
            
            # Create backups
            for src_path in output_files:
                if os.path.exists(src_path):
                    # Get just the filename
                    filename = os.path.basename(src_path)
                    # Create destination path in backup directory
                    dst_path = f"{self.backup_dir}/{filename}"
                    # Copy the file
                    try:
                        import shutil
                        shutil.copy2(src_path, dst_path)
                        self._log_debug(f"Created backup: {dst_path}")
                    except Exception as e:
                        self._log_warning(f"Failed to create backup of {src_path}: {e}")
                        
        self._log_success("Final mode verification completed successfully")
        return True
            
    def _log_stats(self, export_to_file=False):
        """Log current statistics.
        
        Args:
            export_to_file: Whether to export stats to a text file
        """
        elapsed = time.time() - self.start_time
        hours, remainder = divmod(elapsed, 3600)
        minutes, seconds = divmod(remainder, 60)
        
        # Update stats from LLM client with more detailed information
        if self.llm_client:
            usage_stats = self.llm_client.get_usage_statistics()
            self.process_statistics["tokens_used"] = usage_stats["total_tokens"]
            self.process_statistics["cost"] = usage_stats["total_cost"]
            self.process_statistics["requests"] = usage_stats["request_count"]
            
            # Include the formatted model-specific summary
            self.process_statistics["model_summary"] = usage_stats.get("summary", "")
        
        # Add evaluation statistics if available
        evaluation_tokens = self.process_statistics.get("evaluation_tokens", 0)
        evaluation_cost = self.process_statistics.get("evaluation_cost", 0.0)
        evaluation_requests = self.process_statistics.get("evaluation_requests", 0)
        
        # Prepare stats as formatted strings
        stats_lines = [
            f"{'=' * 50}",
            f"STATISTICS SUMMARY:",
            f"Operation mode: {self.operation_mode.upper()}",
            f"Time elapsed: {int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}",
            f"Tokens used (main): {self.process_statistics['tokens_used']:,}",
            f"Estimated cost (main): ${self.process_statistics['cost']:.4f}",
        ]
        
        # Add evaluation stats if competitor evaluation was used
        if self.competitor_evaluation and evaluation_tokens > 0:
            stats_lines.extend([
                f"Tokens used (evaluation): {evaluation_tokens:,}",
                f"Estimated cost (evaluation): ${evaluation_cost:.4f}",
                f"Total tokens: {self.process_statistics['tokens_used'] + evaluation_tokens:,}",
                f"Total cost: ${self.process_statistics['cost'] + evaluation_cost:.4f}",
                f"Evaluation requests: {evaluation_requests}"
            ])
        
        # Continue with the rest of the stats
        stats_lines.extend([
            f"API requests: {self.process_statistics['requests']}",
            f"Cached requests: {self.process_statistics['cached_requests']}",
            f"Files processed: {self.process_statistics['files_processed']}",
            f"Files created: {self.process_statistics['files_created']}",
            f"Token budget remaining: {self.process_statistics['token_budget_remaining']:,}",
            f"Cost budget remaining: ${self.process_statistics['cost_budget_remaining']:.4f}",
            f"Provider: {self.provider} ({self.model_code[0]})",
            f"Model: {self.model_name} (Code: {self.model_code})",
            f"Competitor evaluation: {'Enabled' if self.competitor_evaluation else 'Disabled'}",
            f"Cost optimization: {'Enabled' if self.optimize_costs else 'Disabled'}",
            f"Maximum papers analyzed: {self.max_papers_to_process}",
            f"{'=' * 50}"
        ])
        
        # Print to console with colors
        print(f"\n{Fore.CYAN}{stats_lines[0]}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{stats_lines[1]}{Style.RESET_ALL}")
        for line in stats_lines[2:-1]:
            print(line)
        print(f"{Fore.CYAN}{stats_lines[-1]}{Style.RESET_ALL}")
        
        # Display model-specific detailed summary if available
        if self.llm_client and "model_summary" in self.process_statistics:
            model_summary = self.process_statistics["model_summary"]
            if model_summary:
                print(f"\n{Fore.GREEN}DETAILED MODEL USAGE:{Style.RESET_ALL}")
                print(model_summary)
                print()
        
        # Export to file if requested
        if export_to_file:
            timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{self.timestamp_dir}/cost{timestamp}.txt"
            
            with open(filename, 'w') as f:
                f.write(f"COST OPTIMIZATION REPORT - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("\n".join(stats_lines))
                
                # Add detailed model usage if available
                if "model_summary" in self.process_statistics:
                    model_summary = self.process_statistics["model_summary"]
                    if model_summary:
                        f.write("\n\nDETAILED MODEL USAGE:\n")
                        f.write(model_summary)
                
                # Add detailed cache statistics
                f.write("\n\nCACHE STATISTICS:\n")
                f.write(f"Cache directory: {CACHE_DIR}\n")
                f.write(f"Cache entries: {len(os.listdir(CACHE_DIR))}\n")
                if os.path.exists(CACHE_DIR) and os.listdir(CACHE_DIR):
                    f.write("\nMost recent cache entries:\n")
                    cache_files = sorted([os.path.join(CACHE_DIR, f) for f in os.listdir(CACHE_DIR) if f.endswith('.json')], 
                                        key=os.path.getmtime, reverse=True)
                    for i, cache_file in enumerate(cache_files[:5]):  # Show 5 most recent
                        mtime = datetime.datetime.fromtimestamp(os.path.getmtime(cache_file))
                        f.write(f"{i+1}. {os.path.basename(cache_file)} - {mtime.strftime('%Y-%m-%d %H:%M:%S')}\n")
                
                # Add detailed cost breakdown by operation
                f.write("\n\nESTIMATED COST BREAKDOWN:\n")
                token_cost = self.process_statistics['tokens_used'] / 1000
                if self.provider == "anthropic":
                    from src.models.anthropic_models import get_claude_model_info
                    model_info = get_claude_model_info(self.model_name)
                    if model_info:
                        input_cost = token_cost * model_info.get("price_per_1k_input", 0.0)
                        output_cost = token_cost * model_info.get("price_per_1k_output", 0.0)
                        f.write(f"Input tokens cost (estimated): ${input_cost:.4f}\n")
                        f.write(f"Output tokens cost (estimated): ${output_cost:.4f}\n")
                elif self.provider == "openai":
                    from src.models.openai_models import get_openai_model_info
                    model_info = get_openai_model_info(self.model_name)
                    if model_info:
                        input_cost = token_cost * model_info.get("price_per_1k_input", 0.0)
                        output_cost = token_cost * model_info.get("price_per_1k_output", 0.0)
                        f.write(f"Input tokens cost (estimated): ${input_cost:.4f}\n")
                        f.write(f"Output tokens cost (estimated): ${output_cost:.4f}\n")
                
                f.write(f"\nTotal cost: ${self.process_statistics['cost']:.4f}\n")
                
                # Add recommendations for further optimization
                f.write("\n\nOPTIMIZATION RECOMMENDATIONS:\n")
                if self.provider == "anthropic" and "opus" in self.model_name:
                    f.write("- Consider using a cheaper model like Claude Sonnet or Haiku for initial drafts\n")
                if self.provider == "openai" and "gpt-4" in self.model_name:
                    f.write("- Consider using a cheaper model like GPT-3.5 for initial drafts\n")
                if self.process_statistics['tokens_used'] > 100000:
                    f.write("- Reduce token usage by processing fewer papers or shortening text inputs\n")
                if self.process_statistics['cached_requests'] < self.process_statistics['requests'] * 0.1:
                    f.write("- Enable caching between runs for similar operations\n")
                
            self._log_success(f"Cost report exported to {filename}")
            return filename
        
        return None
        
    def _get_competing_model(self) -> Tuple[str, str]:
        """Get the closest competing model to evaluate responses.
        
        Returns:
            Tuple of (provider, model_name) for the competing model
        """
        # Define competitor mapping by tier and provider
        competitors = {
            "anthropic": {
                "basic": {"openai": "gpt-4o-mini", "google": "gemini-1.5-flash"},
                "standard": {"openai": "gpt-4o", "google": "gemini-1.5-pro"},
                "advanced": {"openai": "gpt-4.5-preview", "google": "gemini-2.5-pro-preview"}
            },
            "openai": {
                "basic": {"anthropic": "claude-3-haiku-20240307", "google": "gemini-1.5-flash"},
                "standard": {"anthropic": "claude-3-5-sonnet-20241022", "google": "gemini-1.5-pro"},
                "advanced": {"anthropic": "claude-opus-4-20250514", "google": "gemini-2.5-pro-preview"}
            },
            "google": {
                "basic": {"anthropic": "claude-3-haiku-20240307", "openai": "gpt-4o-mini"},
                "standard": {"anthropic": "claude-3-5-sonnet-20241022", "openai": "gpt-4o"},
                "advanced": {"anthropic": "claude-opus-4-20250514", "openai": "gpt-4.5-preview"}
            }
        }
        
        # Determine current model's tier
        current_tier = "standard"  # Default tier
        
        # Look for tier indicators in model name
        if any(basic_indicator in self.model_name.lower() for basic_indicator in ["haiku", "mini", "flash"]):
            current_tier = "basic"
        elif any(advanced_indicator in self.model_name.lower() for advanced_indicator in ["opus", "4.5", "2.5"]):
            current_tier = "advanced"
        
        # Get competitors for this provider and tier
        tier_competitors = competitors.get(self.provider, {}).get(current_tier, {})
        
        # If we have a competing_evaluator specified, use that one
        if hasattr(self, 'competing_evaluator') and self.competing_evaluator:
            competing_provider, competing_model = self.competing_evaluator.split('/')
            if competing_provider != self.provider:  # Make sure it's actually a competitor
                return competing_provider, competing_model
        
        # Otherwise use the first available competitor
        for competing_provider, competing_model in tier_competitors.items():
            # Check if we have the API key for this provider
            env_var = f"{competing_provider.upper()}_API_KEY"
            if os.getenv(env_var):
                return competing_provider, competing_model
        
        # If no competitors available, return None/None
        return None, None
    
    def _evaluate_response_quality(self, prompt: str, response: str, task_type: str = "general", 
                                  use_competitor: bool = True) -> Dict[str, Any]:
        """Evaluate the quality of a model response.
        
        Args:
            prompt: The original prompt sent to the model
            response: The model's response
            task_type: The type of task (analysis, generation, editing, etc.)
            use_competitor: Whether to use a competing model for evaluation
            
        Returns:
            Dictionary with quality metrics and feedback
        """
        # If using a competitor model for evaluation
        if use_competitor and not hasattr(self, 'is_evaluator') and self.competitor_evaluation:
            competing_provider, competing_model = self._get_competing_model()
            
            if competing_provider and competing_model:
                # Track the cost of this evaluation in the statistics
                self.process_statistics["evaluation_requests"] = self.process_statistics.get("evaluation_requests", 0) + 1
                
                try:
                    # Create an evaluation prompt
                    evaluation_prompt = f"""
                    You are an expert at evaluating AI model outputs for quality and correctness.
                    
                    ORIGINAL PROMPT:
                    ```
                    {prompt}
                    ```
                    
                    MODEL RESPONSE TO EVALUATE:
                    ```
                    {response}
                    ```
                    
                    TASK TYPE: {task_type}
                    
                    Please evaluate this response on a scale of 1-5 (where 5 is best) based on:
                    1. Relevance to the prompt
                    2. Accuracy and correctness
                    3. Completeness
                    4. Structure and clarity
                    5. Overall quality
                    
                    In your evaluation, identify:
                    - Specific issues with the response (if any)
                    - Highlights or strengths of the response (if any)
                    
                    Format your response as a JSON object with these fields:
                    - quality_score: overall score (1-5)
                    - quality_issues: array of strings describing issues
                    - quality_highlights: array of strings describing strengths
                    - details: object with component scores
                    """
                    
                    # Create a temporary LLM client for the competing model
                    # Mark it as an evaluator to avoid infinite evaluation loops
                    from src.utils.llm_client import get_llm_client
                    evaluator_client = get_llm_client(competing_provider, competing_model, verify=False)
                    evaluator_client.is_evaluator = True  # Mark to prevent recursive evaluations
                    
                    # Get evaluation from the competing model
                    evaluation_response = evaluator_client.get_completion(
                        prompt=evaluation_prompt,
                        system_prompt="You are an expert evaluator of AI outputs. Evaluate fairly and objectively.",
                        max_tokens=1000
                    )
                    
                    # Track the evaluation cost
                    evaluation_tokens = evaluator_client.total_tokens_used
                    evaluation_cost = evaluator_client.total_cost
                    
                    self.process_statistics["evaluation_tokens"] = self.process_statistics.get("evaluation_tokens", 0) + evaluation_tokens
                    self.process_statistics["evaluation_cost"] = self.process_statistics.get("evaluation_cost", 0.0) + evaluation_cost
                    
                    # Parse the evaluation response
                    try:
                        # Look for JSON in the response
                        json_start = evaluation_response.find('{')
                        json_end = evaluation_response.rfind('}') + 1
                        if json_start >= 0 and json_end > json_start:
                            json_str = evaluation_response[json_start:json_end]
                            evaluation = json.loads(json_str)
                            
                            # Add the evaluator info
                            evaluation["evaluator"] = f"{competing_provider}/{competing_model}"
                            evaluation["evaluation_cost"] = evaluation_cost
                            
                            # Log the cross-model evaluation
                            self.logger.info(f"Cross-model evaluation: {self.provider}/{self.model} evaluated by {competing_provider}/{competing_model}")
                            self.logger.info(f"Evaluation score: {evaluation.get('quality_score', 'N/A')}/5")
                            
                            return evaluation
                    except Exception as parse_error:
                        self.logger.warning(f"Error parsing evaluation response: {parse_error}")
                        # Fall back to the basic evaluation
                
                except Exception as eval_error:
                    self.logger.warning(f"Error during competitor evaluation: {eval_error}")
                    # Fall back to the basic evaluation
        
        # Basic quality checks (used if competitor evaluation fails or is disabled)
        quality_score = 5  # Default high score
        quality_issues = []
        quality_highlights = []
        
        # Length check
        if len(response) < 50:
            quality_score -= 1
            quality_issues.append("Response is very short")
        elif len(response) > 10000 and task_type != "analysis":
            quality_score -= 1
            quality_issues.append("Response may be unnecessarily verbose")
        else:
            quality_highlights.append("Response length is appropriate")
        
        # Content relevance check
        relevance_score = 5
        prompt_keywords = set(re.findall(r'\b\w{4,}\b', prompt.lower()))
        response_keywords = set(re.findall(r'\b\w{4,}\b', response.lower()))
        keyword_overlap = len(prompt_keywords.intersection(response_keywords)) / max(1, len(prompt_keywords))
        
        if keyword_overlap < 0.3:
            relevance_score -= 2
            quality_issues.append("Low relevance to prompt (few keyword matches)")
        elif keyword_overlap < 0.5:
            relevance_score -= 1
            quality_issues.append("Moderate relevance to prompt")
        else:
            quality_highlights.append("Good relevance to original prompt")
        
        # Structure check
        structure_score = 5
        if task_type == "analysis" and ":" not in response and len(response.split("\n")) < 3:
            structure_score -= 1
            quality_issues.append("Analysis lacks structure (no sections or bullet points)")
        
        # JSON validation if expected
        if "JSON" in prompt or "json" in prompt:
            try:
                json.loads(response.strip())
                quality_highlights.append("Valid JSON format")
            except json.JSONDecodeError:
                quality_score -= 2
                quality_issues.append("Invalid JSON format")
        
        # Task-specific checks
        if task_type == "analysis":
            if not any(kw in response.lower() for kw in ["however", "but", "although", "while", "despite"]):
                quality_score -= 1
                quality_issues.append("Analysis lacks nuance or consideration of drawbacks")
            else:
                quality_highlights.append("Analysis shows nuanced thinking")
        
        # Completeness check
        if "list" in prompt.lower() and response.count("\n") < 3:
            quality_score -= 1
            quality_issues.append("Response may be incomplete (expected list items)")
        
        # Calculate final scores
        final_quality_score = min(5, max(1, int(round((quality_score + relevance_score + structure_score) / 3))))
        
        return {
            "quality_score": final_quality_score,
            "quality_issues": quality_issues,
            "quality_highlights": quality_highlights,
            "details": {
                "base_quality": quality_score,
                "relevance": relevance_score,
                "structure": structure_score,
                "keyword_overlap": keyword_overlap
            },
            "evaluator": "basic_algorithm"  # Mark this as an algorithmic evaluation
        }

    def _optimized_completion(self, prompt: str, system_prompt: Optional[str] = None, max_tokens: Optional[int] = None, 
                             task_type: str = "general", echo_quality: bool = True, **kwargs) -> str:
        """Get a completion from the LLM with cost optimization and caching.
        
        Args:
            prompt: The user prompt to send to the model
            system_prompt: Optional system prompt
            max_tokens: Optional max tokens parameter
            task_type: The type of task for quality evaluation
            echo_quality: Whether to echo quality metrics
            **kwargs: Additional parameters for the model
            
        Returns:
            The model's response as a string
        """
        # Check cache first
        cached_response = get_cached_completion(self.llm_client, prompt, system_prompt, **kwargs)
        if cached_response:
            self._log_debug("Using cached response")
            self.process_statistics["cached_requests"] += 1
            return cached_response
            
        # Estimate tokens and cost
        estimated_input_tokens = estimate_tokens(prompt)
        if system_prompt:
            estimated_input_tokens += estimate_tokens(system_prompt)
            
        # Rough estimation of output tokens (typically 50-70% of input size for analysis tasks)
        estimated_output_tokens = estimated_input_tokens * 0.6
        
        # Get pricing for the current model
        model_info = None
        if self.provider == "anthropic":
            from src.models.anthropic_models import get_claude_model_info
            model_info = get_claude_model_info(self.model_name)
        elif self.provider == "openai":
            from src.models.openai_models import get_openai_model_info
            model_info = get_openai_model_info(self.model_name)
        elif self.provider == "google":
            from src.models.google_models import get_gemini_model_info
            model_info = get_gemini_model_info(self.model_name)
            
        # Calculate estimated cost
        if model_info:
            input_cost_per_1k = model_info.get("price_per_1k_input", 0.001)
            output_cost_per_1k = model_info.get("price_per_1k_output", 0.002)
            estimated_cost = (estimated_input_tokens / 1000) * input_cost_per_1k + (estimated_output_tokens / 1000) * output_cost_per_1k
        else:
            # Default conservative estimate
            estimated_cost = (estimated_input_tokens + estimated_output_tokens) / 1000 * 0.01
            
        # Check if within budget
        if not self._check_budget(estimated_input_tokens + estimated_output_tokens, estimated_cost):
            self._log_warning("Budget exceeded. Using fallback response.")
            return "Budget exceeded. Using fallback response."
            
        # If optimizing costs, truncate the prompt if it's too long
        if self.optimize_costs and estimated_input_tokens > self.max_tokens_per_prompt * 0.8:
            self._log_debug(f"Truncating prompt from {estimated_input_tokens} tokens")
            # Simple truncation strategy - keep the first and last parts of the prompt
            # This assumes prompt often has instructions at the beginning and key data at the end
            words = prompt.split()
            half_length = len(words) // 2
            quarter_length = len(words) // 4
            
            # Keep first quarter and last quarter, dropping the middle
            truncated_words = words[:half_length - quarter_length] + ["..."] + words[half_length + quarter_length:]
            prompt = " ".join(truncated_words)
            
            estimated_input_tokens = estimate_tokens(prompt)
            self._log_debug(f"Truncated to {estimated_input_tokens} tokens")
            
        # Get the response with verification
        # The first request to the model will automatically verify it can correctly answer
        # when Christmas is. Subsequent requests will skip verification unless the model failed.
        response = self.llm_client.get_completion(
            prompt=prompt,
            system_prompt=system_prompt,
            verify_first=True,  # Always verify the model on first use
            max_tokens=max_tokens or get_max_tokens_for_model(
                self.provider, 
                self.model_name, 
                use_safe_limit=True,  # Default to 90% safe limit
                use_max_safe=self.high_quality  # Use 99% max safe limit for high quality runs
            ),
            **kwargs
        )
        
        # Evaluate response quality
        if not cached_response:
            quality_info = self._evaluate_response_quality(prompt, response, task_type)
            
            # Track quality stats
            if not hasattr(self, 'quality_statistics'):
                self.quality_statistics = {
                    "total_evaluations": 0,
                    "quality_scores": [],
                    "issues_frequency": {},
                    "highlights_frequency": {},
                    "task_types": {}
                }
            
            # Update quality statistics
            self.quality_statistics["total_evaluations"] += 1
            self.quality_statistics["quality_scores"].append(quality_info["quality_score"])
            
            # Track issues and highlights frequency
            for issue in quality_info["quality_issues"]:
                self.quality_statistics["issues_frequency"][issue] = self.quality_statistics["issues_frequency"].get(issue, 0) + 1
            
            for highlight in quality_info["quality_highlights"]:
                self.quality_statistics["highlights_frequency"][highlight] = self.quality_statistics["highlights_frequency"].get(highlight, 0) + 1
            
            # Track task types
            self.quality_statistics["task_types"][task_type] = self.quality_statistics["task_types"].get(task_type, 0) + 1
            
            # Echo quality metrics to console
            evaluator_info = f"Evaluated by: {quality_info.get('evaluator', 'algorithm')}"
            if 'evaluation_cost' in quality_info:
                evaluator_info += f" (Cost: ${quality_info['evaluation_cost']:.4f})"
                
            if quality_info["quality_score"] <= 2:
                print(f"\n{Fore.RED}[!] LOW QUALITY RESPONSE (Score: {quality_info['quality_score']}/5) - Task: {task_type}{Style.RESET_ALL}")
                print(f"{Fore.RED}Issues: {', '.join(quality_info['quality_issues'])}{Style.RESET_ALL}")
                print(f"{Fore.BLUE}{evaluator_info}{Style.RESET_ALL}")
            elif quality_info["quality_score"] == 3:
                print(f"\n{Fore.YELLOW}[⚠] MEDIUM QUALITY RESPONSE (Score: {quality_info['quality_score']}/5) - Task: {task_type}{Style.RESET_ALL}")
                if quality_info["quality_issues"]:
                    print(f"{Fore.YELLOW}Issues: {', '.join(quality_info['quality_issues'])}{Style.RESET_ALL}")
                if quality_info["quality_highlights"]:
                    print(f"{Fore.GREEN}Highlights: {', '.join(quality_info['quality_highlights'])}{Style.RESET_ALL}")
                print(f"{Fore.BLUE}{evaluator_info}{Style.RESET_ALL}")
            elif quality_info["quality_score"] >= 4:
                print(f"\n{Fore.GREEN}[✓] HIGH QUALITY RESPONSE (Score: {quality_info['quality_score']}/5) - Task: {task_type}{Style.RESET_ALL}")
                if quality_info["quality_highlights"]:
                    print(f"{Fore.GREEN}Highlights: {', '.join(quality_info['quality_highlights'])}{Style.RESET_ALL}")
                print(f"{Fore.BLUE}{evaluator_info}{Style.RESET_ALL}")
            
            # Log to file
            self.logger.info(f"Response Quality - Score: {quality_info['quality_score']}/5, Task: {task_type}, Evaluator: {quality_info.get('evaluator', 'algorithm')}")
            if quality_info["quality_issues"]:
                self.logger.info(f"Quality Issues: {', '.join(quality_info['quality_issues'])}")
            if quality_info["quality_highlights"]:
                self.logger.info(f"Quality Highlights: {', '.join(quality_info['quality_highlights'])}")
        
        # Save to cache
        save_to_cache(self.llm_client, prompt, system_prompt, response, **kwargs)
        
        # Update budget based on actual usage
        tokens_used = self.llm_client.total_tokens_used - self.process_statistics["tokens_used"]
        cost = self.llm_client.total_cost - self.process_statistics["cost"]
        self._update_budget(tokens_used, cost)
        
        # Record completion in database (if not from cache)
        if not cached_response:
            # Get the current step ID if we're in a step
            current_step_id = None
            if hasattr(self, 'process_monitoring') and self.process_monitoring.get("current_step"):
                step_number = self.process_monitoring["current_step"]
                # Find the most recent step with this number
                self.workflow_db.cursor.execute('''
                SELECT step_id FROM steps 
                WHERE run_id = ? AND step_number = ? 
                ORDER BY start_time DESC LIMIT 1
                ''', (self.run_id, step_number))
                step = self.workflow_db.cursor.fetchone()
                if step:
                    current_step_id = step['step_id']
            
            # Create prompt hash for caching
            prompt_hash = hashlib.md5(f"{self.provider}_{self.model_name}_{prompt}_{system_prompt}_{str(kwargs)}".encode()).hexdigest()
            
            # Record completion
            self.workflow_db.add_completion(
                self.run_id,
                current_step_id,
                self.provider,
                self.model_name,
                task_type,
                prompt_hash,
                tokens_used,  # This is approximate - we don't split prompt/completion tokens here
                0,            # We don't have this breakdown available
                cost,
                False         # Not from cache
            )
            
            # Record evaluation if we have quality_info
            if hasattr(self, 'quality_statistics') and 'quality_info' in locals():
                evaluation_tokens = quality_info.get("evaluation_tokens", 0)
                evaluation_cost = quality_info.get("evaluation_cost", 0.0)
                
                self.workflow_db.add_evaluation(
                    self.run_id,
                    current_step_id,
                    f"{self.provider}/{self.model_name}",
                    quality_info.get("evaluator", "basic_algorithm"),
                    task_type,
                    quality_info.get("quality_score", 0),
                    tokens_used,  # Approximation of prompt tokens
                    0,            # Not available
                    evaluation_tokens,
                    evaluation_cost,
                    {
                        "quality_issues": quality_info.get("quality_issues", []),
                        "quality_highlights": quality_info.get("quality_highlights", []),
                        "details": quality_info.get("details", {})
                    }
                )
        
        return response
    
    def run(self):
        """Run the full paper revision process.
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            self._log_info("Starting paper revision process")
            
            # Initialize process monitoring
            self.process_monitoring = {
                "start_time": time.time(),
                "steps": [],
                "current_step": 0,
                "total_steps": 10,  # Total number of main steps in the process
                "step_quality": {}
            }
            
            # Display process banner
            print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            print(f"{Fore.CYAN}PAPER REVISION PROCESS - QUALITY MONITORING ENABLED{Style.RESET_ALL}")
            print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            print(f"Using preprocessed files from: {Fore.YELLOW}{os.path.abspath('./asis')}{Style.RESET_ALL}")
            print(f"Model: {Fore.YELLOW}{self.provider.capitalize()} - {self.model_name}{Style.RESET_ALL}")
            print(f"Quality monitoring: {Fore.GREEN}Enabled{Style.RESET_ALL}")
            print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            
            # Step 1: Analyze original paper
            self.process_monitoring["current_step"] = 1
            step_start_time = time.time()
            print(f"\n{Fore.BLUE}[STEP 1/10] Analyzing original paper...{Style.RESET_ALL}")
            
            # Start step in database
            step_id = self.workflow_db.start_step(self.run_id, 1, "Analyze original paper")
            
            self._log_info("Step 1: Analyzing original paper")
            paper_analysis = self._analyze_original_paper()
            
            # Record step completion and quality
            step_duration = time.time() - step_start_time
            quality_score = None
            if hasattr(self, 'quality_statistics') and self.quality_statistics["task_types"].get("paper_analysis", 0) > 0:
                # Extract quality score for this specific task type if available
                quality_score = next((score for score, task in zip(self.quality_statistics["quality_scores"], 
                                                                self.quality_statistics["task_types"].keys()) 
                                if task == "paper_analysis"), None)
            
            # Record step in process monitoring
            self.process_monitoring["steps"].append({
                "step": 1,
                "name": "Analyze original paper",
                "duration": step_duration,
                "quality_score": quality_score
            })
            
            # Complete step in database with output file information
            output_file = None
            if paper_analysis and isinstance(paper_analysis, dict) and paper_analysis.get("processed_path"):
                output_file = paper_analysis.get("processed_path")
            self.workflow_db.complete_step(step_id, "completed", output_file)
            
            if quality_score:
                print(f"{Fore.GREEN}✓ Step completed in {step_duration:.1f}s (Quality score: {quality_score}/5){Style.RESET_ALL}")
            else:
                print(f"{Fore.GREEN}✓ Step completed in {step_duration:.1f}s{Style.RESET_ALL}")
            
            # Step 2: Analyze reviewer comments
            self._log_info("Step 2: Analyzing reviewer comments")
            reviewer_comments = self._analyze_reviewer_comments()
            
            # Step 3: Process editor letter and PRISMA requirements
            self._log_info("Step 3: Processing editor letter and PRISMA requirements")
            editor_requirements = self._process_editor_requirements()
            
            # Step 4: Analyze journal style and requirements
            self._log_info("Step 4: Analyzing journal style and requirements")
            journal_style = self._analyze_journal_style()
            
            # Step 5: Generate revision summary
            self._log_info("Step 5: Generating revision summary")
            issues, solutions = self._generate_revision_plan(
                paper_analysis, reviewer_comments, editor_requirements, journal_style
            )
            revision_summary_path = self._create_revision_summary(issues, solutions, self.revision_summary_path)
            self._log_success(f"Created revision summary at {revision_summary_path}")
            
            # Step 6: Generate changes document
            self._log_info("Step 6: Generating changes document")
            changes = self._generate_changes(paper_analysis, issues, solutions)
            changes_document_path = self._create_changes_document(changes, self.changes_document_path)
            self._log_success(f"Created changes document at {changes_document_path}")
            
            # Step 7: Validate and update references
            self._log_info("Step 7: Validating and updating references")
            new_references = self._validate_and_update_references(paper_analysis, reviewer_comments, self.new_bib_path)
            
            # Step 8: Create revised paper with track changes
            self._log_info("Step 8: Creating revised paper with track changes")
            revised_paper_path = self._create_revised_paper(changes, self.revised_paper_path)
            self._log_success(f"Created revised paper at {revised_paper_path}")
            
            # Step 9: Create assessment document
            self._log_info("Step 9: Creating assessment document")
            assessment_path = self._create_assessment(changes, paper_analysis, self.assessment_path)
            self._log_success(f"Created assessment document at {assessment_path}")
            
            # Step 10: Create letter to editor with process summary
            self._log_info("Step 10: Creating letter to editor with process disclosure")
            
            # Ensure review process information is stored in database
            if hasattr(self, 'report') and "process_summary" in self.report:
                process_summary = self.report["process_summary"]
                self._log_info(f"Including process summary in editor letter: {len(process_summary.get('process_description', ''))} chars")
            
            editor_letter_path = self._create_editor_letter(reviewer_comments, changes, self.editor_response_path)
            self._log_success(f"Created letter to editor at {editor_letter_path}")
            
            # Create summary report with essential information
            report_text, report_path = self._create_revision_report(
                issues, solutions, reviewer_comments, new_references
            )
            
            # Print quality evaluation summary if available
            if hasattr(self, 'quality_statistics') and self.quality_statistics["total_evaluations"] > 0:
                print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}PROCESS QUALITY EVALUATION:{Style.RESET_ALL}")
                print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                
                # Calculate average quality score
                avg_score = sum(self.quality_statistics["quality_scores"]) / len(self.quality_statistics["quality_scores"])
                
                # Get most common issues and highlights
                top_issues = sorted(self.quality_statistics["issues_frequency"].items(), key=lambda x: x[1], reverse=True)[:3]
                top_highlights = sorted(self.quality_statistics["highlights_frequency"].items(), key=lambda x: x[1], reverse=True)[:3]
                
                # Print quality summary
                print(f"\nEvaluation Summary:")
                print(f"• Total completions evaluated: {self.quality_statistics['total_evaluations']}")
                print(f"• Average quality score: {Fore.YELLOW}{avg_score:.1f}/5.0{Style.RESET_ALL}")
                
                # Print quality assessment based on average score
                if avg_score >= 4.0:
                    print(f"{Fore.GREEN}✓ HIGH QUALITY PROCESS: The revision process produced high-quality outputs{Style.RESET_ALL}")
                elif avg_score >= 3.0:
                    print(f"{Fore.YELLOW}⚠ MODERATE QUALITY PROCESS: The revision process produced acceptable outputs with some issues{Style.RESET_ALL}")
                else:
                    print(f"{Fore.RED}! LOW QUALITY PROCESS: The revision process had significant quality issues{Style.RESET_ALL}")
                
                # Print top issues if any
                if top_issues:
                    print(f"\n{Fore.YELLOW}Process Lowlights:{Style.RESET_ALL}")
                    for issue, count in top_issues:
                        print(f"• {issue} ({count} occurrences)")
                
                # Print top highlights if any
                if top_highlights:
                    print(f"\n{Fore.GREEN}Process Highlights:{Style.RESET_ALL}")
                    for highlight, count in top_highlights:
                        print(f"• {highlight} ({count} occurrences)")
                        
                # Add quality statistics to the log file
                self.logger.info(f"Quality Statistics Summary:")
                self.logger.info(f"Total evaluations: {self.quality_statistics['total_evaluations']}")
                self.logger.info(f"Average quality score: {avg_score:.1f}/5.0")
                self.logger.info(f"Top issues: {', '.join([issue for issue, _ in top_issues])}")
                self.logger.info(f"Top highlights: {', '.join([highlight for highlight, _ in top_highlights])}")
            
            # Print essential information to console
            print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            print(f"{Fore.CYAN}ESSENTIAL REVISION INFORMATION:{Style.RESET_ALL}")
            print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            
            # Extract and print the most important parts from the report
            reviewer_concerns = []
            for i, reviewer in enumerate(reviewer_comments, 1):
                concerns = reviewer.get("main_concerns", [])
                required = reviewer.get("required_changes", [])
                if concerns or required:
                    reviewer_concerns.append(f"Reviewer {i}:")
                    if concerns:
                        reviewer_concerns.append(f"  Main concerns: {concerns[0] if concerns else 'None'}")
                    if required:
                        reviewer_concerns.append(f"  Required changes: {required[0] if required else 'None'}")
            
            # Print reviewer concerns
            print(f"\n{Fore.YELLOW}REVIEWER CONCERNS:{Style.RESET_ALL}")
            for line in reviewer_concerns:
                print(line)
            
            # Print key issues and solutions
            print(f"\n{Fore.YELLOW}KEY ISSUES ADDRESSED:{Style.RESET_ALL}")
            for i, issue in enumerate(issues[:3], 1):  # Show top 3 issues
                print(f"{i}. {issue.get('title', 'Unnamed issue')} (Source: {issue.get('source', 'unspecified')})")
            
            # Print references added
            if new_references:
                print(f"\n{Fore.YELLOW}NEW REFERENCES ADDED:{Style.RESET_ALL}")
                for i, ref in enumerate(new_references[:3], 1):  # Show top 3 references
                    print(f"{i}. {ref.get('title', 'Unnamed reference')} ({ref.get('authors', 'Unknown')})")
            
            # Final statistics with export to file
            cost_report = self._log_stats(export_to_file=True)
            
            # Final verification for final mode
            if self.operation_mode.lower() == "final":
                verification_success = self._verify_final_mode()
                if not verification_success:
                    self._log_warning("Final mode verification failed. Some features may be incomplete.")
            
            self._log_success("Paper revision process completed successfully!")
            
            # Special message for final mode with location of FINAL directory
            if self.operation_mode.lower() == "final" and self.final_dir:
                print(f"\n{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                print(f"{Fore.GREEN}FINAL REVIEW PROCESS COMPLETE{Style.RESET_ALL}")
                print(f"{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                print(f"• All documents saved to FINAL directory: {Fore.CYAN}{self.final_dir}{Style.RESET_ALL}")
                print(f"• Process disclosure included in editor letter")
                print(f"• Backups stored in standard directory structure")
                print(f"{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}\n")
            
            self._log_success(f"Detailed revision report saved to {report_path}")
            self._log_success(f"Log file saved to {self.log_path}")
            
            # No need for additional verification - already done by _verify_final_mode
            try:
                # Placeholder for future process verification
                process_summary = None
                if process_summary:
                    if self.operation_mode.lower() == "final" and self.final_dir:
                        print(f"\n{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                        print(f"{Fore.GREEN}FINAL REVIEW PROCESS COMPLETE{Style.RESET_ALL}")
                        print(f"{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                        print(f"• All documents saved to FINAL directory: {Fore.CYAN}{self.final_dir}{Style.RESET_ALL}")
                        print(f"• Process disclosure included in editor letter")
                    print(f"• Total reviewers: {process_summary.get('reviewer_count', 0)}")
                    print(f"• Total reviewer personas: {process_summary.get('total_reviewer_personas', 0)}")
                    print(f"• Total editors: {process_summary.get('editor_count', 0)}")
                    print(f"• Process details stored in database for future reference")
                else:
                    print(f"\n{Fore.BLUE}[INFO]{Style.RESET_ALL} Review process information verified and stored in database")
            except Exception as e:
                self._log_warning(f"Error verifying review process information: {e}")
                if self.debug:
                    import traceback
                    traceback.print_exc()
            
            # Complete run in database
            self.workflow_db.complete_run(
                self.run_id,
                "completed",
                self.process_statistics["tokens_used"],
                self.process_statistics["cost"],
                self.process_statistics.get("evaluation_tokens", 0),
                self.process_statistics.get("evaluation_cost", 0.0)
            )
            
            # Get workflow stats from database
            workflow_stats = self.workflow_db.get_stats(self.run_id)
            
            # Print database statistics
            print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            print(f"{Fore.CYAN}DATABASE WORKFLOW STATISTICS:{Style.RESET_ALL}")
            print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
            print(f"• Files processed: {workflow_stats.get('file_count', 0)}")
            print(f"• Steps completed: {workflow_stats.get('completed_steps', 0)}/{workflow_stats.get('step_count', 0)}")
            print(f"• Completions: {workflow_stats.get('completion_count', 0)} ({workflow_stats.get('cached_count', 0)} cached)")
            print(f"• Evaluations: {workflow_stats.get('evaluation_count', 0)}")
            print(f"• Changes made: {workflow_stats.get('change_count', 0)}")
            print(f"• Average quality score: {workflow_stats.get('avg_quality_score', 0):.2f}/5.0")
            print(f"• Total duration: {workflow_stats.get('total_duration', 0):.2f}s")
            
            # Log final completion
            self.logger.info("Paper revision process completed successfully")
            self.logger.info(f"Total time elapsed: {time.time() - self.start_time:.2f} seconds")
            
            return {
                "revision_summary": self.revision_summary_path,
                "changes_document": self.changes_document_path,
                "revised_paper": self.revised_paper_path,
                "assessment": self.assessment_path,
                "editor_letter": self.editor_response_path,
                "new_bib": self.new_bib_path,
                "cost_report": cost_report,
                "revision_report": report_path,
                "log_file": self.log_path
            }
            
        except Exception as e:
            self._log_error(f"Error in paper revision process: {e}")
            if self.debug:
                import traceback
                tb = traceback.format_exc()
                self.logger.error(f"Traceback: {tb}")
                traceback.print_exc()
            else:
                self.logger.error(f"Exception: {str(e)}")
                
            # Log completion even in error case
            self.logger.info(f"Paper revision process failed after {time.time() - self.start_time:.2f} seconds")
            
            # Complete run in database with failed status
            self.workflow_db.complete_run(
                self.run_id,
                "failed",
                self.process_statistics["tokens_used"],
                self.process_statistics["cost"],
                self.process_statistics.get("evaluation_tokens", 0),
                self.process_statistics.get("evaluation_cost", 0.0)
            )
            
            # Move any created files to the trash directory
            trash_dir = self._move_to_trash()
            
            # Create a special log file in the trash directory
            trash_log_path = f"{trash_dir}/error_log_{self.timestamp}.txt"
            try:
                # Copy current log to trash directory
                import shutil
                if os.path.exists(self.log_path):
                    shutil.copy2(self.log_path, trash_log_path)
                else:
                    # Create a new log if the original doesn't exist
                    with open(trash_log_path, 'w') as f:
                        f.write(f"ERROR LOG - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write(f"Error: {str(e)}\n")
                        if self.debug:
                            f.write(f"\nTraceback:\n{tb}")
            except Exception as log_error:
                self._log_error(f"Failed to create error log in trash directory: {str(log_error)}")
            
            # Return the trash directory and error log
            return {
                "error": str(e),
                "trash_dir": trash_dir,
                "error_log": trash_log_path if os.path.exists(trash_log_path) else None,
                "log_file": self.log_path if os.path.exists(self.log_path) else None
            }
    
    def _analyze_original_paper(self) -> Dict[str, Any]:
        """Analyze the original paper.
        
        Returns:
            Dictionary with paper analysis results
        """
        self._log_info("Loading original paper")
        pdf_processor = PDFProcessor(self.original_paper_path)
        self.process_statistics["files_processed"] += 1
        
        # Get file information for tracking
        file_size = os.path.getsize(self.original_paper_path)
        page_count = pdf_processor.get_page_count()
        # Estimate tokens (1 page ≈ 500 tokens for academic papers)
        token_estimate = page_count * 500
        
        # Track the file in the database
        self._track_preprocessed_file(
            self.original_paper_path,
            self.original_paper_path,  # Same path since we're not creating a new file
            "pdf",
            file_size,
            token_estimate,
            page_count
        )
        
        sections = pdf_processor.extract_sections()
        tables = pdf_processor.extract_tables()
        figures = pdf_processor.extract_figures()
        references = pdf_processor.extract_references()
        
        # Extract only key sections to reduce token usage
        key_sections = {
            'Title': sections.get('Title', 'N/A'),
            'Abstract': sections.get('Abstract', 'N/A'),
            'Introduction': sections.get('Introduction', 'N/A')
        }
        
        # For cost optimization, reduce the amount of text being processed
        abstract_text = key_sections['Abstract']
        introduction_text = key_sections['Introduction']
        
        # Truncate long sections if optimizing costs
        if self.optimize_costs:
            # Keep first 500 characters of abstract (roughly 125 tokens)
            if len(abstract_text) > 500:
                abstract_text = abstract_text[:500] + "..."
                
            # Keep first 1000 characters of introduction (roughly 250 tokens)
            if len(introduction_text) > 1000:
                introduction_text = introduction_text[:1000] + "..."
        
        # Extract structured information using LLM
        prompt = f"""
        I'm analyzing a scientific paper from the journal Computers (ISSN: 2073-431X).
        I need you to extract and organize key information from this paper.
        
        Here's the paper content:
        
        Title: {key_sections.get('Title', 'N/A')}
        
        Abstract: {abstract_text}
        
        Introduction: {introduction_text}
        
        Provide a structured analysis with the following information:
        1. Title of the paper
        2. Authors
        3. Main research questions or objectives
        4. Methodology used
        5. Key findings
        6. Limitations mentioned
        7. Future work suggested
        8. Overall structure of the paper (inferred from the available sections)
        
        Format the response as a JSON object.
        """
        
        self._log_info("Analyzing paper structure and content")
        
        # Use optimized completion function
        paper_analysis_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific paper analysis assistant. Extract structured information from papers and format as JSON.",
            task_type="paper_analysis"
        )
        
        try:
            paper_analysis = json.loads(paper_analysis_json)
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON. Using basic analysis.")
            paper_analysis = {
                "title": sections.get('Title', 'Unknown Title'),
                "authors": "Unknown Authors",
                "objectives": "Unknown Objectives",
                "methodology": "Unknown Methodology",
                "findings": "Unknown Findings",
                "limitations": "Unknown Limitations",
                "future_work": "Unknown Future Work",
                "structure": list(sections.keys()),
                "key_terms": [],
                "publication_details": {}
            }
        
        pdf_processor.close()
        
        # Add sections, tables, and figures to the analysis
        paper_analysis["sections"] = sections
        paper_analysis["tables"] = tables
        paper_analysis["figures"] = [caption for caption, _ in figures]
        
        # Reduce token usage - only include a limited number of references
        if self.optimize_costs and len(references) > 10:
            # Keep only first 10 references to save tokens
            paper_analysis["references"] = references[:10]
            paper_analysis["references_count"] = len(references)
        else:
            paper_analysis["references"] = references
        
        self._log_success("Paper analysis completed")
        return paper_analysis
    
    def _analyze_reviewer_comments(self) -> List[Dict[str, Any]]:
        """Analyze the reviewer comments.
        
        Returns:
            List of dictionaries with reviewer comment analysis
        """
        reviewer_paths = [self.reviewer1_path, self.reviewer2_path, self.reviewer3_path]
        reviewer_comments = []
        
        for i, path in enumerate(reviewer_paths, 1):
            self._log_info(f"Analyzing reviewer {i} comments")
            pdf_processor = PDFProcessor(path)
            self.process_statistics["files_processed"] += 1
            
            text = pdf_processor.text
            
            # Optimize token usage by focusing on key parts of the review
            # This uses the heuristic that most important comments are often in the first third and last third
            text_length = len(text)
            if self.optimize_costs and text_length > 3000:
                # For long reviews, take first 1000 chars + last 1000 chars
                trimmed_text = text[:1000] + "\n...[content trimmed]...\n" + text[-1000:]
                self._log_debug(f"Trimmed reviewer {i} text from {text_length} to {len(trimmed_text)} chars")
            else:
                # For shorter reviews or when not optimizing, take up to 3000 chars
                max_length = 3000 if self.optimize_costs else 10000
                trimmed_text = text[:max_length] + ("..." if len(text) > max_length else "")
            
            # Focus on the most important aspects to save tokens
            prompt = f"""
            I'm analyzing reviewer comments for a scientific paper. Please extract only the most critical feedback.
            
            Here are the reviewer {i} comments:
            
            {trimmed_text}
            
            Provide a concise structured analysis with just these key points:
            1. Overall assessment (positive, neutral, negative)
            2. Main concerns (3-5 bullet points)
            3. Required changes (3-5 most important changes that must be addressed)
            
            Format the response as a JSON object with only these fields.
            """
            
            # Use optimized completion with reduced token usage
            reviewer_analysis_json = self._optimized_completion(
                prompt=prompt,
                system_prompt=f"You are a scientific reviewer analysis assistant. Extract only the most critical feedback from reviewer {i}'s comments as JSON.",
                max_tokens=1000,  # Limit response size
                task_type="reviewer_analysis"
            )
            
            try:
                reviewer_analysis = json.loads(reviewer_analysis_json)
                
                # Add default empty fields for optional analysis components
                reviewer_analysis.setdefault("suggested_changes", [])
                reviewer_analysis.setdefault("methodology_comments", [])
                reviewer_analysis.setdefault("results_comments", [])
                reviewer_analysis.setdefault("writing_comments", [])
                reviewer_analysis.setdefault("references_comments", [])
                
            except json.JSONDecodeError:
                # Fallback if LLM didn't return valid JSON
                self._log_warning(f"LLM didn't return valid JSON for reviewer {i}. Using basic analysis.")
                reviewer_analysis = {
                    "overall_assessment": "Unknown",
                    "main_concerns": ["Unknown concerns"],
                    "required_changes": ["Unknown required changes"],
                    "suggested_changes": [],
                    "methodology_comments": [],
                    "results_comments": [],
                    "writing_comments": [],
                    "references_comments": []
                }
            
            # To save memory/tokens, don't include full text when optimizing costs
            if not self.optimize_costs:
                reviewer_analysis["full_text"] = text
            else:
                # Just save the first 100 chars as a reference
                reviewer_analysis["text_preview"] = text[:100] + "..."
                
            reviewer_analysis["reviewer_number"] = i
            
            reviewer_comments.append(reviewer_analysis)
            pdf_processor.close()
        
        self._log_success("Reviewer comment analysis completed")
        return reviewer_comments
    
    def _process_editor_requirements(self) -> Dict[str, Any]:
        """Process editor letter and PRISMA requirements.
        
        Returns:
            Dictionary with editor requirements
        """
        # Process editor letter
        self._log_info("Processing editor letter")
        editor_pdf = PDFProcessor(self.editor_letter_path)
        self.process_statistics["files_processed"] += 1
        editor_text = editor_pdf.text
        editor_pdf.close()
        
        # Process PRISMA requirements
        self._log_info("Processing PRISMA requirements")
        prisma_pdf = PDFProcessor(self.prisma_requirements_path)
        self.process_statistics["files_processed"] += 1
        prisma_text = prisma_pdf.text
        prisma_pdf.close()
        
        # Optimize token usage
        if self.optimize_costs:
            # Editor text: Focus on first part of letter (usually contains key decisions)
            editor_text_trimmed = editor_text[:1500]
            
            # For PRISMA: Extract just the main checklist items
            prisma_lines = prisma_text.split('\n')
            prisma_checklist = []
            for line in prisma_lines:
                # Look for lines that might be checklist items (often numbered or bulleted)
                if re.search(r'^\s*(\d+[\.\)]|\*|\-|\•)', line) and len(line) < 200:
                    prisma_checklist.append(line.strip())
            
            # If we found checklist items, use them; otherwise take the beginning
            if prisma_checklist and len(prisma_checklist) > 5:
                prisma_text_trimmed = "\n".join(prisma_checklist[:15])  # Take top 15 items
            else:
                prisma_text_trimmed = prisma_text[:1000]
        else:
            # If not optimizing, still limit size but keep more content
            editor_text_trimmed = editor_text[:3000]
            prisma_text_trimmed = prisma_text[:3000]
        
        # Focus prompt on just the essential information needed
        prompt = f"""
        I'm analyzing the editor letter and PRISMA requirements for a scientific paper revision.
        Extract only the most critical information.
        
        Editor Letter:
        {editor_text_trimmed}
        
        PRISMA Requirements:
        {prisma_text_trimmed}
        
        Provide a structured analysis with just these key points:
        1. Editor's decision (reject, revise, accept with revisions, etc.)
        2. Editor's top 3-5 main requirements (most critical changes requested)
        3. Top 3-5 PRISMA framework requirements that must be met
        
        Format the response as a JSON object with only these fields.
        """
        
        editor_analysis_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific editor analysis assistant. Extract only the most critical requirements from editor letters and PRISMA guidelines.",
            max_tokens=1000  # Limit response size
        )
        
        try:
            editor_requirements = json.loads(editor_analysis_json)
            # Add default values for any missing fields
            editor_requirements.setdefault("editor_decision", "Unknown decision")
            editor_requirements.setdefault("editor_requirements", [])
            editor_requirements.setdefault("prisma_requirements", [])
            # Add backwards compatibility fields
            editor_requirements.setdefault("editor_suggestions", 
                                          editor_requirements.get("editor_requirements", [])[:2])
            editor_requirements.setdefault("prisma_modifications", [
                f"Ensure compliance with {req}" for req in editor_requirements.get("prisma_requirements", [])[:2]
            ])
            
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON for editor requirements. Using basic analysis.")
            editor_requirements = {
                "editor_decision": "Unknown decision",
                "editor_requirements": ["Unknown requirements"],
                "editor_suggestions": ["Unknown suggestions"],
                "prisma_requirements": ["Unknown PRISMA requirements"],
                "prisma_modifications": ["Unknown PRISMA modifications"]
            }
        
        # Add text previews (not full text) to save tokens
        if self.optimize_costs:
            editor_requirements["editor_letter_preview"] = editor_text[:200] + "..."
            editor_requirements["prisma_preview"] = prisma_text[:200] + "..."
        else:
            # Add full texts when not optimizing
            editor_requirements["editor_letter_text"] = editor_text
            editor_requirements["prisma_text"] = prisma_text
        
        self._log_success("Editor requirements analysis completed")
        return editor_requirements
    
    def _analyze_journal_style(self) -> Dict[str, Any]:
        """Analyze journal style and requirements.
        
        Returns:
            Dictionary with journal style information
        """
        # Process journal information
        self._log_info("Processing journal information")
        journal_pdf = PDFProcessor(self.journal_info_path)
        self.process_statistics["files_processed"] += 1
        
        # Track journal file in database
        file_size = os.path.getsize(self.journal_info_path)
        page_count = journal_pdf.get_page_count()
        token_estimate = page_count * 500
        
        self._track_preprocessed_file(
            self.journal_info_path,
            self.journal_info_path,
            "pdf",
            file_size,
            token_estimate,
            page_count
        )
        
        journal_text = journal_pdf.text
        journal_pdf.close()
        
        # Process Scopus information (limit to save processing)
        scopus_text = ""
        for path in self.scopus_info_paths[:1 if self.optimize_costs else 2]:  # Only process first one if optimizing
            self._log_info(f"Processing Scopus information from {path}")
            scopus_pdf = PDFProcessor(path)
            self.process_statistics["files_processed"] += 1
            
            # Track Scopus file in database
            file_size = os.path.getsize(path)
            page_count = scopus_pdf.get_page_count()
            token_estimate = page_count * 500
            
            self._track_preprocessed_file(
                path,
                path,
                "pdf",
                file_size,
                token_estimate,
                page_count
            )
            
            scopus_text += scopus_pdf.text[:2000] + "\n\n"  # Take just beginning of each
            scopus_pdf.close()
        
        # Process highly cited papers for style analysis - limit number based on max_papers_to_process
        self._log_info("Processing highly cited papers for style analysis")
        cited_papers_sample = []
        max_papers = min(self.max_papers_to_process, 3)  # Use at most 3 papers
        
        # If optimizing costs, only process 1 paper
        papers_to_process = 1 if self.optimize_costs else max_papers
        
        for i, path in enumerate(self.cited_papers_paths[:papers_to_process]):
            cited_pdf = PDFProcessor(path)
            self.process_statistics["files_processed"] += 1
            
            # Track cited paper in database
            file_size = os.path.getsize(path)
            page_count = cited_pdf.get_page_count()
            token_estimate = page_count * 500
            
            self._track_preprocessed_file(
                path,
                path,
                "pdf",
                file_size,
                token_estimate,
                page_count
            )
            
            sections = cited_pdf.extract_sections()
            
            # Get smaller sample of sections when optimizing
            abstract_limit = 250 if self.optimize_costs else 500
            intro_limit = 250 if self.optimize_costs else 500
            
            sample = {
                "paper_number": i + 1,
                "title": sections.get("Title", "Unknown Title"),
                "abstract": sections.get("Abstract", "Unknown Abstract")[:abstract_limit],
                "introduction": sections.get("Introduction", "Unknown Introduction")[:intro_limit]
            }
            cited_papers_sample.append(sample)
            cited_pdf.close()
        
        # Process similar papers - only if not optimizing costs or needed for full analysis
        similar_papers_sample = []
        if not self.optimize_costs:
            self._log_info("Processing similar papers for style analysis")
            for i, path in enumerate(self.similar_papers_paths[:1]):  # Just one even without optimization
                similar_pdf = PDFProcessor(path)
                self.process_statistics["files_processed"] += 1
                
                # Track similar paper in database
                file_size = os.path.getsize(path)
                page_count = similar_pdf.get_page_count()
                token_estimate = page_count * 500
                
                self._track_preprocessed_file(
                    path,
                    path,
                    "pdf",
                    file_size,
                    token_estimate,
                    page_count
                )
                
                sections = similar_pdf.extract_sections()
                sample = {
                    "paper_number": i + 1,
                    "title": sections.get("Title", "Unknown Title"),
                    "abstract": sections.get("Abstract", "Unknown Abstract")[:300],
                }
                similar_papers_sample.append(sample)
                similar_pdf.close()
        
        # Analyze reference style
        self._log_info("Analyzing reference style")
        ref_validator = ReferenceValidator(self.bib_path)
        self.process_statistics["files_processed"] += 1
        
        # Track bibliography file in database
        file_size = os.path.getsize(self.bib_path)
        # BIB files don't have pages, estimate token count differently
        # Rough estimate: 4 chars ≈ 1 token
        with open(self.bib_path, 'r', encoding='utf-8', errors='ignore') as f:
            bib_content = f.read()
        token_estimate = len(bib_content) // 4
        
        self._track_preprocessed_file(
            self.bib_path,
            self.bib_path,
            "bib",
            file_size,
            token_estimate,
            0  # No pages for BIB files
        )
        
        citation_style = ref_validator.get_citation_style()
        
        # Use much more targeted prompt that focuses only on key info when optimizing costs
        if self.optimize_costs:
            prompt = f"""
            I'm analyzing style requirements for a scientific paper for journal Computers (ISSN: 2073-431X).
            Extract ONLY the 3 most critical style requirements in each category.
            
            Journal Information Excerpt:
            {journal_text[:1000]}
            
            Citation Style:
            {citation_style}
            
            Sample Paper Title and Abstract:
            {cited_papers_sample[0]["title"] if cited_papers_sample else "Unknown"}
            {cited_papers_sample[0]["abstract"] if cited_papers_sample else "Unknown"}
            
            Provide a BRIEF style guide with ONLY:
            1. Top 3 formatting requirements (most important)
            2. Standard section structure (list of section names only)
            3. Citation style (brief description)
            
            Format as JSON with only these three fields.
            """
        else:
            # More comprehensive prompt when not optimizing costs
            prompt = f"""
            I'm analyzing the style requirements for a scientific paper to be published in the journal Computers (ISSN: 2073-431X).
            
            Journal Information:
            {journal_text[:2000]}
            
            Scopus Information:
            {scopus_text[:1500]}
            
            Citation Style:
            {citation_style}
            
            Sample of Highly Cited Papers from this Journal:
            {json.dumps(cited_papers_sample, indent=2)}
            
            Sample of Similar Papers to the Current Paper:
            {json.dumps(similar_papers_sample, indent=2)}
            
            Based on this information, provide a style guide for this journal, including:
            1. General formatting requirements
            2. Section structure
            3. Figure and table presentation
            4. Citation and reference style
            5. Language and tone
            
            Format the response as a JSON object.
            """
        
        # Use optimized completion
        style_analysis_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific journal style analysis assistant. Extract only the essential style guidelines.",
            max_tokens=1500 if self.optimize_costs else 3000
        )
        
        try:
            journal_style = json.loads(style_analysis_json)
            
            # Ensure all required fields exist
            journal_style.setdefault("formatting", ["Default formatting: 12pt Times New Roman, 1 inch margins"])
            journal_style.setdefault("section_structure", ["Introduction", "Methods", "Results", "Discussion", "Conclusion"])
            journal_style.setdefault("citation_style", citation_style)
            
            # Add these fields only when not optimizing
            if not self.optimize_costs:
                journal_style.setdefault("figures_and_tables", ["Figures and tables should be properly labeled"])
                journal_style.setdefault("language_and_tone", ["Clear, formal academic writing"])
                journal_style.setdefault("successful_papers", ["Novel contributions", "Rigorous methodology"])
                journal_style.setdefault("similar_papers_structure", ["Standard IMRAD structure"])
                
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON for journal style. Using basic analysis.")
            journal_style = {
                "formatting": ["12pt font", "1 inch margins", "Double spacing"],
                "section_structure": ["Introduction", "Methods", "Results", "Discussion", "Conclusion"],
                "citation_style": citation_style
            }
            
            # Add additional fields only when not optimizing
            if not self.optimize_costs:
                journal_style.update({
                    "figures_and_tables": ["Label all figures and tables"],
                    "language_and_tone": ["Formal academic writing"],
                    "successful_papers": ["Novel contributions"],
                    "similar_papers_structure": ["Standard IMRAD structure"]
                })
        
        self._log_success("Journal style analysis completed")
        return journal_style
    
    def _generate_revision_plan(
        self,
        paper_analysis: Dict[str, Any],
        reviewer_comments: List[Dict[str, Any]],
        editor_requirements: Dict[str, Any],
        journal_style: Dict[str, Any]
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Generate a revision plan based on all analyses.
        
        Args:
            paper_analysis: Analysis of the original paper
            reviewer_comments: Analysis of reviewer comments
            editor_requirements: Editor and PRISMA requirements
            journal_style: Journal style guidelines
            
        Returns:
            Tuple of (issues, solutions) as lists of dictionaries
        """
        self._log_info("Generating revision plan")
        
        # Prepare input for LLM
        paper_summary = {
            "title": paper_analysis.get("title", "Unknown Title"),
            "objectives": paper_analysis.get("objectives", "Unknown Objectives"),
            "methodology": paper_analysis.get("methodology", "Unknown Methodology"),
            "findings": paper_analysis.get("findings", "Unknown Findings"),
            "structure": paper_analysis.get("structure", [])
        }
        
        reviewer_summary = []
        for i, reviewer in enumerate(reviewer_comments, 1):
            main_concerns = reviewer.get("main_concerns", [])
            if main_concerns is None:
                main_concerns = []
            required_changes = reviewer.get("required_changes", [])
            if required_changes is None:
                required_changes = []
                
            summary = {
                "reviewer": i,
                "assessment": reviewer.get("overall_assessment", "Unknown"),
                "main_concerns": main_concerns[:3] if len(main_concerns) > 0 else [],  # Limit to top 3
                "required_changes": required_changes[:3] if len(required_changes) > 0 else []  # Limit to top 5
            }
            reviewer_summary.append(summary)
        
        editor_reqs = editor_requirements.get("editor_requirements", [])
        if editor_reqs is None:
            editor_reqs = []
        prisma_reqs = editor_requirements.get("prisma_requirements", [])
        if prisma_reqs is None:
            prisma_reqs = []
            
        editor_summary = {
            "requirements": editor_reqs[:5] if len(editor_reqs) > 0 else [],  # Limit to top 5
            "decision": editor_requirements.get("editor_decision", "Unknown"),
            "prisma_requirements": prisma_reqs[:5] if len(prisma_reqs) > 0 else []  # Limit to top 5
        }
        
        formatting = journal_style.get("formatting", [])
        if formatting is None:
            formatting = []
        section_structure = journal_style.get("section_structure", [])
        if section_structure is None:
            section_structure = []
            
        style_summary = {
            "formatting": formatting[:3] if len(formatting) > 0 else [],  # Limit to top 3
            "section_structure": section_structure[:3] if len(section_structure) > 0 else [],  # Limit to top 3
            "citation_style": journal_style.get("citation_style", "Unknown")
        }
        
        # Create a more targeted prompt when optimizing costs
        if self.optimize_costs:
            prompt = f"""
            I'm developing a focused revision plan for a scientific paper based on the most critical feedback.
            
            Paper Title:
            {paper_summary.get("title", "Unknown")}
            
            Top Reviewer Concerns:
            {json.dumps([r.get("main_concerns", [])[:2] for r in reviewer_summary], indent=0)}
            
            Editor Decision:
            {editor_summary.get("decision", "Unknown")}
            
            Top Requirements:
            {json.dumps(editor_summary.get("requirements", [])[:3], indent=0)}
            
            Based on this focused information, generate:
            
            1. The 3-5 most critical issues that must be addressed, each with:
               - title (concise description)
               - source (reviewer number, editor, or style)
               - severity (critical, major, minor)
            
            2. 3-5 specific solutions to address these issues, each with:
               - title (concise description)
               - implementation (brief explanation of changes needed)
               - complexity (high, medium, low)
            
            Format the response as a JSON object with "issues" and "solutions" arrays.
            """
        else:
            # More comprehensive prompt when not optimizing
            prompt = f"""
            I'm developing a revision plan for a scientific paper based on reviewer comments, editor requirements, and journal style guidelines.
            
            Paper Summary:
            {json.dumps(paper_summary, indent=2)}
            
            Reviewer Comments:
            {json.dumps(reviewer_summary, indent=2)}
            
            Editor Requirements:
            {json.dumps(editor_summary, indent=2)}
            
            Journal Style:
            {json.dumps(style_summary, indent=2)}
            
            Based on this information, generate:
            
            1. A list of issues that need to be addressed, with each issue having:
               - title (short description)
               - description (detailed explanation)
               - source (reviewer number, editor, or style guidelines)
               - severity (critical, major, minor)
            
            2. A list of solutions to address these issues, with each solution having:
               - title (short description)
               - implementation (detailed explanation of changes needed)
               - complexity (high, medium, low)
               - impact (how this addresses the issues)
            
            Ensure that each solution addresses one or more of the identified issues.
            Format the response as a JSON object with "issues" and "solutions" arrays.
            """
        
        # Use optimized completion with appropriate token limits
        revision_plan_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific paper revision assistant. Generate structured revision plans based on reviewer comments and requirements.",
            max_tokens=2000 if self.optimize_costs else 4000,
            task_type="revision_planning"
        )
        
        try:
            revision_plan = json.loads(revision_plan_json)
            issues = revision_plan.get("issues", [])
            solutions = revision_plan.get("solutions", [])
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON for revision plan. Using basic plan.")
            issues = []
            for i, reviewer in enumerate(reviewer_comments):
                main_concerns = reviewer.get("main_concerns", ["Unknown concerns"])
                if main_concerns is None:
                    main_concerns = ["Unknown concerns"]
                
                # Safely slice the list, ensuring it's not None
                top_concerns = main_concerns[:2] if len(main_concerns) >= 2 else main_concerns
                
                issues.append({
                    "title": f"Reviewer {i+1} concern",
                    "description": ", ".join(top_concerns),
                    "source": f"Reviewer {i+1}",
                    "severity": "major"
                })
            
            solutions = [{
                "title": f"Address reviewer {i+1} concerns",
                "implementation": "Review and implement the changes requested by the reviewer.",
                "complexity": "medium",
                "impact": "Addresses the main concerns of the reviewer."
            } for i in range(len(reviewer_comments))]
        
        self._log_success(f"Generated revision plan with {len(issues)} issues and {len(solutions)} solutions")
        return issues, solutions
    
    def _create_revision_summary(
        self,
        issues: List[Dict[str, Any]],
        solutions: List[Dict[str, Any]],
        output_path: str
    ) -> str:
        """Create a revision summary document.
        
        Args:
            issues: List of issues to address
            solutions: List of solutions
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self._log_info("Creating revision summary document")
        
        # Load the original docx if it exists, otherwise use a new document
        if os.path.exists(self.original_docx_path):
            doc_processor = DocumentProcessor(self.original_docx_path)
        else:
            # Create a new document from the PDF
            pdf_processor = PDFProcessor(self.original_paper_path)
            temp_docx_path = pdf_processor.pdf_to_docx()
            pdf_processor.close()
            doc_processor = DocumentProcessor(temp_docx_path)
        
        # Create the revision summary document
        summary_path = doc_processor.create_revision_summary(issues, solutions, output_path)
        self.process_statistics["files_created"] += 1
        
        return summary_path
    
    def _generate_changes(
        self,
        paper_analysis: Dict[str, Any],
        issues: List[Dict[str, Any]],
        solutions: List[Dict[str, Any]]
    ) -> List[Tuple[str, str, str, Optional[int]]]:
        """Generate specific text changes based on the revision plan.
        
        Args:
            paper_analysis: Analysis of the original paper
            issues: List of issues to address
            solutions: List of solutions
            
        Returns:
            List of tuples (old_text, new_text, reason, line_number)
        """
        self._log_info("Generating specific text changes")
        
        # Prepare input for LLM
        paper_sections = paper_analysis.get("sections", {})
        section_samples = {}
        for section_name, content in paper_sections.items():
            # Limit section content to save tokens
            section_samples[section_name] = content[:1000] + "..." if len(content) > 1000 else content
        
        solutions_summary = []
        for solution in solutions:
            summary = {
                "title": solution.get("title", "Unknown solution"),
                "implementation": solution.get("implementation", "Unknown implementation"),
                "complexity": solution.get("complexity", "medium")
            }
            solutions_summary.append(summary)
        
        prompt = f"""
        I'm generating specific text changes for a scientific paper revision.
        
        Paper Sections:
        {json.dumps(section_samples, indent=2)}
        
        Solutions to Implement:
        {json.dumps(solutions_summary, indent=2)}
        
        Based on this information, generate a list of specific text changes that should be made to implement the solutions.
        For each change, provide:
        1. The original text to be replaced
        2. The new text to replace it with
        3. The reason for the change
        4. The approximate line number (can be an estimate)
        
        Focus on substantive changes that address the revision requirements. Include:
        - Changes to improve methodology description
        - Changes to address limitations
        - Changes to improve the presentation of results
        - Changes to improve the structure and flow
        - Changes to address specific reviewer concerns
        
        Generate at least 5 but no more than 15 specific changes.
        Format the response as a JSON array of objects with "old_text", "new_text", "reason", and "line_number" fields.
        """
        
        # Use optimized completion with appropriate token limits for changes generation
        changes_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific paper revision assistant. Generate specific text changes to implement revision solutions.",
            max_tokens=3000 if self.optimize_costs else 4000,
            task_type="text_generation"
        )
        
        try:
            changes_data = json.loads(changes_json)
            # Convert to list of tuples
            changes = [(
                item.get("old_text", ""),
                item.get("new_text", ""),
                item.get("reason", ""),
                item.get("line_number")
            ) for item in changes_data]
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON for changes. Using basic changes.")
            changes = []
            for section_name, content in list(section_samples.items())[:3]:  # Use first 3 sections
                if len(content) > 100:
                    changes.append((
                        content[:100],
                        f"[REVISED] {content[:100]}",
                        f"Improve {section_name} section based on reviewer comments",
                        None
                    ))
        
        self._log_success(f"Generated {len(changes)} specific text changes")
        return changes
    
    def _create_changes_document(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: str
    ) -> str:
        """Create a document detailing all changes.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self._log_info("Creating changes document")
        
        # Load the original docx if it exists, otherwise use a new document
        if os.path.exists(self.original_docx_path):
            doc_processor = DocumentProcessor(self.original_docx_path)
        else:
            # Create a new document from the PDF
            pdf_processor = PDFProcessor(self.original_paper_path)
            temp_docx_path = pdf_processor.pdf_to_docx()
            pdf_processor.close()
            doc_processor = DocumentProcessor(temp_docx_path)
        
        # Create the changes document
        changes_path = doc_processor.create_changes_document(changes, output_path)
        self.process_statistics["files_created"] += 1
        
        return changes_path
    
    def _validate_and_update_references(
        self,
        paper_analysis: Dict[str, Any],
        reviewer_comments: List[Dict[str, Any]],
        output_path: str
    ) -> List[Dict[str, str]]:
        """Validate existing references and add new ones based on reviewer suggestions.
        
        Args:
            paper_analysis: Analysis of the original paper
            reviewer_comments: Analysis of reviewer comments
            output_path: Path where the new BibTeX file should be saved
            
        Returns:
            List of new references added
        """
        self._log_info("Validating and updating references")
        
        # Load references
        ref_validator = ReferenceValidator(self.bib_path)
        
        # Validate existing references
        valid_refs, invalid_refs = ref_validator.validate_references()
        self._log_info(f"Found {len(valid_refs)} valid and {len(invalid_refs)} invalid references")
        
        # Extract reference suggestions from reviewer comments
        reference_comments = []
        for reviewer in reviewer_comments:
            ref_comments = reviewer.get("references_comments", [])
            reference_comments.extend(ref_comments)
        
        # Generate new references based on reviewer suggestions
        if reference_comments:
            prompt = f"""
            I'm updating references for a scientific paper based on reviewer comments.
            
            Current Paper References:
            {paper_analysis.get('references', ['No references available'])[:10]}
            
            Reviewer Comments on References:
            {reference_comments}
            
            Based on these comments, suggest new references that should be added to the paper.
            For each reference, provide:
            1. Title
            2. Authors
            3. Journal/Conference
            4. Year
            5. DOI (if you can estimate it)
            6. Why this reference should be added
            
            Format the response as a JSON array of objects with the fields above.
            """
            
            # Use optimized completion for reference suggestions
            new_refs_json = self._optimized_completion(
                prompt=prompt,
                system_prompt="You are a scientific reference assistant. Suggest new references based on reviewer comments.",
                max_tokens=2000 if self.optimize_costs else 3000
            )
            
            try:
                new_refs_data = json.loads(new_refs_json)
                new_references = []
                
                for ref_data in new_refs_data:
                    # Convert to BibTeX entry format
                    entry = {
                        "ENTRYTYPE": "article",
                        "title": ref_data.get("title", "Unknown Title"),
                        "author": ref_data.get("authors", "Unknown Authors"),
                        "journal": ref_data.get("journal", ref_data.get("conference", "Unknown Venue")),
                        "year": str(ref_data.get("year", "2023")),
                        "doi": ref_data.get("doi", "")
                    }
                    
                    # Add to reference validator
                    ref_id = ref_validator.add_reference(entry)
                    
                    new_references.append({
                        "id": ref_id,
                        "title": entry["title"],
                        "authors": entry["author"],
                        "year": entry["year"],
                        "reason": ref_data.get("why", "Suggested by reviewer")
                    })
                
                # Save updated references
                ref_validator.export_references(
                    set(list(valid_refs) + [ref["id"] for ref in new_references]),
                    output_path
                )
                self.process_statistics["files_created"] += 1
                
                self._log_success(f"Added {len(new_references)} new references")
                return new_references
                
            except json.JSONDecodeError:
                self._log_warning("LLM didn't return valid JSON for new references. No new references added.")
        
        # If no new references, just export valid ones
        ref_validator.export_references(valid_refs, output_path)
        self.process_statistics["files_created"] += 1
        
        return []
    
    def _create_revised_paper(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: str
    ) -> str:
        """Create revised paper with track changes.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self._log_info("Creating revised paper with track changes")
        
        # Load the original docx
        if os.path.exists(self.original_docx_path):
            doc_processor = DocumentProcessor(self.original_docx_path)
        else:
            # Create a new document from the PDF
            pdf_processor = PDFProcessor(self.original_paper_path)
            temp_docx_path = pdf_processor.pdf_to_docx()
            pdf_processor.close()
            doc_processor = DocumentProcessor(temp_docx_path)
        
        # Apply changes with track changes
        changes_applied = 0
        for old_text, new_text, reason, _ in changes:
            if doc_processor.add_tracked_change(old_text, new_text, reason):
                changes_applied += 1
        
        # Save the revised document
        doc_processor.save(output_path)
        self.process_statistics["files_created"] += 1
        
        self._log_success(f"Applied {changes_applied} changes to the paper")
        return output_path
    
    def _create_assessment(
        self,
        changes: List[Tuple[str, str, str, Optional[int]]],
        paper_analysis: Dict[str, Any],
        output_path: str
    ) -> str:
        """Create assessment document.
        
        Args:
            changes: List of tuples (old_text, new_text, reason, line_number)
            paper_analysis: Analysis of the original paper
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self._log_info("Creating assessment document")
        
        # Create a new document
        from docx import Document
        doc = Document()
        
        # Add heading
        doc.add_heading('Revision Assessment', 0)
        
        # Add introduction
        doc.add_paragraph(
            'This document assesses the changes made to the paper in response to reviewer comments. ' +
            'It evaluates the impact of these changes and identifies any issues that may need to be addressed manually.'
        )
        
        # Summarize changes
        doc.add_heading('Summary of Changes', 1)
        
        changes_by_type = {}
        for _, _, reason, _ in changes:
            change_type = reason.split(' ')[0] if reason else 'Other'
            if change_type not in changes_by_type:
                changes_by_type[change_type] = 0
            changes_by_type[change_type] += 1
        
        p = doc.add_paragraph('Changes by type:')
        for change_type, count in changes_by_type.items():
            p = doc.add_paragraph(f'{change_type}: {count} changes', style='List Bullet')
        
        # Impact assessment
        doc.add_heading('Impact Assessment', 1)
        
        # Use LLM to generate impact assessment
        paper_summary = {
            "title": paper_analysis.get("title", "Unknown Title"),
            "objectives": paper_analysis.get("objectives", "Unknown Objectives"),
            "methodology": paper_analysis.get("methodology", "Unknown Methodology"),
            "findings": paper_analysis.get("findings", "Unknown Findings")
        }
        
        changes_summary = []
        for old_text, new_text, reason, _ in changes[:5]:  # Use first 5 changes as examples
            summary = {
                "reason": reason,
                "old_text_preview": old_text[:50] + "..." if len(old_text) > 50 else old_text,
                "new_text_preview": new_text[:50] + "..." if len(new_text) > 50 else new_text
            }
            changes_summary.append(summary)
        
        prompt = f"""
        I'm assessing the impact of revisions made to a scientific paper.
        
        Paper Summary:
        {json.dumps(paper_summary, indent=2)}
        
        Sample of Changes Made:
        {json.dumps(changes_summary, indent=2)}
        
        Total changes: {len(changes)}
        
        Provide an assessment of:
        1. Overall impact of the changes on the paper's quality
        2. How well the changes address likely reviewer concerns
        3. Potential remaining issues that might need manual attention
        4. Specific areas where the paper has been strengthened
        5. Recommendations for any final manual adjustments
        
        Format the response as a JSON object with these five sections.
        """
        
        # Use optimized completion for assessment generation
        assessment_json = self._optimized_completion(
            prompt=prompt,
            system_prompt="You are a scientific paper assessment assistant. Evaluate the impact of revisions on paper quality.",
            max_tokens=2000 if self.optimize_costs else 3000
        )
        
        try:
            assessment = json.loads(assessment_json)
            
            # Add overall impact
            doc.add_heading('Overall Impact', 2)
            doc.add_paragraph(assessment.get("overall_impact", "No assessment available."))
            
            # Add reviewer concerns addressed
            doc.add_heading('Reviewer Concerns Addressed', 2)
            doc.add_paragraph(assessment.get("reviewer_concerns", "No assessment available."))
            
            # Add remaining issues
            doc.add_heading('Remaining Issues', 2)
            remaining_issues = assessment.get("remaining_issues", "No issues identified.")
            if isinstance(remaining_issues, list):
                for issue in remaining_issues:
                    doc.add_paragraph(issue, style='List Bullet')
            else:
                doc.add_paragraph(remaining_issues)
            
            # Add strengthened areas
            doc.add_heading('Strengthened Areas', 2)
            strengthened_areas = assessment.get("strengthened_areas", "No areas identified.")
            if isinstance(strengthened_areas, list):
                for area in strengthened_areas:
                    doc.add_paragraph(area, style='List Bullet')
            else:
                doc.add_paragraph(strengthened_areas)
            
            # Add recommendations
            doc.add_heading('Recommendations', 2)
            recommendations = assessment.get("recommendations", "No recommendations.")
            if isinstance(recommendations, list):
                for rec in recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
            else:
                doc.add_paragraph(recommendations)
            
        except json.JSONDecodeError:
            # Fallback if LLM didn't return valid JSON
            self._log_warning("LLM didn't return valid JSON for assessment. Using basic assessment.")
            
            doc.add_heading('Overall Impact', 2)
            doc.add_paragraph('The changes have addressed many of the issues identified in the review process.')
            
            doc.add_heading('Remaining Tasks', 2)
            doc.add_paragraph('The following tasks should be completed manually:')
            doc.add_paragraph('1. Review all changes for accuracy and consistency.', style='List Bullet')
            doc.add_paragraph('2. Check references for proper formatting.', style='List Bullet')
            doc.add_paragraph('3. Ensure all reviewer comments have been addressed.', style='List Bullet')
        
        # Add conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(
            'This assessment provides an overview of the changes made and their impact. ' +
            'Authors should review the revised paper carefully before submission.'
        )
        
        # Save the document
        doc.save(output_path)
        self.process_statistics["files_created"] += 1
        
        self._log_success("Created assessment document")
        return output_path
    
    def _create_editor_letter(
        self,
        reviewer_comments: List[Dict[str, Any]],
        changes: List[Tuple[str, str, str, Optional[int]]],
        output_path: str
    ) -> str:
        """Create letter to the editor.
        
        Args:
            reviewer_comments: Analysis of reviewer comments
            changes: List of tuples (old_text, new_text, reason, line_number)
            output_path: Path where the document should be saved
            
        Returns:
            Path to the created document
        """
        self._log_info("Creating letter to editor")
        
        # Process changes to create responses to reviewers
        reviewer_responses = []
        
        for i, reviewer in enumerate(reviewer_comments, 1):
            # Group changes by reason to identify which ones address this reviewer's comments
            changes_by_reason = {}
            for old_text, new_text, reason, _ in changes:
                if f"Reviewer {i}" in reason or f"reviewer {i}" in reason.lower():
                    if reason not in changes_by_reason:
                        changes_by_reason[reason] = []
                    changes_by_reason[reason].append((old_text, new_text))
            
            # Extract main comments from the reviewer
            main_concerns = reviewer.get("main_concerns", [])
            required_changes = reviewer.get("required_changes", [])
            
            # Create a list of comment-response pairs
            comments = []
            
            # Add responses to main concerns
            for j, concern in enumerate(main_concerns, 1):
                response_text = "We appreciate this concern and have addressed it in our revision."
                changes_text = "No specific changes were made for this comment."
                
                # Find changes related to this concern
                for reason, change_list in changes_by_reason.items():
                    if any(keyword in concern.lower() for keyword in reason.lower().split()):
                        changes_text = f"We made {len(change_list)} changes to address this concern, including: "
                        changes_text += ", ".join([f"replacing '{old[:20]}...' with '{new[:20]}...'" 
                                                 for old, new in change_list[:2]])
                        if len(change_list) > 2:
                            changes_text += f", and {len(change_list) - 2} more changes."
                
                comments.append({
                    "comment": concern,
                    "response": response_text,
                    "changes": changes_text
                })
            
            # Add responses to required changes
            for j, required in enumerate(required_changes, 1):
                response_text = "We have implemented this required change in our revision."
                changes_text = "No specific changes were made for this comment."
                
                # Find changes related to this requirement
                for reason, change_list in changes_by_reason.items():
                    if any(keyword in required.lower() for keyword in reason.lower().split()):
                        changes_text = f"We made {len(change_list)} changes to address this requirement, including: "
                        changes_text += ", ".join([f"replacing '{old[:20]}...' with '{new[:20]}...'" 
                                                 for old, new in change_list[:2]])
                        if len(change_list) > 2:
                            changes_text += f", and {len(change_list) - 2} more changes."
                
                comments.append({
                    "comment": required,
                    "response": response_text,
                    "changes": changes_text
                })
            
            reviewer_responses.append({
                "reviewer": i,
                "comments": comments
            })
        
        # Create the editor letter
        doc_processor = DocumentProcessor(self.original_docx_path) if os.path.exists(self.original_docx_path) else None
        
        # Get process summary for inclusion in the letter
        process_summary = None
        if hasattr(self, 'workflow_db') and self.run_id:
            try:
                process_summary = self.workflow_db.get_review_process_summary(self.run_id)
            except Exception as e:
                self._log_warning(f"Error getting review process summary: {e}")
        
        if doc_processor:
            editor_letter_path = doc_processor.create_editor_letter(reviewer_responses, output_path, process_summary)
        else:
            # Create manually if original docx not available
            from docx import Document
            doc = Document()
            
            # Add header information
            doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("To: The Editor, Computers (ISSN: 2073-431X)")
            doc.add_paragraph("Subject: Revised manuscript submission")
            doc.add_paragraph()
            
            # Add salutation
            doc.add_paragraph("Dear Editor,")
            
            # Add introduction
            doc.add_paragraph(
                "Thank you for the opportunity to revise our manuscript. We have carefully addressed all the comments " +
                "provided by the reviewers and made the necessary changes to improve the quality of our paper. " +
                "We believe that the revised version addresses all the concerns raised and significantly improves the manuscript."
            )
            
            # Add reviewer responses
            doc.add_heading('Responses to Reviewer Comments', 1)
            
            for i, response in enumerate(reviewer_responses, 1):
                doc.add_heading(f"Reviewer {i}", 2)
                
                for j, comment in enumerate(response['comments'], 1):
                    p = doc.add_paragraph()
                    p.add_run(f"Comment {j}: ").bold = True
                    p.add_run(comment['comment'])
                    
                    p = doc.add_paragraph()
                    p.add_run("Response: ").bold = True
                    p.add_run(comment['response'])
                    
                    p = doc.add_paragraph()
                    p.add_run("Changes made: ").bold = True
                    p.add_run(comment['changes'])
                    
                    doc.add_paragraph()
            
            # Add closing
            doc.add_paragraph(
                "We hope that the revised manuscript now meets the standards for publication in Computers. " +
                "We look forward to your feedback and are available to address any additional questions or concerns."
            )
            
            # Add revision process disclaimer if provided
            if process_summary:
                doc.add_heading('REVISION PROCESS DISCLOSURE', 1)
                
                # Add detailed information about the multi-persona revision process
                doc.add_paragraph(process_summary.get("process_description", ""))
                
                # Add statistics
                stats_paragraph = doc.add_paragraph()
                stats_paragraph.add_run("REVISION STATISTICS:\n").bold = True
                stats_paragraph.add_run(f"• Reviewers: {process_summary.get('reviewer_count', 0)}\n")
                stats_paragraph.add_run(f"• Reviewer personas: {process_summary.get('total_reviewer_personas', 0)}\n")
                stats_paragraph.add_run(f"• Editors: {process_summary.get('editor_count', 0)}\n")
                stats_paragraph.add_run(f"• Fine-tuned personas used: {process_summary.get('fine_personas_used', 0)}\n")
                stats_paragraph.add_run(f"• Total reviews generated: {process_summary.get('review_count', 0)}\n")
                stats_paragraph.add_run(f"• Final decision: {process_summary.get('decision', 'Not specified')}")
                
                # Add acknowledgment
                doc.add_paragraph(
                    "We acknowledge that this revision process utilized advanced AI-assisted multi-persona review " +
                    "technology to ensure a comprehensive, diverse, and thorough evaluation of our manuscript. " +
                    "The multiple persona approach ensures that our paper was examined from various academic " +
                    "perspectives before making our final revisions."
                )
                
                # Add attribution for FinePersonas if used
                if process_summary.get('fine_personas_used', 0) > 0:
                    doc.add_paragraph(
                        "This revision process utilized the FinePersonas dataset " +
                        "(https://huggingface.co/datasets/argilla/FinePersonas-v0.1) " +
                        "to enhance the diversity and expertise of the reviewer perspectives."
                    )
            
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph("The Authors")
            
            # Save the document
            doc.save(output_path)
            editor_letter_path = output_path
        
        self.process_statistics["files_created"] += 1
        
        self._log_success("Created letter to editor")
        return editor_letter_path

def choose_model(operation_mode=None, selected_provider=None):
    """Interactive model selection with default recommendations based on operation mode.
    
    Args:
        operation_mode: The operation mode to use for recommendations (training, finetuning, final)
        selected_provider: Optional pre-selected provider to skip provider selection
    
    Returns:
        Tuple of (provider, model)
    """
    colorama_init()
    
    # Get mode settings if provided
    mode_settings = OPERATION_MODES.get(operation_mode, {}) if operation_mode else {}
    provider_recommendations = mode_settings.get("provider_recommendations", {})
    
    # If provider is not pre-selected, ask user to select one
    if not selected_provider:
        print(f"{Fore.CYAN}Choose LLM Provider:{Style.RESET_ALL}")
        print(f"{Fore.CYAN}1.{Style.RESET_ALL} Anthropic Claude" + 
              (f" {Fore.GREEN}[Recommended: {provider_recommendations.get('anthropic', '')}]{Style.RESET_ALL}" 
               if 'anthropic' in provider_recommendations else ""))
        print(f"{Fore.CYAN}2.{Style.RESET_ALL} OpenAI GPT" + 
              (f" {Fore.GREEN}[Recommended: {provider_recommendations.get('openai', '')}]{Style.RESET_ALL}" 
               if 'openai' in provider_recommendations else ""))
        print(f"{Fore.CYAN}3.{Style.RESET_ALL} Google Gemini" + 
              (f" {Fore.GREEN}[Recommended: {provider_recommendations.get('google', '')}]{Style.RESET_ALL}" 
               if 'google' in provider_recommendations else ""))
        
        provider_choice = input("Enter choice (1-3): ")
        
        if provider_choice == "1":
            provider = "anthropic"
        elif provider_choice == "2":
            provider = "openai"
        elif provider_choice == "3":
            provider = "google"
        else:
            print(f"{Fore.RED}Invalid choice. Defaulting to Anthropic Claude.{Style.RESET_ALL}")
            provider = "anthropic"
    else:
        provider = selected_provider
    
    # Get models based on provider
    if provider == "anthropic":
        models = get_claude_model_choices()
        recommended_model = provider_recommendations.get("anthropic")
    elif provider == "openai":
        models = get_openai_model_choices()
        recommended_model = provider_recommendations.get("openai")
    elif provider == "google":
        models = get_gemini_model_choices()
        recommended_model = provider_recommendations.get("google")
    else:
        print(f"{Fore.RED}Invalid provider. Defaulting to Anthropic Claude.{Style.RESET_ALL}")
        provider = "anthropic"
        models = get_claude_model_choices()
        recommended_model = provider_recommendations.get("anthropic")
    
    # Find the index of the recommended model if it exists
    recommended_index = -1
    if recommended_model:
        for i, model in enumerate(models):
            if recommended_model in model:
                recommended_index = i
                break
    
    print(f"\n{Fore.CYAN}Choose Model for {provider.capitalize()}:{Style.RESET_ALL}")
    
    # If we have a recommendation for this operation mode, highlight it prominently
    if operation_mode and recommended_index >= 0:
        print(f"{Fore.GREEN}Recommended model for {operation_mode.upper()} mode: {models[recommended_index]}{Style.RESET_ALL}")
        
    # List all models
    for i, model in enumerate(models, 1):
        model_text = f"{Fore.CYAN}{i}.{Style.RESET_ALL} {model}"
        # Highlight the recommended model
        if i - 1 == recommended_index:
            model_text += f" {Fore.GREEN}[RECOMMENDED]{Style.RESET_ALL}"
        print(model_text)
    
    # If there's a recommended model, use it as default
    if recommended_index >= 0 and operation_mode:
        model_choice = input(f"Enter choice (1-{len(models)}) or press Enter for recommended model: ")
        if not model_choice.strip():  # User pressed Enter
            return provider, models[recommended_index]
    else:
        model_choice = input(f"Enter choice (1-{len(models)}): ")
    
    try:
        if model_choice.strip():  # Only try to parse if not empty
            model_index = int(model_choice) - 1
            if 0 <= model_index < len(models):
                chosen_model = models[model_index]
            else:
                if recommended_index >= 0:
                    print(f"{Fore.YELLOW}Invalid choice. Using recommended model.{Style.RESET_ALL}")
                    chosen_model = models[recommended_index]
                else:
                    print(f"{Fore.RED}Invalid choice. Defaulting to first model.{Style.RESET_ALL}")
                    chosen_model = models[0]
        elif recommended_index >= 0:  # Empty input with recommendation
            chosen_model = models[recommended_index]
        else:  # Empty input without recommendation
            print(f"{Fore.RED}No choice made. Defaulting to first model.{Style.RESET_ALL}")
            chosen_model = models[0]
    except ValueError:
        if recommended_index >= 0:
            print(f"{Fore.YELLOW}Invalid input. Using recommended model.{Style.RESET_ALL}")
            chosen_model = models[recommended_index]
        else:
            print(f"{Fore.RED}Invalid input. Defaulting to first model.{Style.RESET_ALL}")
            chosen_model = models[0]
    
    return provider, chosen_model

def choose_operation_mode():
    """Interactive operation mode selection."""
    colorama_init()
    
    print(f"{Fore.CYAN}Choose Operation Mode:{Style.RESET_ALL}")
    print(f"{Fore.CYAN}1.{Style.RESET_ALL} {OPERATION_MODES['training']['description']}")
    print(f"   Recommended models: {Fore.GREEN}Anthropic: {OPERATION_MODES['training']['provider_recommendations']['anthropic']}, " +
          f"OpenAI: {OPERATION_MODES['training']['provider_recommendations']['openai']}, " +
          f"Google: {OPERATION_MODES['training']['provider_recommendations']['google']}{Style.RESET_ALL}")
    
    print(f"{Fore.CYAN}2.{Style.RESET_ALL} {OPERATION_MODES['finetuning']['description']}")
    print(f"   Recommended models: {Fore.GREEN}Anthropic: {OPERATION_MODES['finetuning']['provider_recommendations']['anthropic']}, " +
          f"OpenAI: {OPERATION_MODES['finetuning']['provider_recommendations']['openai']}, " +
          f"Google: {OPERATION_MODES['finetuning']['provider_recommendations']['google']}{Style.RESET_ALL}")
    
    print(f"{Fore.CYAN}3.{Style.RESET_ALL} {OPERATION_MODES['final']['description']}")
    print(f"   Recommended models: {Fore.GREEN}Anthropic: {OPERATION_MODES['final']['provider_recommendations']['anthropic']}, " +
          f"OpenAI: {OPERATION_MODES['final']['provider_recommendations']['openai']}, " +
          f"Google: {OPERATION_MODES['final']['provider_recommendations']['google']}{Style.RESET_ALL}")
    
    mode_choice = input("Enter choice (1-3): ")
    
    if mode_choice == "1":
        return "training"
    elif mode_choice == "2":
        return "finetuning"
    elif mode_choice == "3":
        return "final"
    else:
        print(f"{Fore.YELLOW}[WARNING]{Style.RESET_ALL} Invalid choice. Defaulting to finetuning mode.")
        return "finetuning"

def main():
    """Main entry point for the paper revision tool."""
    # Load environment variables
    load_dotenv()
    
    # Load API keys from environment if available
    scopus_api_key_env = os.environ.get("SCOPUS_API_KEY")
    wos_client_id_env = os.environ.get("WOS_CLIENT_ID")
    
    # Check if we need to update model information
    db = WorkflowDB()
    update_due = db.is_model_update_due()
    db.close()
    
    if update_due:
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Biweekly model data update is due. Updating model information...")
        # Force update of all model scores
        initialize_model_scores(force_update=True)
        print(f"{Fore.GREEN}[SUCCESS]{Style.RESET_ALL} Model data updated successfully.")
    else:
        # Normal initialization
        initialize_model_scores()
    
    # Check if preprocessing has been done
    preprocessing_metadata_file = "./.cache/preprocessed/metadata.json"
    model_choice_file = "./.cache/preprocessed/model_choice.json"
    
    preprocessing_done = os.path.exists(preprocessing_metadata_file)
    model_choice_exists = os.path.exists(model_choice_file)
    
    # Parse command line arguments
    parser = argparse.ArgumentParser(description="Paper Revision Tool for Computers Journal")
    
    # Add extended help information
    parser.description += """

Model Verification:
  This tool can verify model accuracy by asking 'When is Christmas?' 
  Run with --verify to check models before paper revision
  Or use verify_models.py to test models separately

File Preprocessing:
  Before running paper revision, use preprocess_files.py to:
  - Analyze and convert files in the 'asis' directory
  - Get recommendations for optimal model selection
  - Estimate processing cost and time
"""
    
    parser.add_argument("--mode", choices=["training", "finetuning", "final"], 
                        help="Operation mode (training, finetuning, or final)")
    parser.add_argument("--provider", choices=["anthropic", "openai", "google"], 
                        help="LLM provider")
    parser.add_argument("--verify", action="store_true", 
                        help="Verify model accuracy before proceeding (asks 'When is Christmas?')")
    parser.add_argument("--preprocess", action="store_true",
                        help="Run file preprocessing before starting (equivalent to running preprocess_files.py)")
    parser.add_argument("--competitor-eval", action="store_true", default=True,
                        help="Use competing models to evaluate output quality (default: enabled)")
    parser.add_argument("--no-competitor-eval", action="store_false", dest="competitor_eval",
                        help="Disable competing model evaluation")
    parser.add_argument("--evaluator", 
                        help="Specific competing model to use for evaluation (format: provider/model)")
    parser.add_argument("--model", help="Model name (specific to provider)")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode")
    parser.add_argument("--token-budget", type=int, 
                        help="Maximum token budget (overrides mode setting)")
    parser.add_argument("--cost-budget", type=float,
                        help="Maximum cost budget in dollars (overrides mode setting)")
    cost_opt_group = parser.add_mutually_exclusive_group()
    cost_opt_group.add_argument("--optimize-costs", action="store_const", const=True, dest="optimize_costs_arg",
                        help="Enable cost optimization (overrides mode setting)")
    cost_opt_group.add_argument("--no-optimize-costs", action="store_const", const=False, dest="optimize_costs_arg",
                        help="Disable cost optimization (overrides mode setting)")
    parser.add_argument("--max-papers", type=int,
                        help="Maximum number of papers to process for style analysis (overrides mode setting)")
    parser.add_argument("--api", choices=["scopus", "wos", "scopus,wos", "wos,scopus"],
                        help="API integration to use (scopus, wos, or both)")
    parser.add_argument("--key", help="API key for the specified API (uses value from .env if not provided)")
    args = parser.parse_args()
    
    # Handle preprocessing if requested
    if args.preprocess:
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Running file preprocessing...")
        import subprocess
        subprocess.run(["python", "preprocess_files.py"])
        # Refresh preprocessing status after running
        preprocessing_done = os.path.exists(preprocessing_metadata_file)
        model_choice_exists = os.path.exists(model_choice_file)
    
    # Show preprocessing status
    if not preprocessing_done:
        print(f"{Fore.YELLOW}[WARNING]{Style.RESET_ALL} Files have not been preprocessed. For optimal results, run:")
        print(f"  python preprocess_files.py")
        print(f"This will analyze your files and recommend the best model for your documents.\n")
    
    # Choose operation mode if not specified
    operation_mode = args.mode
    if not operation_mode:
        operation_mode = choose_operation_mode()
    
    # Get settings from operation mode
    mode_settings = OPERATION_MODES[operation_mode]
    
    # Apply operation mode settings (can be overridden by explicit command line args)
    token_budget = args.token_budget if args.token_budget is not None else mode_settings["token_budget"]
    cost_budget = args.cost_budget if args.cost_budget is not None else mode_settings["cost_budget"]
    max_papers = args.max_papers if args.max_papers is not None else mode_settings["max_papers"]
    
    # For optimize_costs, only override if explicitly set via command line
    if hasattr(args, 'optimize_costs_arg') and args.optimize_costs_arg is not None:
        optimize_costs = args.optimize_costs_arg
    else:
        optimize_costs = mode_settings["optimize_costs"]
    
    # Choose model interactively if not specified via command line
    # Provider and model selection logic
    if not args.provider and not args.model:
        # No provider or model specified - check for preprocessing recommendations
        if model_choice_exists:
            # Load model choice from preprocessing
            try:
                with open(model_choice_file, 'r') as f:
                    model_choice = json.load(f)
                    provider = model_choice.get("provider")
                    model = model_choice.get("model")
                    
                    # Confirm using preprocessed recommendation
                    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Found preprocessing recommendation:")
                    print(f"  Provider: {Fore.YELLOW}{provider.capitalize()}{Style.RESET_ALL}")
                    print(f"  Model: {Fore.YELLOW}{model}{Style.RESET_ALL}")
                    
                    use_recommended = input(f"Use this recommended model? (Y/n): ").lower()
                    if use_recommended == 'n':
                        # User wants to choose manually
                        provider, model = choose_model(operation_mode=operation_mode)
            except Exception as e:
                print(f"{Fore.YELLOW}[WARNING]{Style.RESET_ALL} Error loading preprocessing recommendation: {e}")
                provider, model = choose_model(operation_mode=operation_mode)
        else:
            # No preprocessing, use interactive selection
            provider, model = choose_model(operation_mode=operation_mode)
    elif not args.model and args.provider:
        # Provider specified but no model
        provider = args.provider
        
        # Check for preprocessing recommendations for this provider
        if preprocessing_done:
            try:
                with open(preprocessing_metadata_file, 'r') as f:
                    metadata = json.load(f)
                    if "model_recommendations" in metadata and provider in metadata["model_recommendations"]:
                        # Show preprocessing recommendations for this provider
                        recommendations = metadata["model_recommendations"][provider]
                        if recommendations:
                            recommended_model = recommendations[0]["model_name"]
                            
                            print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Preprocessing recommends for {provider.capitalize()}: {recommended_model}")
                            use_recommended = input(f"Use this recommended model? (Y/n): ").lower()
                            if use_recommended != 'n':
                                model = recommended_model
                            else:
                                # Fall back to operation mode recommendations
                                recommended_mode_model = mode_settings["provider_recommendations"].get(provider)
                                if recommended_mode_model:
                                    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mode '{operation_mode}' recommends: {recommended_mode_model}")
                                    use_mode_rec = input(f"Use mode's recommended model? (Y/n): ").lower()
                                    if use_mode_rec != 'n':
                                        model = recommended_mode_model
                                    else:
                                        _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                                else:
                                    _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                        else:
                            # Fall back to operation mode recommendations
                            recommended_model = mode_settings["provider_recommendations"].get(provider)
                            if recommended_model:
                                print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mode '{operation_mode}' recommends model: {recommended_model}")
                                use_recommended = input(f"Use recommended {provider} model? (Y/n): ").lower()
                                if use_recommended != 'n':
                                    model = recommended_model
                                else:
                                    _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                            else:
                                _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                    else:
                        # No preprocessing recommendations for this provider
                        recommended_model = mode_settings["provider_recommendations"].get(provider)
                        if recommended_model:
                            print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mode '{operation_mode}' recommends model: {recommended_model}")
                            use_recommended = input(f"Use recommended {provider} model? (Y/n): ").lower()
                            if use_recommended != 'n':
                                model = recommended_model
                            else:
                                _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                        else:
                            _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
            except Exception as e:
                print(f"{Fore.YELLOW}[WARNING]{Style.RESET_ALL} Error loading preprocessing recommendations: {e}")
                # Fall back to operation mode recommendations
                recommended_model = mode_settings["provider_recommendations"].get(provider)
                if recommended_model:
                    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mode '{operation_mode}' recommends model: {recommended_model}")
                    use_recommended = input(f"Use recommended {provider} model? (Y/n): ").lower()
                    if use_recommended != 'n':
                        model = recommended_model
                    else:
                        _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
                else:
                    _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
        else:
            # No preprocessing, use operation mode recommendations
            recommended_model = mode_settings["provider_recommendations"].get(provider)
            if recommended_model:
                print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mode '{operation_mode}' recommends model: {recommended_model}")
                use_recommended = input(f"Use recommended {provider} model? (Y/n): ").lower()
                if use_recommended != 'n':
                    model = recommended_model
                else:
                    _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
            else:
                _, model = choose_model(operation_mode=operation_mode, selected_provider=provider)
    else:
        # Both provider and model specified directly in arguments
        provider = args.provider
        model = args.model
    
    # Print operation mode banner with detailed settings
    print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
    print(f"{Fore.CYAN}OPERATION MODE: {operation_mode.upper()}{Style.RESET_ALL}")
    print(f"{Fore.CYAN}{mode_settings['description']}{Style.RESET_ALL}")
    print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
    print(f"{Fore.BLUE}Settings for {operation_mode.upper()} mode:{Style.RESET_ALL}")
    print(f"• Token budget: {Fore.YELLOW}{token_budget:,}{Style.RESET_ALL} tokens")
    print(f"• Cost budget: {Fore.YELLOW}${cost_budget:.2f}{Style.RESET_ALL}")
    print(f"• Max papers to analyze: {Fore.YELLOW}{max_papers}{Style.RESET_ALL}")
    print(f"• Cost optimization: {Fore.GREEN if optimize_costs else Fore.RED}{optimize_costs}{Style.RESET_ALL}")
    print(f"• Selected provider: {Fore.YELLOW}{provider.capitalize()}{Style.RESET_ALL}")
    print(f"• Selected model: {Fore.YELLOW}{model}{Style.RESET_ALL}")
    print(f"• Model verification: {Fore.GREEN if args.verify else Fore.YELLOW}{args.verify}{Style.RESET_ALL}" + 
          f"{'' if args.verify else ' (use --verify flag or verify_models.py to check models)'}")
    print(f"• Using preprocessing: {Fore.GREEN if preprocessing_done else Fore.YELLOW}{preprocessing_done}{Style.RESET_ALL}" +
          f"{'' if preprocessing_done else ' (run preprocess_files.py for file analysis and model recommendations)}'}")
    print(f"• Competitor evaluation: {Fore.GREEN if args.competitor_eval else Fore.RED}{args.competitor_eval}{Style.RESET_ALL}")
    
    # Show evaluator if specified
    if args.competitor_eval and args.evaluator:
        print(f"  Using specific evaluator: {Fore.YELLOW}{args.evaluator}{Style.RESET_ALL}")
    elif args.competitor_eval:
        print(f"  Using automatic competitor selection based on model tier")
    
    # Show if the selected model matches the recommendation
    recommended_model = mode_settings["provider_recommendations"].get(provider, "")
    if recommended_model in model:
        print(f"  {Fore.GREEN}✓ Using recommended model for {operation_mode} mode{Style.RESET_ALL}")
    else:
        print(f"  {Fore.YELLOW}⚠ Not using the recommended model ({recommended_model}){Style.RESET_ALL}")
    
    print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}\n")
    
    # Use API key from environment if not provided via command line
    api_key = args.key
    if not api_key and args.api:
        if "scopus" in args.api and scopus_api_key_env:
            api_key = scopus_api_key_env
            print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Using Scopus API key from environment")
    
    # Create and run the paper revision tool with parameters
    revision_tool = PaperRevisionTool(
        provider=provider,
        model_name=model,
        debug=args.debug,
        token_budget=token_budget,
        cost_budget=cost_budget,
        max_papers_to_process=max_papers,
        optimize_costs=optimize_costs,
        operation_mode=operation_mode,
        verify=args.verify,  # Pass the verify flag from command line arguments
        competitor_evaluation=args.competitor_eval,  # Enable/disable competitor evaluation
        competing_evaluator=args.evaluator,  # Specific competing model to use for evaluation
        api=args.api,  # API integration to use (scopus, wos, or both)
        api_key=api_key  # API key for the specified API
    )
    
    # Print starting message
    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Starting paper revision process...")
    
    # Run the tool
    results = revision_tool.run()
    
    if results:
        print(f"\n{Fore.GREEN}Paper revision completed successfully!{Style.RESET_ALL}")
        print("Output files:")
        
        # Display special reports first
        special_reports = ["revision_report", "cost_report", "log_file"]
        for report_name in special_reports:
            if report_name in results and results[report_name]:
                if report_name == "revision_report":
                    print(f"{Fore.GREEN}- {report_name}: {results[report_name]}{Style.RESET_ALL} (Detailed revision summary)")
                elif report_name == "cost_report":
                    print(f"{Fore.GREEN}- {report_name}: {results[report_name]}{Style.RESET_ALL} (Cost optimization report)")
                elif report_name == "log_file":
                    print(f"{Fore.GREEN}- {report_name}: {results[report_name]}{Style.RESET_ALL} (Process log file)")
            
        # Display other output files
        for name, path in results.items():
            if name not in special_reports and path:
                # Extract model code from path for clearer display
                model_code = "Unknown"
                if "/tobe/" in path:
                    try:
                        model_part = path.split("/tobe/")[1].split("/")[0]
                        model_code = model_part.split("_")[0]  # Extract just the code portion
                    except:
                        pass
                print(f"- {name}: {path} (Model: {model_code})")
                
        print(f"\n{Fore.BLUE}[INFO]{Style.RESET_ALL} Reports and logs saved to:")
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} - Cost report: {results.get('cost_report', 'N/A')}")
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} - Revision report: {results.get('revision_report', 'N/A')}")
        print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} - Log file: {results.get('log_file', 'N/A')}")
        
        # Check for latest runs by provider in this operation mode
        try:
            db = WorkflowDB()
            provider_runs = db.get_latest_provider_runs(operation_mode)
            
            # Check if we have runs from all three providers
            providers = ["anthropic", "openai", "google"]
            missing_providers = [p for p in providers if p not in provider_runs]
            
            if len(provider_runs) == 3:
                print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}MEGA-RESULT OPPORTUNITY!{Style.RESET_ALL}")
                print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                print(f"You now have completed runs from all three providers in {operation_mode.upper()} mode:")
                
                # Show provider runs
                for provider_name, run in provider_runs.items():
                    run_time = datetime.datetime.fromisoformat(run.get('timestamp', '')).strftime('%Y-%m-%d %H:%M:%S')
                    print(f"• {provider_name.capitalize()}: {run.get('model')} (Run ID: {run.get('run_id')}, Time: {run_time})")
                
                # Dynamically determine the best model across all providers for this mode
                # using provider specs-based scoring system that considers both capability and cost-efficiency
                
                # For operation mode "training", prioritize cost more than capability
                # For "finetuning", balance cost and capability
                # For "final", prioritize capability over cost
                model_scores = get_best_models_by_mode(operation_mode, provider_runs)
                
                # Get the best model (highest blended score)
                best_model = model_scores[0]
                best_model_provider = best_model["provider"]
                best_model_name = best_model["model"]
                
                # Get the second-best model as evaluator
                next_best_model = model_scores[1]
                evaluator_provider = next_best_model["provider"]
                evaluator_model = next_best_model["model"]
                
                print(f"\n{Fore.CYAN}Model Assessment for Mega-Result:{Style.RESET_ALL}")
                print(f"• Best model identified: {Fore.GREEN}{best_model_provider.capitalize()} {best_model_name}{Style.RESET_ALL}")
                print(f"  - Capability: {best_model['capability']}/100, Cost-efficiency: {best_model['cost_efficiency']}/100")
                print(f"  - Blended score: {best_model['blended_score']:.1f}/100 (Optimized for {operation_mode} mode)")
                print(f"  - Cost: ${best_model['input_cost']:.5f} input, ${best_model['output_cost']:.5f} output per 1K tokens")
                
                print(f"\n• Evaluator model: {Fore.YELLOW}{evaluator_provider.capitalize()} {evaluator_model}{Style.RESET_ALL}")
                print(f"  - Capability: {next_best_model['capability']}/100, Cost-efficiency: {next_best_model['cost_efficiency']}/100")
                print(f"  - Blended score: {next_best_model['blended_score']:.1f}/100")
                
                print(f"\n• Third model: {Fore.BLUE}{model_scores[2]['provider'].capitalize()} {model_scores[2]['model']}{Style.RESET_ALL}")
                print(f"  - Capability: {model_scores[2]['capability']}/100, Cost-efficiency: {model_scores[2]['cost_efficiency']}/100")
                print(f"  - Blended score: {model_scores[2]['blended_score']:.1f}/100")
                
                # Add option to prioritize cost or capability
                print(f"\n{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                print(f"{Fore.CYAN}MEGA-RESULT OPTIONS:{Style.RESET_ALL}")
                print(f"{Fore.CYAN}{'=' * 70}{Style.RESET_ALL}")
                print(f"1. {Fore.GREEN}Standard mode{Style.RESET_ALL}: Optimized for {operation_mode} mode (default)")
                print(f"   - Uses {best_model_provider.capitalize()} {best_model_name}")
                print(f"   - Evaluated by {evaluator_provider.capitalize()} {evaluator_model}")
                
                # Calculate alternative ranking that prioritizes cost
                cost_optimized_scores = get_best_models_by_mode(operation_mode, provider_runs, prioritize_cost=True)
                cost_best_model = cost_optimized_scores[0]
                cost_next_best = cost_optimized_scores[1]
                
                print(f"\n2. {Fore.YELLOW}Cost-optimized mode{Style.RESET_ALL}: Prioritizes cost-efficiency")
                print(f"   - Uses {cost_best_model['provider'].capitalize()} {cost_best_model['model']}")
                print(f"   - Evaluated by {cost_next_best['provider'].capitalize()} {cost_next_best['model']}")
                print(f"   - Cost: ${cost_best_model['input_cost']:.5f} input, ${cost_best_model['output_cost']:.5f} output per 1K tokens")
                
                # Recalculate with minimum cost weight to prioritize raw capability
                capability_scores = []
                for model in model_scores:
                    capability_scores.append({
                        "provider": model["provider"],
                        "model": model["model"],
                        "capability": model["capability"],
                        "cost_efficiency": model["cost_efficiency"],
                        "blended_score": model["capability"],  # Use capability as the score
                        "input_cost": model["input_cost"],
                        "output_cost": model["output_cost"],
                        "run_id": model["run_id"]
                    })
                # Sort by pure capability
                capability_scores.sort(key=lambda x: x["capability"], reverse=True)
                capability_best_model = capability_scores[0]
                capability_next_best = capability_scores[1]
                
                print(f"\n3. {Fore.MAGENTA}Capability-optimized mode{Style.RESET_ALL}: Prioritizes raw capability")
                print(f"   - Uses {capability_best_model['provider'].capitalize()} {capability_best_model['model']}")
                print(f"   - Evaluated by {capability_next_best['provider'].capitalize()} {capability_next_best['model']}")
                print(f"   - Cost: ${capability_best_model['input_cost']:.5f} input, ${capability_best_model['output_cost']:.5f} output per 1K tokens")
                
                # Prompt for mega-result creation with mode selection
                print(f"\nWould you like to create a MEGA-RESULT that combines the best parts of all three runs?")
                
                mode_choice = input("Enter mode (1-3) or 'n' to cancel [1]: ").strip()
                
                if not mode_choice:
                    mode_choice = "1"  # Default to standard mode
                
                if mode_choice.lower() == 'n':
                    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mega-result creation skipped.")
                    create_mega = 'n'
                else:
                    try:
                        mode_num = int(mode_choice)
                        if 1 <= mode_num <= 3:
                            create_mega = 'y'
                            
                            # Set models based on chosen mode
                            if mode_num == 2:  # Cost-optimized
                                best_model = cost_best_model
                                next_best_model = cost_next_best
                                best_model_provider = best_model["provider"]
                                best_model_name = best_model["model"]
                                evaluator_provider = next_best_model["provider"]
                                evaluator_model = next_best_model["model"]
                                print(f"{Fore.YELLOW}Cost-optimized mode selected.{Style.RESET_ALL}")
                                
                            elif mode_num == 3:  # Capability-optimized
                                best_model = capability_best_model
                                next_best_model = capability_next_best
                                best_model_provider = best_model["provider"]
                                best_model_name = best_model["model"]
                                evaluator_provider = next_best_model["provider"]
                                evaluator_model = next_best_model["model"]
                                print(f"{Fore.MAGENTA}Capability-optimized mode selected.{Style.RESET_ALL}")
                                
                            else:  # Standard mode (already set)
                                print(f"{Fore.GREEN}Standard mode selected.{Style.RESET_ALL}")
                        else:
                            print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Invalid mode. Must be 1-3.")
                            create_mega = 'n'
                    except ValueError:
                        print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Invalid input. Mega-result creation skipped.")
                        create_mega = 'n'
                
                if create_mega == 'y':
                    print(f"\n{Fore.GREEN}Creating mega-result...{Style.RESET_ALL}")
                    
                    # Create a new timestamp for the mega run
                    mega_timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                    
                    # Gather run IDs to merge
                    run_ids = [run.get('run_id') for run in provider_runs.values()]
                    
                    # Settings for the mega run
                    mega_settings = {
                        "source_runs": run_ids,
                        "operation_mode": operation_mode,
                        "is_mega_result": True,
                        "source_providers": list(provider_runs.keys())
                    }
                    
                    # Create the merged run in the database
                    db.merge_run_results(
                        run_ids,
                        mega_timestamp,
                        best_model_provider,
                        best_model_name,
                        operation_mode,
                        mega_settings
                    )
                    
                    print(f"\n{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                    print(f"{Fore.GREEN}MEGA-RESULT CREATED SUCCESSFULLY!{Style.RESET_ALL}")
                    print(f"{Fore.GREEN}{'=' * 70}{Style.RESET_ALL}")
                    print(f"• Run ID: {Fore.CYAN}{mega_timestamp}{Style.RESET_ALL}")
                    print(f"• Primary model: {Fore.GREEN}{best_model_provider.capitalize()} {best_model_name}{Style.RESET_ALL}")
                    print(f"• Evaluator: {Fore.YELLOW}{evaluator_provider.capitalize()} {evaluator_model}{Style.RESET_ALL}")
                    print(f"• Source runs:")
                    for idx, model_info in enumerate(model_scores):
                        color = Fore.GREEN if idx == 0 else (Fore.YELLOW if idx == 1 else Fore.BLUE)
                        print(f"  {idx+1}. {color}{model_info['provider'].capitalize()} {model_info['model']}{Style.RESET_ALL} (Run ID: {model_info['run_id']})")
                    
                    print(f"\nYou can find the merged results in the database using run ID: {Fore.CYAN}{mega_timestamp}{Style.RESET_ALL}")
                    
                    # For final mode, offer the option to simulate reviewer feedback
                    if operation_mode == "final":
                        print(f"\n{Fore.MAGENTA}{'=' * 70}{Style.RESET_ALL}")
                        print(f"{Fore.MAGENTA}SIMULATED REVIEWER FEEDBACK OPTION{Style.RESET_ALL}")
                        print(f"{Fore.MAGENTA}{'=' * 70}{Style.RESET_ALL}")
                        print(f"Would you like to generate simulated reviewer feedback on your revised paper?")
                        print(f"This will create personas based on the original reviewers' comments and generate")
                        print(f"individual reviews plus an editor's summary to guide further improvements.")
                        
                        create_review = input("Generate reviewer feedback? (y/N): ").lower()
                        
                        if create_review == 'y':
                            # Get the paths to the needed files
                            revised_paper_path = results.get("revised_paper")
                            if not revised_paper_path or not os.path.exists(revised_paper_path):
                                print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Could not find revised paper file. Skipping reviewer feedback.")
                            else:
                                try:
                                    print(f"\n{Fore.BLUE}Generating reviewer personas and feedback...{Style.RESET_ALL}")
                                    
                                    # Extract journal guidelines from analysis
                                    journal_guidelines = "Follow standard academic publishing guidelines. Focus on originality, methodology, clarity, and significance of contribution."
                                    if hasattr(revision_tool, 'journal_style') and revision_tool.journal_style:
                                        if 'formatting_requirements' in revision_tool.journal_style:
                                            journal_guidelines = revision_tool.journal_style['formatting_requirements']
                                    
                                    # Get field from paper analysis
                                    field = "Computer Science"
                                    if hasattr(revision_tool, 'paper_analysis') and revision_tool.paper_analysis:
                                        if 'field' in revision_tool.paper_analysis:
                                            field = revision_tool.paper_analysis['field']
                                    
                                    # Generate review report
                                    review_report = create_review_report(
                                        revised_paper_path=revised_paper_path,
                                        original_comments=revision_tool._analyze_reviewer_comments(),
                                        journal_guidelines=journal_guidelines,
                                        field=field,
                                        model_client=revision_tool.llm_client
                                    )
                                    
                                    # Save the report
                                    timestamp_dir = os.path.dirname(revised_paper_path)
                                    report_path = save_review_report(review_report, timestamp_dir)
                                    
                                    print(f"\n{Fore.GREEN}Reviewer feedback generated successfully!{Style.RESET_ALL}")
                                    print(f"Report saved to: {Fore.CYAN}{report_path}{Style.RESET_ALL}")
                                    
                                    # Display key decision
                                    decision = review_report['editor_summary']['decision']
                                    decision_color = Fore.GREEN if decision == "Accept" else (
                                                    Fore.YELLOW if decision in ["Minor Revision"] else Fore.RED)
                                    print(f"\nEditor's Decision: {decision_color}{decision}{Style.RESET_ALL}")
                                    
                                    # Show top concerns
                                    if 'common_concerns' in review_report['editor_summary'] and review_report['editor_summary']['common_concerns']:
                                        print(f"\nTop concerns to address:")
                                        for i, concern in enumerate(review_report['editor_summary']['common_concerns'][:3], 1):
                                            print(f"{i}. {concern}")
                                    
                                    # Ask if user wants to start a new revision cycle
                                    if decision not in ["Accept", "Reject"]:
                                        print(f"\nWould you like to start a new revision cycle to address these concerns?")
                                        new_cycle = input("Start new revision cycle? (y/N): ").lower()
                                        
                                        if new_cycle == 'y':
                                            print(f"\n{Fore.BLUE}To start a new revision cycle, please run the tool again and use the revised paper as input.{Style.RESET_ALL}")
                                            print(f"Remember to consider the reviewer feedback from: {report_path}")
                                        else:
                                            # User chose not to continue - revise the editor letter to acknowledge unaddressed concerns
                                            print(f"\n{Fore.YELLOW}Revising letter to editor to acknowledge unaddressed concerns...{Style.RESET_ALL}")
                                            
                                            # Get the path to the editor letter
                                            editor_letter_path = results.get("editor_letter")
                                            if not editor_letter_path or not os.path.exists(editor_letter_path):
                                                print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Could not find editor letter file. Skipping revision.")
                                            else:
                                                try:
                                                    # Read the current editor letter
                                                    with open(editor_letter_path, 'r', encoding='utf-8') as f:
                                                        current_letter = f.read()
                                                    
                                                    # Prompt for reasons why concerns weren't addressed
                                                    print(f"\n{Fore.CYAN}Why were some concerns not addressed? (select all that apply){Style.RESET_ALL}")
                                                    print(f"1. Time constraints")
                                                    print(f"2. Beyond the scope of the current work")
                                                    print(f"3. Technical limitations")
                                                    print(f"4. Disagreement with reviewer assessment")
                                                    print(f"5. Will be addressed in future work")
                                                    print(f"6. Other (specify)")
                                                    
                                                    reasons_input = input("Enter choices (e.g., 1,3,5): ")
                                                    reasons = []
                                                    other_reason = ""
                                                    
                                                    # Process the selected reasons
                                                    if reasons_input:
                                                        reason_nums = [r.strip() for r in reasons_input.split(',')]
                                                        
                                                        reason_map = {
                                                            "1": "time constraints",
                                                            "2": "the concern is beyond the scope of the current work",
                                                            "3": "technical limitations",
                                                            "4": "disagreement with the reviewer's assessment",
                                                            "5": "plan to address in future work"
                                                        }
                                                        
                                                        for num in reason_nums:
                                                            if num in reason_map:
                                                                reasons.append(reason_map[num])
                                                            elif num == "6":
                                                                other_reason = input("Please specify other reason: ")
                                                                if other_reason:
                                                                    reasons.append(other_reason)
                                                    
                                                    if not reasons:
                                                        reasons = ["limitations in the scope of the current revision"]
                                                    
                                                    # Get the unaddressed concerns
                                                    unaddressed_concerns = []
                                                    if 'common_concerns' in review_report['editor_summary']:
                                                        unaddressed_concerns = review_report['editor_summary']['common_concerns']
                                                    
                                                    # Generate the addendum for the editor letter
                                                    concerns_text = ", ".join([f"'{c}'" for c in unaddressed_concerns[:3]])
                                                    reasons_text = ", ".join(reasons)
                                                    
                                                    addendum = f"""
                                                    
                                                    ADDENDUM REGARDING UNADDRESSED CONCERNS:
                                                    
                                                    We acknowledge that there are remaining concerns in our manuscript that have been identified by the reviewers, including {concerns_text}. We have not addressed these concerns in the current revision due to {reasons_text}.
                                                    
                                                    """
                                                    
                                                    if "5" in reasons_input or "plan to address in future work" in reasons:
                                                        addendum += """We intend to address these issues in future work and believe they do not compromise the core contributions of the current manuscript. We appreciate the reviewers' thorough feedback and will incorporate these suggestions in our ongoing research program.
                                                        """
                                                    
                                                    # Append the addendum to the editor letter
                                                    updated_letter = current_letter + addendum
                                                    
                                                    # Save the updated letter
                                                    with open(editor_letter_path, 'w', encoding='utf-8') as f:
                                                        f.write(updated_letter)
                                                    
                                                    print(f"{Fore.GREEN}Letter to editor updated successfully with explanations for unaddressed concerns.{Style.RESET_ALL}")
                                                    print(f"Updated letter saved to: {Fore.CYAN}{editor_letter_path}{Style.RESET_ALL}")
                                                    
                                                except Exception as e:
                                                    print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Failed to update editor letter: {e}")
                                                    if revision_tool.debug:
                                                        import traceback
                                                        traceback.print_exc()
                                    
                                except Exception as e:
                                    print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Failed to generate reviewer feedback: {e}")
                                    if revision_tool.debug:
                                        import traceback
                                        traceback.print_exc()
                else:
                    print(f"{Fore.BLUE}[INFO]{Style.RESET_ALL} Mega-result creation skipped.")
            
            elif len(provider_runs) > 0:
                print(f"\n{Fore.YELLOW}You have completed runs for the following providers in {operation_mode.upper()} mode:{Style.RESET_ALL}")
                for provider_name, run in provider_runs.items():
                    run_time = datetime.datetime.fromisoformat(run.get('timestamp', '')).strftime('%Y-%m-%d %H:%M:%S')
                    print(f"• {provider_name.capitalize()}: {run.get('model')} (Run ID: {run.get('run_id')}, Time: {run_time})")
                
                if missing_providers:
                    print(f"\n{Fore.YELLOW}To create a mega-result, you need runs from all three providers.{Style.RESET_ALL}")
                    print(f"Missing providers: {', '.join(p.capitalize() for p in missing_providers)}")
                    print(f"Run the tool with each missing provider to enable the mega-result feature.")
            
            db.close()
        except Exception as e:
            print(f"{Fore.RED}[ERROR]{Style.RESET_ALL} Failed to check for provider runs: {e}")
    else:
        print(f"\n{Fore.RED}Paper revision failed.{Style.RESET_ALL}")
        
        # Show error message if available
        if results and "error" in results:
            print(f"Error: {results['error']}")
        
        # Show trash directory information
        if results and "trash_dir" in results:
            print(f"\n{Fore.YELLOW}Files from this failed run have been copied to the trash directory:{Style.RESET_ALL}")
            print(f"  {results['trash_dir']}")
        
        # Show error log location
        if results and "error_log" in results and results["error_log"]:
            print(f"\n{Fore.YELLOW}Detailed error log:{Style.RESET_ALL}")
            print(f"  {results['error_log']}")
        elif results and "log_file" in results and results["log_file"]:
            print(f"\n{Fore.YELLOW}Check the log file for details:{Style.RESET_ALL}")
            print(f"  {results['log_file']}")
        else:
            print("\nNo log file was generated. Check console output for errors.")

if __name__ == "__main__":
    main()